"""
terminal/tools.py — Read-only tool implementations for the NSE Agent Adda.

Each tool returns a plain dict (JSON-serialisable).  Tools must NOT mutate any
data, execute shell commands, or access the network beyond approved sources.
"""

from __future__ import annotations

import json
import os
import re
import sqlite3
import warnings
from difflib import SequenceMatcher
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Any

import pandas as pd

# ── F&O data and options analysis ────────────────────────────────────────────
from terminal.fno_data import (
    fetch_live_option_chain,
    fetch_live_futures,
    load_and_store_latest as _fno_load_latest,
    get_available_dates as _fno_available_dates,
    days_to_expiry,
    get_expiry_dates,
    get_lot_size,
)
from terminal.options_analysis import (
    analyze_option_chain,
    analyze_futures,
    build_strategy,
    recommend_strategies,
    STRATEGY_CATALOG,
    analyze_buying_opportunity,
    scan_options_buying_opportunities,
    calc_expected_move,
    theta_decay_profile,
    calc_chain_ivs,
)

# ── Web research module (screener.in, Yahoo Finance, multi-site search) ───────
from terminal.web_research import (
    scrape_screener_in,
    search_yahoo_finance,
    multi_source_web_search,
    comprehensive_stock_research,
)

# ── Deep Search Engine ────────────────────────────────────────────────────────
from terminal.search_engine import (
    search_nse_announcements,
    search_corporate_actions,
    search_insider_trades,
    search_bse_filings,
    search_shareholding_analysis,
    search_analyst_coverage,
    search_concall_transcripts,
    search_sector_news,
    search_social_buzz,
    search_broker_research,
    search_mf_holdings,
    deep_search,
)

# ── Forensic accounting suite ─────────────────────────────────────────────────
from terminal.forensics import run_forensic_analysis, screen_forensic_watchlist
from terminal.postgres_tools import (
    audit_postgres_coverage,
    ensure_postgres_schema,
    get_data_source_manifest,
    get_postgres_health,
    load_historical_eod_to_postgres,
    load_intraday_ohlcv_to_postgres,
)
from terminal.report_context import (
    compare_reports,
    get_last_report,
    list_generated_reports,
    open_report,
    read_report,
    summarize_report,
)
from terminal.symbol_search import project_legacy_result
from terminal.symbol_search import resolve as _hybrid_resolve_symbol
from terminal.entity_resolution import (
    detect_non_symbol_terms,
    resolve_company_alias,
    resolve_index_or_stock,
    resolve_stock_entity,
    validate_requested_symbols,
)
from terminal.situation_assessment import (
    assess_user_situation,
    request_clarification,
    resolve_conversation_reference,
    resolve_entity_context,
    validate_intent_evidence_plan,
)
from terminal.results_tools import (
    discover_financial_filings,
    get_latest_results,
    ingest_financial_filing,
    parse_financial_filing,
    parse_pdf_filing as parse_results_pdf_filing,
    parse_xbrl_filing,
    reconcile_filing_facts,
    summarize_latest_results,
)
from terminal.youtube import analyze_youtube_channel_latest, analyze_youtube_video, list_youtube_channels
from terminal.evidence_gate import (
    build_evidence_matrix,
    render_missing_evidence_block,
    validate_answer_against_evidence,
    validate_required_tools_executed,
)
from terminal.fno_composite import (
    get_cost_of_carry as get_composite_cost_of_carry,
    get_fno_overview,
    get_futures_basis as get_composite_futures_basis,
    get_max_pain as get_composite_max_pain,
    get_option_chain_summary,
    get_pcr_summary,
    get_top_oi_strikes,
    recommend_options_strategy,
)
from terminal.company_evidence_tools import (
    audit_company_search,
    get_company_evidence_coverage,
    promote_company_evidence_to_postgres,
    search_company_filings as audit_search_company_filings,
    search_company_official_sources,
)

# ── Seasonal / macro modules ──────────────────────────────────────────────────
import sys as _sys
_sys.path.insert(0, str(Path(__file__).parent.parent))
from seasonal_heat_calendar import build_seasonal_heat_calendar, get_all_seasonal_signals
from economic_cycle import detect_economic_cycle_phase

# ── Chart module ─────────────────────────────────────────────────────────────
from terminal.charts import render_chart, render_html_chart, chart_summary
from terminal.visual_scan.command import run_visual_scan

# ── Intraday screener engine ──────────────────────────────────────────────────
from terminal.intraday import (
    compute_all as _compute_intraday_all,
    get_intraday_analysis,
    run_intraday_screener as _run_intraday_screener,
    get_intraday_candles,
    key_levels as _intraday_key_levels,
    _quiet_yf_download,
    run_all_signals as _run_intraday_all_signals,
)
from terminal.intraday_storage import persist_intraday_bars, persist_intraday_snapshot

# ── Paths ─────────────────────────────────────────────────────────────────────
ROOT      = Path(__file__).parent.parent
DB_PATH   = ROOT / "data" / "sector_rotation_tracker.db"
STOCK_CSV = ROOT / "data" / "nse_sec_full_data.csv"
INDEX_CSV = ROOT / "data" / "nse_index_data.csv"
# PG-SCAN-FALLBACK: local mapping used when NSE API blocks the constituents call
INDEX_MAPPING_CSV = ROOT / "data" / "index_stock_mapping.csv"
GLOBAL_INDEX_CSV = ROOT / "data" / "global_indices.csv"
GLOBAL_CORR_CSV  = ROOT / "data" / "global_correlations.csv"
REPORTS   = ROOT / "reports"
PG_DSN    = os.environ.get("AGENT_ADDA_PG_DSN") or os.environ.get("PG_DSN") or "dbname=nse_market user=nse_admin host=/tmp"
_INDEX_REFERENCE_ALIAS_CACHE: dict[str, str] | None = None
IST_TZ = "Asia/Kolkata"

try:
    from dotenv import load_dotenv  # type: ignore
    load_dotenv(ROOT / ".env")
except Exception:
    pass


# ─────────────────────────────────────────────────────────────────────────────
# Internal helpers
# ─────────────────────────────────────────────────────────────────────────────

def _db_conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def _legacy_sqlite_fallbacks_enabled() -> bool:
    return os.environ.get("AGENT_ADDA_ENABLE_SQLITE_FALLBACKS", "").strip().lower() in {"1", "true", "yes"}


def _pg_conn():
    import psycopg2
    return psycopg2.connect(PG_DSN)


def _pg_fetchall(sql: str, params: tuple | list | None = None) -> list[tuple]:
    conn = _pg_conn()
    try:
        with conn.cursor() as cur:
            cur.execute(sql, params or ())
            return cur.fetchall()
    finally:
        conn.close()


def _pg_read_df(sql: str, params: tuple | list | None = None) -> pd.DataFrame:
    conn = _pg_conn()
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            return pd.read_sql_query(sql, conn, params=params)
    finally:
        conn.close()


def _latest_snapshot_date() -> str:
    try:
        rows = _pg_fetchall("SELECT MAX(snapshot_date)::text FROM scores.stage_snapshots")
        if rows and rows[0][0]:
            return rows[0][0]
    except Exception:
        pass
    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
        return "N/A"
    conn = _db_conn()
    row = conn.execute("SELECT MAX(snapshot_date) FROM stage_snapshots").fetchone()
    conn.close()
    return row[0] if row and row[0] else "N/A"


def _load_price_history(symbol: str, days: int = 400) -> pd.DataFrame:
    cutoff = (date.today() - timedelta(days=days)).isoformat()
    try:
        df = _pg_read_df(
            """
            SELECT symbol AS "SYMBOL",
                   trade_date AS "TIMESTAMP",
                   open AS "OPEN",
                   high AS "HIGH",
                   low AS "LOW",
                   close AS "CLOSE",
                   volume AS "TOTTRDQTY"
            FROM market.equity_eod
            WHERE symbol = %s
              AND trade_date >= %s
            ORDER BY trade_date
            """,
            (symbol, cutoff),
        )
        if not df.empty:
            out = _normalise_price_history_frame(symbol, df)
            out.attrs["data_source"] = "PostgreSQL market.equity_eod"
            return out
    except Exception:
        pass

    if not STOCK_CSV.exists():
        return pd.DataFrame()
    df = pd.read_csv(
        STOCK_CSV,
        usecols=["SYMBOL", "TIMESTAMP", "OPEN", "HIGH", "LOW", "CLOSE", "TOTTRDQTY"],
        low_memory=False,
    )
    df = df[(df["SYMBOL"] == symbol) & (df["TIMESTAMP"] >= cutoff)]
    df["TIMESTAMP"] = pd.to_datetime(df["TIMESTAMP"])
    for c in ["OPEN", "HIGH", "LOW", "CLOSE", "TOTTRDQTY"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")
    out = df.sort_values("TIMESTAMP")
    out.attrs["data_source"] = "EOD CSV"
    return out


def _normalise_price_history_frame(symbol: str, df: pd.DataFrame) -> pd.DataFrame:
    if df.empty:
        return pd.DataFrame()
    out = df.copy()
    if "TIMESTAMP" not in out.columns:
        if "Date" in out.columns:
            out["TIMESTAMP"] = out["Date"]
        elif out.index.name or not isinstance(out.index, pd.RangeIndex):
            out["TIMESTAMP"] = out.index
    rename = {
        "Open": "OPEN",
        "High": "HIGH",
        "Low": "LOW",
        "Close": "CLOSE",
        "Volume": "TOTTRDQTY",
    }
    out = out.rename(columns={k: v for k, v in rename.items() if k in out.columns})
    out["SYMBOL"] = symbol
    required = ["SYMBOL", "TIMESTAMP", "OPEN", "HIGH", "LOW", "CLOSE", "TOTTRDQTY"]
    missing = [col for col in required if col not in out.columns]
    if missing:
        return pd.DataFrame()
    out = out[required].copy()
    out["TIMESTAMP"] = pd.to_datetime(out["TIMESTAMP"], errors="coerce")
    for col in ["OPEN", "HIGH", "LOW", "CLOSE", "TOTTRDQTY"]:
        out[col] = pd.to_numeric(out[col], errors="coerce")
    out = out.dropna(subset=["TIMESTAMP", "OPEN", "HIGH", "LOW", "CLOSE"])
    return out.sort_values("TIMESTAMP")


def _fetch_on_demand_price_history(symbol: str, days: int = 400) -> pd.DataFrame:
    """Fetch missing EOD price history on demand using yfinance NSE ticker."""
    try:
        import yfinance as yf
    except Exception:
        return pd.DataFrame()
    period = "2y" if days <= 500 else "5y"
    try:
        fetched = _quiet_yf_download(
            yf,
            f"{symbol}.NS",
            period=period,
            interval="1d",
            auto_adjust=False,
            progress=False,
            threads=False,
        )
    except Exception:
        return pd.DataFrame()
    if isinstance(fetched.columns, pd.MultiIndex):
        fetched.columns = [str(col[0]) for col in fetched.columns]
    return _normalise_price_history_frame(symbol, fetched.tail(max(days, 260)))


def _load_on_demand_price_history(symbol: str, days: int = 400) -> pd.DataFrame:
    """Read previously fetched on-demand EOD bars from PostgreSQL."""
    try:
        import psycopg2
    except Exception:
        return pd.DataFrame()
    dsn = os.environ.get("AGENT_ADDA_PG_DSN") or os.environ.get("PG_DSN") or "dbname=nse_market user=nse_admin host=/tmp"
    cutoff = (date.today() - timedelta(days=days)).isoformat()
    conn = None
    try:
        conn = psycopg2.connect(dsn)
        with warnings.catch_warnings():
            warnings.simplefilter("ignore", UserWarning)
            df = pd.read_sql_query(
                """
                SELECT symbol AS "SYMBOL",
                       trade_date AS "TIMESTAMP",
                       open AS "OPEN",
                       high AS "HIGH",
                       low AS "LOW",
                       close AS "CLOSE",
                       volume AS "TOTTRDQTY"
                FROM on_demand.eod_price_history
                WHERE symbol = %s AND trade_date >= %s
                ORDER BY trade_date
                """,
                conn,
                params=(symbol, cutoff),
            )
    except Exception:
        return pd.DataFrame()
    finally:
        if conn:
            conn.close()
    return _normalise_price_history_frame(symbol, df)


def persist_on_demand_eod_history(symbol: str, df: pd.DataFrame) -> dict:
    """Best-effort load of on-demand EOD bars into PostgreSQL."""
    bars = _normalise_price_history_frame(symbol, df)
    if bars.empty:
        return {"ok": False, "rows_inserted": 0, "reason": "no_valid_bars"}
    try:
        import psycopg2
        from psycopg2.extras import Json, execute_values
    except Exception as exc:
        return {"ok": False, "rows_inserted": 0, "reason": f"postgres_unavailable: {exc}"}

    dsn = os.environ.get("AGENT_ADDA_PG_DSN") or os.environ.get("PG_DSN") or "dbname=nse_market user=nse_admin host=/tmp"
    rows = [
        (
            str(row["SYMBOL"]).upper(),
            row["TIMESTAMP"].date(),
            float(row["OPEN"]),
            float(row["HIGH"]),
            float(row["LOW"]),
            float(row["CLOSE"]),
            int(row["TOTTRDQTY"]) if pd.notna(row["TOTTRDQTY"]) else None,
            "yfinance",
            Json({"source": "on-demand chat fetch"}),
        )
        for _, row in bars.iterrows()
    ]
    sql = """
        CREATE SCHEMA IF NOT EXISTS on_demand;
        CREATE TABLE IF NOT EXISTS on_demand.eod_price_history (
            symbol      TEXT NOT NULL,
            trade_date  DATE NOT NULL,
            open        NUMERIC(18,6),
            high        NUMERIC(18,6),
            low         NUMERIC(18,6),
            close       NUMERIC(18,6),
            volume      BIGINT,
            source      TEXT NOT NULL,
            raw_json    JSONB NOT NULL DEFAULT '{}'::jsonb,
            loaded_at   TIMESTAMPTZ NOT NULL DEFAULT now(),
            PRIMARY KEY (symbol, trade_date, source)
        );
        CREATE INDEX IF NOT EXISTS idx_on_demand_eod_symbol_date
            ON on_demand.eod_price_history (symbol, trade_date DESC);
    """
    insert_sql = """
        INSERT INTO on_demand.eod_price_history
            (symbol, trade_date, open, high, low, close, volume, source, raw_json)
        VALUES %s
        ON CONFLICT (symbol, trade_date, source) DO UPDATE SET
            open = EXCLUDED.open,
            high = EXCLUDED.high,
            low = EXCLUDED.low,
            close = EXCLUDED.close,
            volume = EXCLUDED.volume,
            raw_json = EXCLUDED.raw_json,
            loaded_at = now()
    """
    conn = None
    try:
        conn = psycopg2.connect(dsn)
        with conn.cursor() as cur:
            cur.execute(sql)
            execute_values(cur, insert_sql, rows, page_size=500)
        conn.commit()
        return {"ok": True, "rows_inserted": len(rows), "schema": "on_demand", "table": "eod_price_history"}
    except Exception as exc:
        if conn:
            conn.rollback()
        return {"ok": False, "rows_inserted": 0, "reason": str(exc)}
    finally:
        if conn:
            conn.close()


def _compute_rsi(closes: pd.Series, period: int = 14) -> float:
    if len(closes) < period + 1:
        return float("nan")
    delta = closes.diff().dropna()
    gain  = delta.clip(lower=0).ewm(com=period - 1, adjust=False).mean()
    loss  = (-delta.clip(upper=0)).ewm(com=period - 1, adjust=False).mean()
    rs    = gain / loss.replace(0, 1e-9)
    return round(float(100 - 100 / (1 + rs.iloc[-1])), 1)


def _compute_adx(grp: pd.DataFrame, period: int = 14) -> float:
    if len(grp) < period + 2:
        return 0.0
    h, l, c = grp["HIGH"].values, grp["LOW"].values, grp["CLOSE"].values
    tr  = [max(h[i]-l[i], abs(h[i]-c[i-1]), abs(l[i]-c[i-1])) for i in range(1, len(c))]
    pdm = [max(h[i]-h[i-1], 0) if (h[i]-h[i-1]) > (l[i-1]-l[i]) else 0 for i in range(1, len(h))]
    ndm = [max(l[i-1]-l[i], 0) if (l[i-1]-l[i]) > (h[i]-h[i-1]) else 0 for i in range(1, len(l))]
    atr = pd.Series(tr).ewm(span=period, adjust=False).mean()
    pdi = 100 * pd.Series(pdm).ewm(span=period, adjust=False).mean() / atr.replace(0, 1e-9)
    ndi = 100 * pd.Series(ndm).ewm(span=period, adjust=False).mean() / atr.replace(0, 1e-9)
    dx  = 100 * abs(pdi - ndi) / (pdi + ndi).replace(0, 1e-9)
    return round(float(dx.ewm(span=period, adjust=False).mean().iloc[-1]), 1)


def _compute_macd_signal(closes: pd.Series) -> str:
    if len(closes) < 26:
        return "N/A"
    ema12  = closes.ewm(span=12, adjust=False).mean()
    ema26  = closes.ewm(span=26, adjust=False).mean()
    macd   = ema12 - ema26
    signal = macd.ewm(span=9, adjust=False).mean()
    hist   = float(macd.iloc[-1] - signal.iloc[-1])
    return "bullish" if hist > 0 else "bearish"


def _safe_float(v: Any, digits: int = 2) -> float | None:
    """Return a rounded plain float, or None for missing/non-numeric values."""
    if v is None:
        return None
    try:
        fv = float(v)
    except (TypeError, ValueError):
        return None
    if pd.isna(fv):
        return None
    return round(fv, digits)


def _supertrend(grp: pd.DataFrame, period: int = 10, mult: float = 3.0) -> str | None:
    grp = grp.tail(60)
    if len(grp) < 20:
        return None
    h, l, c = grp["HIGH"].values, grp["LOW"].values, grp["CLOSE"].values
    tr  = [max(h[i]-l[i], abs(h[i]-c[i-1]), abs(l[i]-c[i-1])) for i in range(1, len(c))]
    atr = pd.Series(tr).ewm(span=period, adjust=False).mean().values
    ub  = [(h[i]+l[i])/2 + mult*atr[i-1] for i in range(1, len(c))]
    lb  = [(h[i]+l[i])/2 - mult*atr[i-1] for i in range(1, len(c))]
    st, direction = ub[0], 1
    for i in range(1, len(ub)):
        if c[i] > st:
            direction = 1
            st = lb[i]
        else:
            direction = -1
            st = ub[i]
    return "BUY" if direction == 1 else "SELL"


# F&O index alias mapping — maps common user names to NSE derivatives symbols
_FO_INDEX_ALIASES: dict[str, str] = {
    # MIDCPNIFTY (Nifty Midcap Select)
    "NIFTY MIDCAP":            "MIDCPNIFTY",
    "NIFTY MIDCAP 50":         "MIDCPNIFTY",
    "NIFTY MIDCAP SELECT":     "MIDCPNIFTY",
    "NIFTY MIDCAP100":         "MIDCPNIFTY",
    "NIFTY MIDCAP 100":        "MIDCPNIFTY",
    "MIDCAP NIFTY":            "MIDCPNIFTY",
    "MIDCPNIFTY":              "MIDCPNIFTY",
    # BANKNIFTY
    "NIFTY BANK":              "BANKNIFTY",
    "BANK NIFTY":              "BANKNIFTY",
    "BANKNIFTY":               "BANKNIFTY",
    # FINNIFTY
    "NIFTY FINANCIAL":         "FINNIFTY",
    "NIFTY FINANCIAL SERVICES": "FINNIFTY",
    "NIFTY FIN":               "FINNIFTY",
    "FINNIFTY":                "FINNIFTY",
    # NIFTY
    "NIFTY":                   "NIFTY",
    "NIFTY 50":                "NIFTY",
    "NIFTY50":                 "NIFTY",
    # NIFTY NEXT 50
    "NIFTY NEXT 50":           "NIFTYNXT50",
    "NIFTY NXT 50":            "NIFTYNXT50",
    "NIFTYNXT50":              "NIFTYNXT50",
}

_COMMON_STOCK_ALIASES: dict[str, str] = {
    "BAJAJ FINANCE": "BAJFINANCE",
    "BAJAJ FIN": "BAJFINANCE",
    "BAJAJ FINSERV": "BAJAJFINSV",
    "HDFC BANK": "HDFCBANK",
    "ICICI BANK": "ICICIBANK",
    "KOTAK BANK": "KOTAKBANK",
    "TATA STEEL": "TATASTEEL",
    "TATA MOTORS": "TATAMOTORS",
    "USL": "UNITDSPR",
    "UNITED SPIRITS": "UNITDSPR",
    "UNITED SPIRITS LIMITED": "UNITDSPR",
    "DIAGEO INDIA": "UNITDSPR",
    # Frequently-typed alternative forms whose default fuzzy matches resolve to
    # the wrong (smaller / unrelated) company. Surfaced by the live 200-case
    # regression run — keep these explicit so company names like "Bharat
    # Petroleum" do not get matched to CHENNPETRO, etc.
    "BHARAT PETROLEUM": "BPCL",
    "BHARAT PETROLEUM CORPORATION": "BPCL",
    "MAHINDRA AND MAHINDRA": "M&M",
    "MAHINDRA & MAHINDRA": "M&M",
    "HINDUSTAN LEVER": "HINDUNILVR",
    "HUL": "HINDUNILVR",
    "HINDUSTAN UNILEVER": "HINDUNILVR",
    "STATE BANK OF INDIA": "SBIN",
    "STATE BANK": "SBIN",
    "ASIAN PAINTS": "ASIANPAINT",
    "POWER GRID": "POWERGRID",
    "POWER GRID CORPORATION": "POWERGRID",
    "ADANI PORTS": "ADANIPORTS",
    "ADANI ENTERPRISES": "ADANIENT",
    "RELIANCE INDUSTRIES": "RELIANCE",
    "TATA CONSULTANCY": "TCS",
    "TATA CONSULTANCY SERVICES": "TCS",
    "TATA INVESTMENT": "TATAINVEST",
    "TATA INVESTMENT CORPORATION": "TATAINVEST",
    "BHARAT FORGE": "BHARATFORG",
    "MARUTI SUZUKI": "MARUTI",
    "SUN PHARMA": "SUNPHARMA",
    "DR REDDY": "DRREDDY",
    "DR REDDYS": "DRREDDY",
    "DIXON TECH": "DIXON",
    "DIXON TECHNOLOGIES": "DIXON",
    "LARSEN AND TOUBRO": "LT",
    "LARSEN & TOUBRO": "LT",
    "PREMIER ENERGIES": "PREMIERENE",
    "HINDUSTAN AERONAUTICS": "HAL",
    "BAJAJ AUTO": "BAJAJ-AUTO",
    # PG-ALIAS-APOLLO: NSE ticker `APOLLO` is Apollo Micro Systems Ltd
    # (defence electronics). The ref.instruments table currently mislabels
    # it as "Apollo Tyres Limited" (Tyres is actually `APOLLOTYRE`), so the
    # name-based resolver never finds Apollo Micro Systems without these
    # explicit aliases.
    "APOLLO MICRO": "APOLLO",
    "APOLLO MICROSYSTEMS": "APOLLO",
    "APOLLO MICRO SYSTEMS": "APOLLO",
    "APOLLO MICRO SYSTEMS LIMITED": "APOLLO",
    "APOLLO MICRO SYSTEMS LTD": "APOLLO",
    "APOLLO TYRES": "APOLLOTYRE",
    "APOLLO HOSPITALS": "APOLLOHOSP",
    "APOLLO HOSPITALS ENTERPRISE": "APOLLOHOSP",
}

_SYMBOL_CONTEXT_TOKENS: set[str] = {
    "ANALYSIS",
    "ASSESSMENT",
    "BSE",
    "CATALYST",
    "CATALYSTS",
    "CONCALL",
    "CONCALLS",
    "CONTEXT",
    "DIVIDEND",
    "GROWTH",
    "HOLDING",
    "HOLDINGS",
    "INSIDER",
    "NEWS",
    "OUTLOOK",
    "REPORT",
    "RESULT",
    "RESULTS",
    "SEARCH",
    "SHAREHOLDING",
    "SOCIAL",
    "STRATEGY",
    "STRATEGIES",
}


_GENERIC_NAME_TOKENS: frozenset[str] = frozenset({
    # Generic business / industry words that appear in many company names.
    # Never register these as single-token aliases — they collide across
    # dozens of issuers (e.g. "INVEST" would otherwise resolve to whichever
    # company name was scanned first).
    "AUTO", "BANK", "BHARAT", "CEMENT", "COAL", "COMPANIES", "COMPANY",
    "CORP", "CORPORATION", "ELECTRIC", "ELECTRICALS", "ELECTRONICS",
    "ENERGIES", "ENERGY", "ENTERPRISE", "ENTERPRISES", "FINANCE",
    "FINANCIAL", "FINSERV", "FOODS", "GAS", "GLOBAL", "GROUP", "GROWTH",
    "HINDUSTAN", "HOLDING", "HOLDINGS", "HOTEL", "HOTELS", "INC", "INDIA",
    "INDIAN", "INDUSTRIES", "INDUSTRY", "INFRA", "INFRASTRUCTURE",
    "INTERNATIONAL", "INVEST", "INVESTMENT", "INVESTMENTS", "LEVER",
    "LIMITED", "LTD", "MANUFACTURING", "MOTOR", "MOTORS", "NATIONAL",
    "NETWORK", "NETWORKS", "PHARMA", "PHARMACEUTICALS", "POWER",
    "PRIVATE", "PRODUCTS", "PROJECTS", "PUBLIC", "SERVICES", "SOLUTIONS",
    "STEEL", "SYSTEMS", "TECH", "TECHNOLOGIES", "TECHNOLOGY",
})


def _register_name_tokens(mapping: dict[str, str], name_u: str, sym_u: str) -> None:
    """Register multi-word company-name tokens as fuzzy aliases.

    Skips generic English/business words that would otherwise collide across
    issuers (e.g. "INVEST" should not resolve to AUTHUM just because that
    name happens to be scanned first).
    """
    tokens = re.sub(r"[^A-Z0-9 ]", "", name_u).split()
    for t in tokens:
        if len(t) < 4 or t in _GENERIC_NAME_TOKENS:
            continue
        mapping.setdefault(t, sym_u)


def _all_symbols_map() -> dict[str, str]:
    """Return {normalized_name: symbol, symbol: symbol} for fuzzy resolution."""
    # Start with F&O index aliases — always available
    mapping: dict[str, str] = dict(_FO_INDEX_ALIASES)
    mapping.update(_COMMON_STOCK_ALIASES)

    try:
        rows = _pg_fetchall(
            """
            SELECT symbol, company_name FROM ref.instruments
            UNION
            SELECT symbol, company_name FROM scores.mv_latest_snapshot
            """
        )
        for sym, name in rows:
            if not sym:
                continue
            sym_u = str(sym).upper()
            mapping[sym_u] = sym_u
            if name:
                name_u = str(name).upper()
                mapping[name_u] = sym_u
                _register_name_tokens(mapping, name_u, sym_u)
        return mapping
    except Exception:
        pass

    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
        return mapping
    conn = _db_conn()
    rows = conn.execute(
        "SELECT DISTINCT symbol, company_name FROM stage_snapshots "
        "WHERE snapshot_date=(SELECT MAX(snapshot_date) FROM stage_snapshots)"
    ).fetchall()
    conn.close()
    for sym, name in rows:
        mapping[sym.upper()] = sym.upper()
        if name:
            mapping[name.upper()] = sym.upper()
            _register_name_tokens(mapping, name.upper(), sym.upper())
    return mapping


def _lookup_key(value: str) -> str:
    """Normalize symbol/name text for local symbol matching."""
    return re.sub(r"[^A-Z0-9]", "", str(value or "").upper())


def _resolve_local_symbol(query: str) -> dict:
    """Resolve a symbol/name from local DB aliases without network access."""
    result = _hybrid_resolve_symbol(
        query,
        alias_map=_all_symbols_map(),
        use_trigram=True,
    )
    return project_legacy_result(result)


def _suggest_local_symbols(query: str, limit: int = 5) -> list[str]:
    """Return likely local symbol suggestions without canonicalizing the query."""
    q_key = _lookup_key(query)
    if not q_key or len(q_key) < 4:
        return []
    suggestions: list[tuple[float, str]] = []
    for key, sym in _all_symbols_map().items():
        key_norm = _lookup_key(key)
        sym_norm = _lookup_key(sym)
        if not key_norm or key_norm != sym_norm:
            continue
        ratio = SequenceMatcher(None, q_key, key_norm).ratio()
        if ratio >= 0.80 or q_key[:4] == key_norm[:4]:
            suggestions.append((ratio, str(sym).upper()))
    suggestions.sort(key=lambda item: item[0], reverse=True)
    return list(dict.fromkeys(sym for _, sym in suggestions[:limit]))


def _canonical_symbol(symbol: str) -> str:
    resolved = _resolve_local_symbol(symbol)
    return str(resolved.get("symbol") or symbol).strip().upper()


# PG-YFIN-FALLBACK: Yahoo Finance company-name → NSE-ticker resolver.
# Cached per-process so repeated lookups for the same name don't re-hit
# the Yahoo search endpoint. Only EQUITY quoteType on NSI/NSE is accepted
# — we deliberately ignore BSE-only and foreign-listed results so the
# downstream agent never receives a non-NSE ticker.
_YFIN_SEARCH_CACHE: dict[str, tuple[str | None, str | None]] = {}


def _yahoo_search_nse_symbol(query: str) -> tuple[str | None, str | None]:
    """Resolve a free-text company name to an NSE ticker via Yahoo Finance.

    Returns ``(symbol, company_name)`` or ``(None, None)`` if no NSE-listed
    EQUITY match is found. The ``.NS`` suffix is stripped from the Yahoo
    symbol before it's returned (e.g. ``APOLLO.NS`` → ``APOLLO``).
    """
    key = (query or "").strip().upper()
    if not key or len(key) < 3:
        return (None, None)
    if key in _YFIN_SEARCH_CACHE:
        return _YFIN_SEARCH_CACHE[key]
    try:
        from yfinance import Search  # type: ignore[import-not-found]
    except Exception:
        _YFIN_SEARCH_CACHE[key] = (None, None)
        return (None, None)
    try:
        quotes = Search(query, max_results=10).quotes or []
    except Exception:
        _YFIN_SEARCH_CACHE[key] = (None, None)
        return (None, None)
    for q in quotes:
        if not isinstance(q, dict):
            continue
        if (q.get("quoteType") or "").upper() != "EQUITY":
            continue
        sym_raw = str(q.get("symbol") or "").strip().upper()
        # NSE India tickers end in ".NS" (exchange=NSI); BSE end in ".BO".
        if not sym_raw.endswith(".NS"):
            continue
        nse_sym = sym_raw[:-3]
        if not re.fullmatch(r"[A-Z0-9&-]{2,16}", nse_sym):
            continue
        name = q.get("longname") or q.get("shortname") or None
        _YFIN_SEARCH_CACHE[key] = (nse_sym, name)
        return (nse_sym, name)
    _YFIN_SEARCH_CACHE[key] = (None, None)
    return (None, None)


# ─────────────────────────────────────────────────────────────────────────────
# Tool functions (all return dict)
# ─────────────────────────────────────────────────────────────────────────────

def resolve_symbol(query: str) -> dict:
    """Resolve a company name / partial name / alias to its NSE symbol.

    AA-HSR-4: this is now a **thin wrapper** over
    :func:`terminal.symbol_search.resolve`. The wrapper preserves the
    legacy return-dict shape so existing callers and tests are unchanged,
    while augmenting the payload with the new ``score`` and
    ``confidence_band`` fields. The orchestrator is responsible for the
    dict + isolated-typo + trigram tiers; we keep the NSE live-search and
    NSE quote-equity fallbacks here so they remain unit-testable and
    network-isolated.

    AA-HSR-5: structured telemetry is emitted inside the hybrid resolver
    itself (see :mod:`terminal.symbol_search.telemetry`), gated by the
    ``NSE_SYMBOL_RESOLUTION_TELEMETRY`` env var so tests never pollute
    ``logs/symbol_resolution.jsonl``.
    """
    import requests as _req
    q = query.strip().upper()
    # Guard: refuse to resolve well-known analytics/screener tokens that the
    # LLM sometimes mistakes for tickers (e.g. "high RS stocks" → "RS").
    # Tells the model exactly which screener to call instead.
    _CONCEPT_TOKENS = {
        "RS", "RSI", "PE", "PB", "EPS", "ROE", "ROCE", "EBITDA", "CAGR",
        "ATH", "ATL", "IV", "OI", "PCR", "VCP", "ORB", "BB", "MACD",
        "VWAP", "FII", "DII", "MF", "AMC", "CANSLIM", "CAN-SLIM",
        "MOMENTUM", "BREAKOUT", "BREAKOUTS", "LEADERS", "BASING",
        "TURNAROUND", "GAINERS", "LOSERS", "MOVERS",
        # Added more tokens that the symbol extractor used to mistake for
        # tickers when the keyword router missed the screener intent.
        "OVERSOLD", "TIGHT", "BULK", "MOST", "GAP", "BOLLINGER",
        "REVIEW", "TOP", "HIGH", "LOW", "NEW", "BEST", "WORST",
        "STRONG", "WEAK", "BUY", "SELL", "HOLD",
        "GROWTH", "STRATEGY", "STRATEGIES", "OUTLOOK", "REPORT", "RESULTS",
        # Time/calendar/event tokens — never a ticker.
        "DUE", "TOMORROW", "YESTERDAY", "TODAY", "TONIGHT",
        "MONDAY", "TUESDAY", "WEDNESDAY", "THURSDAY", "FRIDAY",
        "WEEK", "WEEKLY", "MONTH", "MONTHLY", "YEAR", "YEARLY",
        "UPCOMING", "FORTHCOMING", "RECENT", "RECENTLY",
        "REPORTING", "REPORTED", "ANNOUNCED", "FILED", "POSTED",
        "EARNINGS", "DIVIDEND", "DIVIDENDS", "AGM", "RIGHTS", "SPLIT", "BONUS",
    }
    if q in _CONCEPT_TOKENS or q in _GENERIC_NAME_TOKENS:
        return {
            "symbol": None,
            "confidence": "none",
            "confidence_band": "none",
            "score": 0.0,
            "method": "none",
            "query": query,
            "error": (
                f"'{query}' is a market concept / screener keyword, not an NSE ticker. "
                "Call run_screener_query (e.g. screen_type='high_rs', 'momentum_52w', "
                "'stage2', 'turnaround') or search_market_knowledge instead."
            ),
        }

    # NEW PRIMARY PATH: hybrid resolver (dict + isolated typo + trigram).
    try:
        rich = _hybrid_resolve_symbol(
            query,
            alias_map=_all_symbols_map(),
            use_trigram=True,
        )
    except Exception:
        rich = None

    if rich is not None and rich.symbol:
        payload: dict = {
            "symbol":           rich.symbol,
            "confidence":       rich.legacy_confidence,
            "confidence_band":  rich.confidence_band,
            "score":            float(rich.score),
            "method":           rich.method,
            "query":            query,
        }
        if rich.matched:
            payload["matched"] = rich.matched
        if rich.candidates:
            payload["candidates"] = list(dict.fromkeys(c.symbol for c in rich.candidates))[:5]
        return payload

    # Search-context guard: prose like "growth strategy outlook for Reliance"
    # should not be resolved as a ticker if the leading token is a context
    # word and the rest didn't parse to a known company.
    q_tokens = re.sub(r"[^A-Z0-9 ]", " ", q).split()
    if len(q_tokens) > 1 and any(tok in _SYMBOL_CONTEXT_TOKENS for tok in q_tokens):
        return {
            "symbol": None,
            "confidence": "none",
            "confidence_band": "none",
            "score": 0.0,
            "method": "none",
            "query": query,
            "error": f"'{query}' contains search/report context, not a resolvable NSE symbol.",
        }
    exact_ticker_query = bool(re.fullmatch(r"[A-Z0-9&-]{2,12}", query.strip()))
    suggestions = _suggest_local_symbols(query)

    # Fall back to NSE live search API
    try:
        s   = _get_live_session()
        url = f"https://www.nseindia.com/api/search?q={_req.utils.quote(query)}&type=equity"
        r   = s.get(url, timeout=10)
        r.raise_for_status()
        payload_search = r.json()
        results = payload_search.get("results", []) if isinstance(payload_search, dict) else []
        if results:
            search_suggestions = [
                str(x.get("symbol") or "").strip().upper()
                for x in results[:5]
                if isinstance(x, dict) and x.get("symbol")
            ]
            suggestions = list(dict.fromkeys([*suggestions, *search_suggestions]))[:5]
            # Prefer exact ticker match when the user passed a ticker-shaped
            # query, otherwise NSE search ranks by relevance and a substring
            # match can land first (e.g. RELIANCE → RELIANCEPOWER).
            q_up = query.strip().upper()
            top = None
            for r_ in results:
                if isinstance(r_, dict) and (r_.get("symbol") or "").upper() == q_up:
                    top = r_
                    break
            if top is None and not exact_ticker_query:
                top = next((r_ for r_ in results if isinstance(r_, dict)), None)
            if top is None:
                raise ValueError("No exact NSE search result for ticker-shaped query")
            return {
                "symbol":           top.get("symbol"),
                "name":             top.get("symbol_info"),
                "confidence":       "nse-search",
                "confidence_band":  "medium",
                "score":            0.70,
                "method":           "live_api",
                "query":            query,
                "candidates": [x.get("symbol") for x in results[:5] if isinstance(x, dict)],
            }
    except Exception:
        pass

    # Secondary fallback: NSE quote-equity endpoint. /api/search is frequently
    # rate-limited or returns HTTP 500; the per-symbol quote endpoint is far
    # more reliable. Only attempt if the query already looks like a ticker.
    q_clean = query.strip().upper()
    if re.fullmatch(r"[A-Z0-9&-]{2,12}", q_clean):
        try:
            s = _get_live_session()
            url = f"https://www.nseindia.com/api/quote-equity?symbol={_req.utils.quote(q_clean)}"
            r = s.get(url, timeout=10)
            if r.ok is True:
                payload_quote = r.json()
                info = (payload_quote.get("info") or {}) if isinstance(payload_quote, dict) else {}
                resolved = (info.get("symbol") or "").strip().upper()
                if resolved == q_clean:
                    return {
                        "symbol":           resolved,
                        "name":             info.get("companyName") or info.get("symbol_info"),
                        "confidence":       "nse-quote",
                        "confidence_band":  "high",
                        "score":            0.90,
                        "method":           "live_api",
                        "query":            query,
                        "candidates":       [resolved],
                    }
        except Exception:
            pass

    # PG-YFIN-FALLBACK: Yahoo Finance search — last-resort when local DB,
    # NSE /api/search, and NSE /api/quote-equity all fail (common when the
    # name doesn't appear in local data and NSE blocks our cookies).
    # Yahoo's search returns ``APOLLO.NS`` (NSE India) which we strip to
    # the bare NSE ticker. Restricted to EQUITY quoteType + NSI/NSE exchange
    # so we never leak BSE-only or US-listed tickers into the agent.
    try:
        yfin_symbol, yfin_name = _yahoo_search_nse_symbol(query)
    except Exception:
        yfin_symbol, yfin_name = None, None
    if yfin_symbol:
        return {
            "symbol":           yfin_symbol,
            "name":             yfin_name,
            "confidence":       "yfin-search",
            "confidence_band":  "medium",
            "score":            0.65,
            "method":           "yahoo_finance",
            "query":            query,
            "candidates":       [yfin_symbol],
        }

    result = {
        "symbol": None,
        "confidence": "none",
        "confidence_band": "none",
        "score": 0.0,
        "method": "none",
        "query": query,
        "error": f"No exact NSE symbol found for '{query}'" if exact_ticker_query else f"No NSE symbol found for '{query}'",
    }
    if suggestions:
        result["candidates"] = suggestions[:5]
        result["suggestion"] = suggestions[0]
    return result


_STAGE_SNAPSHOT_COLS = [
    "company_name", "stage", "stage_score", "investment_score", "price",
    "rsi", "relative_strength", "change_1d_pct", "change_1w_pct", "change_1m_pct",
    "market_cap_cat", "sector", "trading_signal", "trend_signal", "supertrend_state",
    "supertrend_value", "technical_score", "fundamental_score", "narrative", "stance",
    "can_slim_score", "minervini_score", "enhanced_fund_score", "earnings_quality",
    "sales_growth", "financial_strength", "institutional_backing",
]


def _stage_snapshot_missing_fields(snap: dict[str, Any]) -> list[str]:
    """Critical fields whose absence prevents trustworthy stage analysis.

    Advisory score fields (fundamental_score, can_slim_score, minervini_score)
    are NOT included here — they are computed by separate offline pipelines and
    often legitimately absent for many symbols. They go into
    ``missing_optional_scores`` instead so renderers don't surface them as a
    red-flag MISSING EVIDENCE block.
    """
    missing: list[str] = []
    for field in ("stage", "stage_score", "price", "rsi", "relative_strength", "technical_score", "trading_signal"):
        if snap.get(field) is None:
            missing.append(field)
    return missing


def _stage_snapshot_missing_optional_scores(snap: dict[str, Any]) -> list[str]:
    optional: list[str] = []
    if snap.get("fundamental_score") is None and snap.get("enhanced_fund_score") is None:
        optional.append("fundamental_score")
    for field in ("can_slim_score", "minervini_score"):
        if snap.get(field) is None:
            optional.append(field)
    return optional


def _finalize_stage_snapshot(sym: str, snap: dict[str, Any], snapshot_date: str) -> dict[str, Any]:
    snap["symbol"] = sym
    snap["snapshot_date"] = snapshot_date
    rs = snap.get("relative_strength")
    if rs is not None:
        snap["rs_pct"] = normalize_relative_strength_pct(rs)
    snap["missing_evidence"] = _stage_snapshot_missing_fields(snap)
    snap["missing_optional_scores"] = _stage_snapshot_missing_optional_scores(snap)
    snap["evidence_coverage"] = "complete" if not snap["missing_evidence"] else "partial"
    return snap


def _read_stage_snapshot_row(sym: str, snapshot_date: str) -> dict[str, Any] | None:
    try:
        rows = _pg_fetchall(
            """
            SELECT company_name, stage, stage_score, investment_score, price,
                   rsi, relative_strength, change_1d_pct, change_1w_pct, change_1m_pct,
                   market_cap_cat, sector, trading_signal, trend_signal, supertrend_state,
                   supertrend_value, technical_score, fundamental_score, narrative, stance,
                   can_slim_score, minervini_score, enhanced_fund_score, earnings_quality,
                   sales_growth, financial_strength, institutional_backing
            FROM scores.stage_snapshots
            WHERE symbol = %s AND snapshot_date = %s
            """,
            (sym, snapshot_date),
        )
        if rows:
            snap = {"data_source": "PostgreSQL scores.stage_snapshots"}
            snap.update(dict(zip(_STAGE_SNAPSHOT_COLS, rows[0])))
            return _finalize_stage_snapshot(sym, snap, snapshot_date)
    except Exception:
        pass

    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
        return None
    conn = _db_conn()
    row = conn.execute(
        "SELECT company_name, stage, stage_score, investment_score, price, "
        "rsi, relative_strength, change_1d_pct, change_1w_pct, change_1m_pct, "
        "market_cap_cat, sector, trading_signal, trend_signal, supertrend_state, "
        "supertrend_value, technical_score, fundamental_score, narrative, stance, "
        "can_slim_score, minervini_score, enhanced_fund_score, earnings_quality, "
        "sales_growth, financial_strength, institutional_backing "
        "FROM stage_snapshots "
        "WHERE symbol=? AND snapshot_date=?",
        (sym, snapshot_date),
    ).fetchone()
    conn.close()
    if not row:
        return None
    snap = {"data_source": "stage_snapshots DB"}
    snap.update(dict(zip(_STAGE_SNAPSHOT_COLS, row)))
    return _finalize_stage_snapshot(sym, snap, snapshot_date)


def _pct_change_from_history(sym: str, days: int) -> float | None:
    hist = _load_price_history(sym, max(days + 30, 80))
    if hist.empty:
        hist = _load_on_demand_price_history(sym, max(days + 30, 80))
    if hist.empty:
        return None
    hist = _normalise_price_history_frame(sym, hist).sort_values("TIMESTAMP")
    if len(hist) < 2:
        return None
    latest = float(hist["CLOSE"].iloc[-1])
    prior = float(hist["CLOSE"].iloc[max(0, len(hist) - days - 1)])
    return round((latest / prior - 1) * 100, 2) if prior else None


def _derive_stage_from_technicals(tech: dict[str, Any]) -> tuple[str, float]:
    price = tech.get("price")
    sma50 = tech.get("sma50")
    sma200 = tech.get("sma200")
    pct_from_52h = tech.get("pct_from_52h")
    score = float(tech.get("technical_score") or 0) / 100
    if all(isinstance(v, (int, float)) for v in (price, sma50, sma200)):
        if price > sma50 > sma200 and (pct_from_52h is None or pct_from_52h > -25):
            return "STAGE_2", round(min(0.99, max(0.5, score)), 4)
        if price < sma50 < sma200:
            return "STAGE_4", round(min(0.99, max(0.1, 1 - score)), 4)
        if price > sma200 and price < sma50:
            return "STAGE_3", round(min(0.99, max(0.2, score * 0.75)), 4)
        return "STAGE_1", round(min(0.99, max(0.2, score * 0.70)), 4)
    return "UNKNOWN", 0.0


def _on_demand_trading_signal(stage: str, tech: dict[str, Any]) -> tuple[str, str, str]:
    rsi = tech.get("rsi")
    macd = tech.get("macd")
    adx = tech.get("adx")
    supertrend = tech.get("supertrend")
    bullish = macd == "bullish" and tech.get("above_sma50") and tech.get("above_sma200")
    if stage == "STAGE_2" and bullish and isinstance(rsi, (int, float)) and rsi <= 75:
        signal = "BUY"
    elif stage == "STAGE_4":
        signal = "SELL"
    else:
        signal = "HOLD"
    if bullish and isinstance(adx, (int, float)) and adx >= 25:
        trend = "STRONG_BULLISH"
    elif bullish:
        trend = "BULLISH"
    else:
        trend = "MIXED"
    st_state = "BULLISH" if supertrend == "BUY" else ("BEARISH" if supertrend == "SELL" else None)
    return signal, trend, st_state


def _backfill_on_demand_stage_snapshot(sym: str, snapshot_date: str) -> dict[str, Any] | None:
    tech = get_technical_setup(sym)
    if tech.get("error"):
        return None
    live = get_nse_intraday_snapshot(sym)
    stage, stage_score = _derive_stage_from_technicals(tech)
    signal, trend, st_state = _on_demand_trading_signal(stage, tech)
    change_1d = tech.get("chg_pct")
    change_1w = _pct_change_from_history(sym, 5)
    change_1m = _pct_change_from_history(sym, 21)
    relative_strength = change_1m
    price = tech.get("price")
    company_name = live.get("name") if live and not live.get("error") else sym
    sector = live.get("sector") if live and not live.get("error") else None
    market_cap = live.get("market_cap_cr") if live and not live.get("error") else None
    market_cap_cat = None
    if isinstance(market_cap, (int, float)):
        market_cap_cat = "LARGE_CAP" if market_cap >= 50000 else ("MID_CAP" if market_cap >= 5000 else "SMALL_CAP")
    row = {
        "snapshot_date": snapshot_date,
        "symbol": sym,
        "company_name": company_name,
        "stage": stage,
        "stage_score": stage_score,
        "price": price,
        "live_price": live.get("last_price") if live and not live.get("error") else None,
        "technical_score": tech.get("technical_score"),
        "rsi": tech.get("rsi"),
        "trading_signal": signal,
        "trend_signal": trend,
        "relative_strength": relative_strength,
        "change_1d_pct": change_1d,
        "change_1w_pct": change_1w,
        "change_1m_pct": change_1m,
        "market_cap_cat": market_cap_cat,
        "source_csv": "on-demand chat",
        "sector": sector,
        "fundamental_score": None,
        "enhanced_fund_score": None,
        "earnings_quality": None,
        "sales_growth": None,
        "financial_strength": None,
        "institutional_backing": None,
        "can_slim_score": None,
        "minervini_score": None,
        "investment_score": tech.get("technical_score"),
        "fund_details": None,
        "narrative": (
            f"{sym} was backfilled on demand from EOD price history and NSE live metadata. "
            "Fundamental/CANSLIM/Minervini fields were not found in the DB and were not inferred."
        ),
        "stance": "TECHNICAL_ONLY",
        "supertrend_state": st_state,
        "supertrend_value": None,
        "price_date": tech.get("as_of"),
    }

    try:
        from psycopg2.extras import Json
        pg_row = dict(row)
        pg_row["fund_details"] = Json(row.get("fund_details") or {})
        cols = list(pg_row.keys())
        placeholders = ", ".join(["%s"] * len(cols))
        updates = ", ".join(f"{c}=EXCLUDED.{c}" for c in cols if c not in ("snapshot_date", "symbol"))
        conn = _pg_conn()
        try:
            with conn.cursor() as cur:
                cur.execute(
                    f"""
                    INSERT INTO scores.stage_snapshots ({', '.join(cols)})
                    VALUES ({placeholders})
                    ON CONFLICT (snapshot_date, symbol) DO UPDATE SET {updates}
                    """,
                    [pg_row[c] for c in cols],
                )
            conn.commit()
        finally:
            conn.close()
    except Exception:
        pass

    if _legacy_sqlite_fallbacks_enabled() and DB_PATH.exists():
        conn = _db_conn()
        conn.execute(
            """INSERT OR REPLACE INTO stage_snapshots
                (snapshot_date, symbol, company_name, stage, stage_score, price, live_price,
                 technical_score, rsi, trading_signal, trend_signal, relative_strength,
                 change_1d_pct, change_1w_pct, change_1m_pct, market_cap_cat, source_csv,
                 sector, fundamental_score, enhanced_fund_score, earnings_quality, sales_growth,
                 financial_strength, institutional_backing, can_slim_score, minervini_score,
                 investment_score, fund_details, narrative, stance, supertrend_state, supertrend_value,
                 price_date)
               VALUES
                (:snapshot_date,:symbol,:company_name,:stage,:stage_score,:price,:live_price,
                 :technical_score,:rsi,:trading_signal,:trend_signal,:relative_strength,
                 :change_1d_pct,:change_1w_pct,:change_1m_pct,:market_cap_cat,:source_csv,
                 :sector,:fundamental_score,:enhanced_fund_score,:earnings_quality,:sales_growth,
                 :financial_strength,:institutional_backing,:can_slim_score,:minervini_score,
                 :investment_score,:fund_details,:narrative,:stance,:supertrend_state,:supertrend_value,
                 :price_date)""",
            row,
        )
        conn.commit()
        conn.close()

    snap = {
        "data_source": "on-demand stage snapshot",
        **{key: row.get(key) for key in _STAGE_SNAPSHOT_COLS if key in row},
        "on_demand_backfill": True,
    }
    return _finalize_stage_snapshot(sym, snap, snapshot_date)


def get_symbol_snapshot(symbol: str) -> dict:
    """Get latest EOD snapshot for a symbol: price, stage, RS, RSI, signals, sector."""
    sym = _canonical_symbol(symbol)
    snap: dict[str, Any] = {"symbol": sym, "data_source": "PostgreSQL scores.stage_snapshots"}

    snapshot_date = _latest_snapshot_date()
    found = _read_stage_snapshot_row(sym, snapshot_date)
    if found:
        return found

    backfilled = _backfill_on_demand_stage_snapshot(sym, snapshot_date)
    if backfilled:
        return backfilled

    snap["error"] = f"{sym} not found in DB snapshot"
    snap["missing_evidence"] = ["stage_snapshot"]
    snap["evidence_coverage"] = "missing"
    snap["snapshot_date"] = snapshot_date
    return snap


def get_technical_setup(symbol: str, days: int = 400) -> dict:
    """Compute technical indicators for a symbol from price history CSV."""
    sym = _canonical_symbol(symbol)
    grp = _load_price_history(sym, days)
    data_source = grp.attrs.get("data_source", "PostgreSQL market.equity_eod") if not grp.empty else "PostgreSQL market.equity_eod"
    pg_persist: dict[str, Any] | None = None
    if grp.empty:
        cached = _load_on_demand_price_history(sym, days)
        if not cached.empty:
            grp = cached
            data_source = "PostgreSQL on-demand EOD"
        else:
            fetched = _fetch_on_demand_price_history(sym, days)
            if not fetched.empty:
                pg_persist = persist_on_demand_eod_history(sym, fetched)
                grp = fetched
                data_source = "on-demand yfinance EOD"

        if grp.empty:
            return {
                "symbol": sym,
                "error": "No price history available; on-demand fetch returned no usable EOD bars",
                "missing_evidence": ["price_history"],
                "data_source": "missing",
            }
    else:
        original_source = grp.attrs.get("data_source", data_source)
        grp = _normalise_price_history_frame(sym, grp)
        grp.attrs["data_source"] = original_source
        if grp.empty:
            return {
                "symbol": sym,
                "error": "No usable price history available after validation",
                "missing_evidence": ["price_history"],
                "data_source": "missing",
            }

    closes = grp["CLOSE"]
    latest = grp.iloc[-1]
    prev   = grp.iloc[-2] if len(grp) > 1 else grp.iloc[-1]

    rsi  = _compute_rsi(closes)
    adx  = _compute_adx(grp)
    macd = _compute_macd_signal(closes)
    st   = _supertrend(grp)

    c   = closes.values
    sma20  = round(float(c[-20:].mean()), 2)  if len(c) >= 20  else None
    sma50  = round(float(c[-50:].mean()), 2)  if len(c) >= 50  else None
    sma200 = round(float(c[-200:].mean()), 2) if len(c) >= 200 else None
    cur    = round(float(latest["CLOSE"]), 2)

    w52_high = round(float(grp["HIGH"].max()), 2)
    w52_low  = round(float(grp["LOW"].min()), 2)
    pct_from_52h = round((cur / w52_high - 1) * 100, 1) if w52_high else None

    avg_vol = round(float(grp["TOTTRDQTY"].tail(20).mean())) if "TOTTRDQTY" in grp else None
    last_vol = int(latest["TOTTRDQTY"]) if pd.notna(latest.get("TOTTRDQTY")) else None
    vol_ratio = round(last_vol / avg_vol, 2) if avg_vol and last_vol else None
    macd_bullish = macd == "bullish"
    supertrend_bullish = st == "BUY"
    technical_score = 0
    technical_score += 10 if sma20 and cur > sma20 else 0
    technical_score += 15 if sma50 and cur > sma50 else 0
    technical_score += 20 if sma200 and cur > sma200 else 0
    technical_score += 20 if isinstance(rsi, (int, float)) and 45 <= rsi <= 75 else 0
    technical_score += 15 if macd_bullish else 0
    technical_score += 15 if supertrend_bullish else 0
    technical_score += 5 if isinstance(adx, (int, float)) and adx >= 25 else 0

    return {
        "symbol":        sym,
        "price":         cur,
        "open":          round(float(latest["OPEN"]), 2),
        "high":          round(float(latest["HIGH"]), 2),
        "low":           round(float(latest["LOW"]), 2),
        "chg_pct":       round((cur / float(prev["CLOSE"]) - 1) * 100, 2),
        "rsi":           rsi,
        "adx":           adx,
        "macd":          macd,
        "supertrend":    st,
        "sma20":         sma20,
        "sma50":         sma50,
        "sma200":        sma200,
        "above_sma20":   (cur > sma20)  if sma20  else None,
        "above_sma50":   (cur > sma50)  if sma50  else None,
        "above_sma200":  (cur > sma200) if sma200 else None,
        "52w_high":      w52_high,
        "52w_low":       w52_low,
        "pct_from_52h":  pct_from_52h,
        "vol_last":      last_vol,
        "vol_avg_20d":   avg_vol,
        "vol_ratio":     vol_ratio,
        "technical_score": technical_score,
        "score_method":  "derived from EOD MA alignment, RSI, MACD, Supertrend, and ADX",
        "data_bars":     len(grp),
        "as_of":         str(pd.to_datetime(latest["TIMESTAMP"]).date()),
        "data_source":   data_source,
        "postgres_persist": pg_persist,
    }


def get_sector_context(sector_or_symbol: str) -> dict:
    """Get sector performance and stock composition context.
    Pass a stock symbol (e.g. 'BHEL') to auto-detect its sector, or a sector name directly."""
    q = _canonical_symbol(sector_or_symbol)
    snap_date = _latest_snapshot_date()

    def _format_sector_rows(sector: str, rows: list[tuple]) -> dict:
        cols = ["symbol","company_name","stage","investment_score","relative_strength",
                "change_1d_pct","change_1w_pct","change_1m_pct","rsi","trading_signal"]
        stocks = [dict(zip(cols, r)) for r in rows]
        s2_count  = sum(1 for s in stocks if s["stage"] == "STAGE_2")
        rs_values = [
            v for v in (normalize_relative_strength_pct(s.get("relative_strength")) for s in stocks)
            if v is not None
        ]
        avg_rs    = round(sum(rs_values) / len(rs_values), 1) if rs_values else None
        avg_1m    = round(sum(float(s["change_1m_pct"] or 0) for s in stocks) / len(stocks), 2)
        buy_sigs  = sum(1 for s in stocks if (s["trading_signal"] or "").startswith("BUY"))
        return {
            "sector":         sector,
            "snapshot_date":  snap_date,
            "total_stocks":   len(stocks),
            "stage2_count":   s2_count,
            "buy_signals":    buy_sigs,
            "avg_rs_pct":     avg_rs,
            "avg_1m_pct":     avg_1m,
            "top5_by_score":  stocks[:5],
            "weakest_3":      stocks[-3:],
            "data_source":    "PostgreSQL scores.stage_snapshots",
        }

    try:
        sym_rows = _pg_fetchall(
            "SELECT sector FROM scores.stage_snapshots WHERE symbol=%s AND snapshot_date=%s",
            (q, snap_date),
        )
        sector = sym_rows[0][0] if sym_rows and sym_rows[0][0] else sector_or_symbol
        rows = _pg_fetchall(
            """
            SELECT symbol, company_name, stage, investment_score, relative_strength,
                   change_1d_pct, change_1w_pct, change_1m_pct, rsi, trading_signal
            FROM scores.stage_snapshots
            WHERE UPPER(sector)=UPPER(%s) AND snapshot_date=%s
            ORDER BY investment_score DESC NULLS LAST
            """,
            (sector, snap_date),
        )
        if not rows:
            rows = _pg_fetchall(
                """
                SELECT symbol, company_name, stage, investment_score, relative_strength,
                       change_1d_pct, change_1w_pct, change_1m_pct, rsi, trading_signal
                FROM scores.stage_snapshots
                WHERE UPPER(sector) LIKE UPPER(%s) AND snapshot_date=%s
                ORDER BY investment_score DESC NULLS LAST
                """,
                (f"%{sector}%", snap_date),
            )
        if rows:
            return _format_sector_rows(sector, rows)
    except Exception:
        pass

    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
        return {"error": "PostgreSQL scores.stage_snapshots unavailable"}

    conn = _db_conn()

    # If it looks like a symbol, resolve its sector first
    sym_row = conn.execute(
        "SELECT sector FROM stage_snapshots WHERE symbol=? AND snapshot_date=?",
        (q, snap_date),
    ).fetchone()
    sector = sym_row[0] if sym_row and sym_row[0] else sector_or_symbol  # preserve original case

    # Case-insensitive sector match — try exact first, then LIKE fuzzy
    rows = conn.execute(
        """SELECT symbol, company_name, stage, investment_score, relative_strength,
                  change_1d_pct, change_1w_pct, change_1m_pct, rsi, trading_signal
           FROM stage_snapshots
           WHERE UPPER(sector)=UPPER(?) AND snapshot_date=?
           ORDER BY investment_score DESC""",
        (sector, snap_date),
    ).fetchall()

    if not rows:
        # Fuzzy fallback: LIKE '%sector%'
        rows = conn.execute(
            """SELECT symbol, company_name, stage, investment_score, relative_strength,
                      change_1d_pct, change_1w_pct, change_1m_pct, rsi, trading_signal
               FROM stage_snapshots
               WHERE UPPER(sector) LIKE UPPER(?) AND snapshot_date=?
               ORDER BY investment_score DESC""",
            (f"%{sector}%", snap_date),
        ).fetchall()
        # Also resolve common abbreviations
        if not rows:
            abbrev_map = {
                "IT": "Information Technology", "TECH": "Information Technology",
                "PHARMA": "Pharma & Healthcare", "HEALTH": "Pharma & Healthcare",
                "BANK": "Banking & Finance", "FINANCE": "Banking & Finance",
                "AUTO": "EV & Auto Ancillaries", "EV": "EV & Auto Ancillaries",
                "FMCG": "FMCG & Consumer Goods", "CONSUMER": "FMCG & Consumer Goods",
                "ENERGY": "Energy - Oil & Gas", "OIL": "Energy - Oil & Gas",
                "POWER": "Energy - Power", "METAL": "Metals & Mining",
                "MINING": "Metals & Mining", "DEFENCE": "Defence & Aerospace",
                "REALTY": "Realty", "REAL ESTATE": "Realty",
            }
            mapped = abbrev_map.get(sector.upper())
            if mapped:
                rows = conn.execute(
                    """SELECT symbol, company_name, stage, investment_score, relative_strength,
                              change_1d_pct, change_1w_pct, change_1m_pct, rsi, trading_signal
                       FROM stage_snapshots
                       WHERE UPPER(sector) LIKE UPPER(?) AND snapshot_date=?
                       ORDER BY investment_score DESC""",
                    (f"%{mapped}%", snap_date),
                ).fetchall()
    conn.close()

    if not rows:
        return {"sector": sector, "error": f"No stocks found for sector '{sector}'"}

    return _format_sector_rows(sector, rows)


def run_screener_query(screen_type: str = "stage2", top_n: int = 10) -> dict:
    """Run a pre-built EOD screener from the DB snapshot data.

    screen_type options:
      Original  : stage2, breakouts, supertrend_buy, strong_buy, new_entrants
      New EOD   : new_highs, momentum_52w, high_rs, turnaround, stage1_base, tight_range, oversold_bounce
    """
    snap_date = _latest_snapshot_date()
    screen_key = screen_type.lower()

    _descriptions = {
        "stage2":           "Stage 2 uptrend stocks — William O'Neil buy zone",
        "breakouts":        "Stage 2 stocks breaking out with rising RS + RSI ≥ 55",
        "supertrend_buy":   "Supertrend BUY state stocks",
        "strong_buy":       "STRONG_BUY signal stocks",
        "new_entrants":     "Stage 2 new entrants in the last 14 days",
        "new_highs":        "Companies creating new highs — latest close within 5% of computed 52-week high",
        "momentum_52w":     "Near 52W high momentum leaders — RS ≥ 1.0, 1M chg > 2%",
        "high_rs":          "Top relative strength leaders — RS ≥ 1.15",
        "turnaround":       "Recovery setups — dip + rising momentum + RS improving",
        "stage1_base":      "Stage 1 basing stocks — coiled, waiting for breakout",
        "tight_range":      "Tight weekly range + RS ≥ 1.0 — VCP-like consolidation",
        "oversold_bounce":  "Stage 2 stocks with RSI < 40 — dip-buy in uptrend",
    }

    _pg_base_cols = (
        "symbol, company_name, stage, stage_score, investment_score, technical_score, price, "
        "COALESCE(relative_strength, change_1m_pct) AS relative_strength, "
        "change_1m_pct, rsi, trading_signal, sector"
    )
    _pg_base_from = "FROM scores.stage_snapshots WHERE snapshot_date=%s"
    pg_query_map: dict[str, str] = {
        "stage2": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_2' ORDER BY investment_score DESC NULLS LAST",
        "breakouts": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_2' AND COALESCE(change_1m_pct,0)>3.0 AND COALESCE(rsi,0) BETWEEN 55 AND 85 AND supertrend_state='BULLISH' ORDER BY investment_score DESC NULLS LAST, change_1m_pct DESC NULLS LAST",
        "supertrend_buy": f"SELECT symbol, company_name, stage, stage_score, investment_score, technical_score, price, COALESCE(relative_strength, change_1d_pct) AS relative_strength, change_1d_pct, rsi, trading_signal, sector {_pg_base_from} AND supertrend_state='BULLISH' AND stage IN ('STAGE_1','STAGE_2') ORDER BY investment_score DESC NULLS LAST, rsi DESC NULLS LAST",
        "strong_buy": f"SELECT {_pg_base_cols} {_pg_base_from} AND trading_signal='STRONG_BUY' AND stage='STAGE_2' AND supertrend_state='BULLISH' ORDER BY investment_score DESC NULLS LAST",
        "new_entrants": "SELECT s.symbol, s.company_name, s.stage, s.stage_score, s.investment_score, s.technical_score, s.price, COALESCE(s.relative_strength, s.change_1m_pct) AS relative_strength, s.change_1m_pct, s.rsi, s.trading_signal, s.sector FROM scores.stage_snapshots s LEFT JOIN scores.stage_changes c ON s.symbol=c.symbol AND c.stage_now='STAGE_2' AND c.stage_prev!='STAGE_2' WHERE s.snapshot_date=%s AND s.stage='STAGE_2' AND (c.change_date >= (%s::date - interval '14 days') OR c.change_date IS NULL) ORDER BY s.investment_score DESC NULLS LAST",
        "new_highs": """
            WITH latest_date AS (
                SELECT MAX(trade_date) AS trade_date FROM market.equity_eod
            ),
            latest AS (
                SELECT e.*
                FROM market.equity_eod e
                JOIN latest_date d ON e.trade_date=d.trade_date
                WHERE e.series='EQ' AND e.close IS NOT NULL AND e.close > 0
            ),
            highs AS (
                SELECT symbol, MAX(high) AS high_52w
                FROM market.equity_eod
                WHERE trade_date >= (SELECT trade_date FROM latest_date) - INTERVAL '370 days'
                  AND series='EQ'
                  AND high IS NOT NULL
                GROUP BY symbol
            ),
            snap AS (
                SELECT *
                FROM scores.stage_snapshots
                WHERE snapshot_date=%s
            )
            SELECT l.symbol,
                   COALESCE(s.company_name, i.company_name, l.symbol) AS company_name,
                   COALESCE(s.stage, 'UNKNOWN') AS stage,
                   s.stage_score,
                   s.investment_score,
                   s.technical_score,
                   l.close AS price,
                   ROUND((l.close / NULLIF(h.high_52w, 0) * 100)::numeric, 2) AS relative_strength,
                   l.change_pct AS change_1d_pct,
                   s.rsi,
                   s.trading_signal,
                   s.sector
            FROM latest l
            JOIN highs h ON h.symbol=l.symbol
            LEFT JOIN snap s ON s.symbol=l.symbol
            LEFT JOIN ref.instruments i ON i.symbol=l.symbol
            WHERE h.high_52w > 0
              AND l.close >= h.high_52w * 0.95
              AND l.symbol !~* '(ETF|LIQUID|LIQID|BEES|GILT|BOND|MON100|MAFANG|MASPTOP|CASHIETF|MONQ)'
              AND COALESCE(s.company_name, i.company_name, l.symbol) !~* '(ETF|LIQUID|LIQID|BOND|GILT)'
            ORDER BY (l.close / NULLIF(h.high_52w, 0)) DESC NULLS LAST,
                     s.investment_score DESC NULLS LAST
        """,
        "momentum_52w": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_2' AND COALESCE(change_1m_pct,0)>5.0 AND COALESCE(rsi,0) BETWEEN 50 AND 85 AND supertrend_state='BULLISH' ORDER BY change_1m_pct DESC NULLS LAST, investment_score DESC NULLS LAST",
        "high_rs": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage IN ('STAGE_2','STAGE_1') AND COALESCE(change_1m_pct,0)>8.0 AND COALESCE(rsi,0)>=55 ORDER BY change_1m_pct DESC NULLS LAST, investment_score DESC NULLS LAST",
        "turnaround": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage IN ('STAGE_1','STAGE_2') AND COALESCE(change_1m_pct,0)>5.0 AND COALESCE(rsi,0) BETWEEN 40 AND 65 AND COALESCE(investment_score,0)<60 ORDER BY change_1m_pct DESC NULLS LAST, investment_score DESC NULLS LAST",
        "stage1_base": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_1' AND COALESCE(rsi,0) BETWEEN 40 AND 60 AND ABS(COALESCE(change_1m_pct,0))<5.0 ORDER BY investment_score DESC NULLS LAST, relative_strength DESC NULLS LAST",
        "tight_range": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_2' AND supertrend_state='BULLISH' AND ABS(COALESCE(change_1w_pct,0))<2.0 AND COALESCE(rsi,0) BETWEEN 45 AND 65 ORDER BY investment_score DESC NULLS LAST, rsi DESC NULLS LAST",
        "oversold_bounce": f"SELECT {_pg_base_cols} {_pg_base_from} AND stage='STAGE_2' AND COALESCE(rsi,0)<40 AND supertrend_state='BULLISH' ORDER BY rsi ASC NULLS LAST, investment_score DESC NULLS LAST",
    }
    if screen_key not in pg_query_map:
        available = sorted(pg_query_map.keys())
        return {"error": f"Unknown screener: {screen_type}", "available": available}

    try:
        params = (snap_date, snap_date, top_n) if screen_key == "new_entrants" else (snap_date, top_n)
        rows = _pg_fetchall(pg_query_map[screen_key] + " LIMIT %s", params)
        cols = ["symbol","company_name","stage","stage_score","investment_score","technical_score","price",
                "relative_strength","change","rsi","trading_signal","sector"]
        stocks = []
        for r in rows:
            d = dict(zip(cols, r))
            if d.get("relative_strength") is not None:
                d["rs_pct"] = normalize_relative_strength_pct(d["relative_strength"])
            stocks.append(d)
        return {
            "screen_type":    screen_key,
            "description":    _descriptions.get(screen_key, ""),
            "snapshot_date":  snap_date,
            "count":          len(stocks),
            "results":        stocks,
            "data_source":    "PostgreSQL market.equity_eod + scores.stage_snapshots"
                               if screen_key == "new_highs"
                               else "PostgreSQL scores.stage_snapshots",
        }
    except Exception:
        pass

    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
        return {"error": "PostgreSQL screener snapshot unavailable"}

    conn = _db_conn()

    _base_cols = (
        "symbol, company_name, stage, stage_score, investment_score, technical_score, price, "
        "COALESCE(relative_strength, change_1m_pct) AS relative_strength, "
        "change_1m_pct, rsi, trading_signal, sector"
    )
    _base_from = "FROM stage_snapshots WHERE snapshot_date=?"

    query_map: dict[str, str | tuple] = {
        # ── Original screeners ────────────────────────────────────────────────
        "stage2": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "ORDER BY investment_score DESC"
        ),
        "breakouts": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "AND COALESCE(change_1m_pct, 0) > 3.0 "
            "AND COALESCE(rsi, 0) BETWEEN 55 AND 85 "
            "AND supertrend_state='BULLISH' "
            "ORDER BY investment_score DESC, change_1m_pct DESC"
        ),
        "supertrend_buy": (
            "SELECT symbol, company_name, stage, stage_score, investment_score, technical_score, price, "
            "COALESCE(relative_strength, change_1d_pct) AS relative_strength, change_1d_pct, rsi, trading_signal, sector "
            f"{_base_from} AND supertrend_state='BULLISH' "
            "AND stage IN ('STAGE_1','STAGE_2') "
            "ORDER BY investment_score DESC, rsi DESC"
        ),
        "strong_buy": (
            f"SELECT {_base_cols} {_base_from} AND trading_signal='STRONG_BUY' "
            "AND stage='STAGE_2' AND supertrend_state='BULLISH' "
            "ORDER BY investment_score DESC"
        ),
        "new_entrants": (
            "SELECT s.symbol, s.company_name, s.stage, s.stage_score, s.investment_score, "
            "s.technical_score, s.price, COALESCE(s.relative_strength, s.change_1m_pct) AS relative_strength, "
            "s.change_1m_pct, s.rsi, s.trading_signal, s.sector "
            "FROM stage_snapshots s "
            "LEFT JOIN stage_changes c ON s.symbol=c.symbol AND c.stage_now='STAGE_2' "
            "AND c.stage_prev != 'STAGE_2' "
            "WHERE s.snapshot_date=? AND s.stage='STAGE_2' "
            "AND (c.change_date >= date(?, '-14 days') OR c.change_date IS NULL) "
            "ORDER BY s.investment_score DESC"
        ),

        # ── New EOD screeners ─────────────────────────────────────────────────

        # Stocks within 5% of their 52-week high with positive RS and upward momentum.
        # Classic "buy strength" — trend-following stocks leading the market.
        "new_highs": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "AND COALESCE(change_1m_pct, 0) > 5.0 "
            "ORDER BY change_1m_pct DESC, investment_score DESC"
        ),

        "momentum_52w": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "AND COALESCE(change_1m_pct, 0) > 5.0 "
            "AND COALESCE(rsi, 0) BETWEEN 50 AND 85 "
            "AND supertrend_state='BULLISH' "
            "ORDER BY change_1m_pct DESC, investment_score DESC"
        ),

        # Top RS ranked stocks — market leaders by relative strength score.
        # RS > 1.15 means outperforming 85%+ of the market.
        "high_rs": (
            f"SELECT {_base_cols} {_base_from} "
            "AND stage IN ('STAGE_2', 'STAGE_1') "
            "AND COALESCE(change_1m_pct, 0) > 8.0 "
            "AND COALESCE(rsi, 0) >= 55 "
            "ORDER BY change_1m_pct DESC, investment_score DESC"
        ),

        # Turnaround candidates: down >20% from recent highs but showing recovery.
        # RSI recovering from oversold + positive recent change signals base building.
        "turnaround": (
            f"SELECT {_base_cols} {_base_from} "
            "AND stage IN ('STAGE_1', 'STAGE_2') "
            "AND COALESCE(change_1m_pct, 0) > 5.0 "
            "AND COALESCE(rsi, 0) BETWEEN 40 AND 65 "
            "AND COALESCE(investment_score, 0) < 60 "
            "ORDER BY change_1m_pct DESC, investment_score DESC"
        ),

        # Stage 1 basing stocks — consolidating sideways before a potential breakout.
        # Low RS + flat price action = coiled spring. Watch for volume expansion.
        "stage1_base": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_1' "
            "AND COALESCE(rsi, 0) BETWEEN 40 AND 60 "
            "AND ABS(COALESCE(change_1m_pct, 0)) < 5.0 "
            "ORDER BY investment_score DESC, relative_strength DESC"
        ),

        # Tight range consolidation: low volatility + near recent highs + good RS.
        # VCP-like setup — volatility contraction precedes explosive moves.
        "tight_range": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "AND supertrend_state='BULLISH' "
            "AND ABS(COALESCE(change_1w_pct, 0)) < 2.0 "
            "AND COALESCE(rsi, 0) BETWEEN 45 AND 65 "
            "ORDER BY investment_score DESC, rsi DESC"
        ),

        # Oversold stocks in Stage 2 uptrend — mean-reversion bounce candidates.
        # RSI below 40 in an otherwise bullish stage = potential entry dip.
        "oversold_bounce": (
            f"SELECT {_base_cols} {_base_from} AND stage='STAGE_2' "
            "AND COALESCE(rsi, 0) < 40 "
            "AND supertrend_state='BULLISH' "
            "ORDER BY rsi ASC, investment_score DESC"
        ),
    }

    _multi_param_keys = {"new_entrants"}
    if screen_key not in query_map:
        available = sorted(query_map.keys())
        return {"error": f"Unknown screener: {screen_type}", "available": available}

    sql = query_map[screen_key]
    cols = ["symbol","company_name","stage","stage_score","investment_score","technical_score","price",
            "relative_strength","change","rsi","trading_signal","sector"]

    if screen_key in _multi_param_keys:
        rows = conn.execute(sql, (snap_date, snap_date)).fetchmany(top_n)
    else:
        rows = conn.execute(sql, (snap_date,)).fetchmany(top_n)
    conn.close()

    stocks = []
    for r in rows:
        d = dict(zip(cols, r))
        if d.get("relative_strength") is not None:
            d["rs_pct"] = normalize_relative_strength_pct(d["relative_strength"])
        stocks.append(d)

    return {
        "screen_type":    screen_key,
        "description":    _descriptions.get(screen_key, ""),
        "snapshot_date":  snap_date,
        "count":          len(stocks),
        "results":        stocks,
    }


def _normalize_index_name(index_name: str) -> str:
    """PG-SCAN-FALLBACK: canonicalize index name to NSE form (e.g. 'NIFTY500' -> 'NIFTY 500')."""
    raw = (index_name or "").strip().upper()
    raw = re.sub(r"\s+", " ", raw)
    # Insert a space between "NIFTY" and a trailing number when missing
    m = re.match(r"^NIFTY(\d{2,4})$", raw)
    if m:
        return f"NIFTY {m.group(1)}"
    # Common shorthand like "NIFTY500" embedded with suffixes is left alone
    return raw


def _load_index_constituents_local(index_name: str) -> list[str]:
    """PG-SCAN-FALLBACK: load index constituents from the local CSV mapping.

    Returns an empty list if the CSV is missing or the index has no entries.
    Used as a graceful fallback when the NSE live API is blocked / returns
    non-JSON (stale cookies, anti-bot challenge, off-hours splash page).
    """
    try:
        if not INDEX_MAPPING_CSV.exists():
            return []
        canonical = _normalize_index_name(index_name)
        symbols: list[str] = []
        with INDEX_MAPPING_CSV.open("r", encoding="utf-8") as fh:
            header = fh.readline()  # INDEX_NAME,STOCK_SYMBOL
            del header
            for line in fh:
                parts = line.rstrip("\n").split(",", 1)
                if len(parts) != 2:
                    continue
                idx_name, sym = parts[0].strip().upper(), parts[1].strip().upper()
                if idx_name == canonical and sym:
                    symbols.append(sym)
        return list(dict.fromkeys(symbols))
    except Exception:
        return []


def _fetch_nse_index_constituents(index_name: str) -> list[str]:
    """Fetch live NSE constituents for an equity index.

    The legacy ``equity-stockIndices?index=...`` endpoint that this used
    to call was deprecated by NSE in 2024 and now returns 404 for every
    index. The function therefore reads exclusively from the local
    ``data/index_stock_mapping.csv`` bundled mapping. The 10-second
    cookie-refreshed API attempt was removed because it could only ever
    waste the request budget and slow the caller down.
    """
    canonical = _normalize_index_name(index_name)
    return _load_index_constituents_local(canonical)


def _growth_index_names(index_scope: str) -> list[str]:
    scope = (index_scope or "MIDCAP").strip().upper()
    if "MID" in scope:
        return ["NIFTY MIDCAP 50", "NIFTY MIDCAP 100", "NIFTY MIDCAP 150", "NIFTY MIDCAP SELECT"]
    if "SMALL" in scope:
        return ["NIFTY SMALLCAP 50", "NIFTY SMALLCAP 100", "NIFTY SMALLCAP 250"]
    if "500" in scope:
        return ["NIFTY 500"]
    return [index_scope.strip().upper() or "NIFTY MIDCAP 150"]


def get_long_term_growth_candidates(index_scope: str = "MIDCAP", top_n: int = 12, include_research: bool = True) -> dict:
    """Rank index constituents for long-term growth research using DB evidence."""
    index_names = _growth_index_names(index_scope)
    symbols: list[str] = []
    warnings_list: list[str] = []
    source_indices: list[str] = []

    for index_name in index_names:
        try:
            fetched = _fetch_nse_index_constituents(index_name)
            if fetched:
                source_indices.append(index_name)
                symbols.extend(fetched)
        except Exception as exc:
            warnings_list.append(f"Could not fetch NSE constituents for {index_name}: {exc}")

    symbols = list(dict.fromkeys(symbols))
    if not symbols:
        return {
            "error": f"No live NSE constituents available for {index_scope}",
            "index_scope": index_scope,
            "indices": index_names,
            "warnings": warnings_list,
        }

    limit = max(1, min(int(top_n or 12), 30))
    rows: list[tuple] = []
    snapshot_date = _latest_snapshot_date()
    try:
        placeholders = ",".join(["%s"] * len(symbols))
        rows = _pg_fetchall(
            f"""
            SELECT symbol, company_name, sector, stage, price,
                   investment_score, technical_score, enhanced_fund_score,
                   financial_strength, sales_growth, earnings_quality,
                   can_slim_score, relative_strength, change_1m_pct,
                   rsi, trading_signal
            FROM scores.stage_snapshots
            WHERE snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)
              AND UPPER(symbol) IN ({placeholders})
            ORDER BY
              COALESCE(enhanced_fund_score, 0) DESC NULLS LAST,
              COALESCE(financial_strength, 0) DESC NULLS LAST,
              COALESCE(sales_growth, 0) DESC NULLS LAST,
              COALESCE(investment_score, 0) DESC NULLS LAST,
              COALESCE(relative_strength, 0) DESC NULLS LAST
            LIMIT %s
            """,
            [*symbols, limit],
        )
    except Exception as exc:
        return {
            "error": f"PostgreSQL growth-candidate snapshot unavailable: {exc}",
            "index_scope": index_scope,
            "indices": source_indices or index_names,
            "constituent_count": len(symbols),
            "warnings": warnings_list,
        }

    cols = [
        "symbol", "company_name", "sector", "stage", "price",
        "investment_score", "technical_score", "enhanced_fund_score",
        "financial_strength", "sales_growth", "earnings_quality",
        "can_slim_score", "relative_strength", "change_1m_pct",
        "rsi", "trading_signal",
    ]
    candidates: list[dict] = []
    for row in rows:
        item = dict(zip(cols, row))
        item["rs_pct"] = normalize_relative_strength_pct(item.get("relative_strength"))
        for key in (
            "price", "investment_score", "technical_score", "enhanced_fund_score",
            "financial_strength", "sales_growth", "earnings_quality", "can_slim_score",
            "change_1m_pct", "rsi",
        ):
            item[key] = _safe_float(item.get(key), 2)
        candidates.append(item)

    research_items: list[dict] = []
    if include_research:
        for item in candidates[: min(5, len(candidates))]:
            sym = str(item.get("symbol") or "").strip().upper()
            if not sym:
                continue
            try:
                sr = scrape_screener_in(sym)
                if sr.get("error"):
                    research_items.append({"symbol": sym, "error": sr.get("error"), "missing_evidence": ["screener_fundamentals"]})
                    continue
                ratios = sr.get("ratios") or {}
                research_items.append({
                    "symbol": sym,
                    "source_url": sr.get("source_url"),
                    "market_cap": ratios.get("Market Cap"),
                    "stock_pe": ratios.get("Stock P/E"),
                    "roe": ratios.get("ROE"),
                    "roce": ratios.get("ROCE"),
                    "pros": (sr.get("pros") or [])[:3],
                    "cons": (sr.get("cons") or [])[:2],
                })
            except Exception as exc:
                research_items.append({"symbol": sym, "error": str(exc), "missing_evidence": ["screener_fundamentals"]})

    return {
        "index_scope": index_scope,
        "indices": source_indices or index_names,
        "constituent_count": len(symbols),
        "snapshot_date": snapshot_date,
        "candidate_count": len(candidates),
        "candidates": candidates,
        "research_items": research_items,
        "warnings": warnings_list,
        "source_trail": ["NSE live index constituents", "PostgreSQL scores.stage_snapshots", *(["screener.in"] if include_research else [])],
    }


def normalize_relative_strength_pct(value: Any) -> float | None:
    """Normalize relative strength to percentage points without double scaling."""
    if value is None:
        return None
    try:
        rs = float(value)
    except (TypeError, ValueError):
        return None
    if rs != rs:
        return None
    if abs(rs) <= 2:
        rs *= 100
    return round(rs, 2)


def _fetch_strength_snapshot_rows(symbols: list[str]) -> dict[str, dict]:
    if symbols:
        try:
            placeholders = ",".join(["%s"] * len(symbols))
            rows = _pg_fetchall(
                f"""
                SELECT symbol, company_name, stage, price, change_1m_pct, rsi,
                       relative_strength, can_slim_score, enhanced_fund_score,
                       financial_strength, investment_score, trading_signal, sector
                FROM scores.stage_snapshots
                WHERE snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)
                  AND UPPER(symbol) IN ({placeholders})
                """,
                [s.upper() for s in symbols],
            )
            cols = [
                "symbol", "company_name", "stage", "price", "change_1m_pct", "rsi",
                "relative_strength", "can_slim_score", "enhanced_fund_score",
                "financial_strength", "investment_score", "trading_signal", "sector",
            ]
            if rows:
                return {str(r[0]).upper(): dict(zip(cols, r)) for r in rows}
        except Exception:
            pass

    if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists() or not symbols:
        return {}
    conn = _db_conn()
    snap_date = _latest_snapshot_date()
    placeholders = ",".join("?" for _ in symbols)
    rows = conn.execute(
        f"""
        SELECT symbol, company_name, stage, price, change_1m_pct, rsi,
               relative_strength, can_slim_score, enhanced_fund_score,
               financial_strength, investment_score, trading_signal, sector
        FROM stage_snapshots
        WHERE snapshot_date=? AND UPPER(symbol) IN ({placeholders})
        """,
        [snap_date, *[s.upper() for s in symbols]],
    ).fetchall()
    conn.close()
    cols = [
        "symbol", "company_name", "stage", "price", "change_1m_pct", "rsi",
        "relative_strength", "can_slim_score", "enhanced_fund_score",
        "financial_strength", "investment_score", "trading_signal", "sector",
    ]
    return {str(r[0]).upper(): dict(zip(cols, r)) for r in rows}


def validate_strength_watchlist(symbols: list[str], top_n: int = 20) -> dict:
    """Validate CANSLIM/RS/fundamental/Piotroski strength from auditable sources."""
    syms = []
    for raw in symbols or []:
        sym = re.sub(r"[^A-Za-z0-9&-]", "", str(raw).upper())
        if sym and sym not in syms:
            syms.append(sym)
    syms = syms[: max(1, min(int(top_n or 20), 30))]
    if not syms:
        return {"error": "No symbols provided", "results": []}

    snap_date = _latest_snapshot_date()
    snapshots = _fetch_strength_snapshot_rows(syms)
    results = []

    for sym in syms:
        snap = snapshots.get(sym, {})
        forensic = run_forensic_analysis(sym)
        missing = []

        rs_pct = normalize_relative_strength_pct(snap.get("relative_strength"))
        if not snap:
            missing.append("stage_snapshot")
        if rs_pct is None:
            missing.append("relative_strength")
        for field in ("can_slim_score", "enhanced_fund_score", "financial_strength"):
            if snap.get(field) is None:
                missing.append(field)
        if forensic.get("error"):
            missing.append("forensic")

        piotroski = forensic.get("piotroski", {}) if not forensic.get("error") else {}
        beneish = forensic.get("beneish", {}) if not forensic.get("error") else {}
        altman = forensic.get("altman", {}) if not forensic.get("error") else {}
        risk = forensic.get("overall_risk") if not forensic.get("error") else "unknown"

        can = snap.get("can_slim_score")
        fund = snap.get("enhanced_fund_score")
        piot_score = piotroski.get("score")
        piot_max = piotroski.get("max_possible") or 9

        strength_score = None
        if not missing:
            risk_penalty = {"low": 0, "moderate": -10, "high": -25}.get(str(risk), -15)
            strength_score = round(
                0.25 * (float(can) / 25 * 100)
                + 0.25 * min(max(float(rs_pct), 0), 100)
                + 0.25 * float(fund)
                + 0.25 * (float(piot_score) / float(piot_max) * 100)
                + risk_penalty,
                1,
            )

        if missing:
            verdict = "Insufficient evidence to rank"
        elif risk == "high":
            verdict = "Caution: forensic risk overrides apparent strength"
        elif piot_score is not None and piot_score >= 7:
            verdict = "Validated strength"
        else:
            verdict = "Watchlist strength; confirm with fundamentals and risk controls"

        results.append({
            "symbol": sym,
            "company_name": snap.get("company_name"),
            "snapshot_date": snap_date,
            "stage": snap.get("stage"),
            "price": snap.get("price"),
            "change_1m_pct": snap.get("change_1m_pct"),
            "rsi": snap.get("rsi"),
            "trading_signal": snap.get("trading_signal"),
            "sector": snap.get("sector"),
            "investment_score": snap.get("investment_score"),
            "relative_strength": snap.get("relative_strength"),
            "rs_pct": rs_pct,
            "can_slim_score": can,
            "enhanced_fund_score": fund,
            "financial_strength": snap.get("financial_strength"),
            "piotroski_score": piot_score,
            "piotroski_max": piot_max if piot_score is not None else None,
            "piotroski_strength": piotroski.get("strength"),
            "beneish_score": beneish.get("score"),
            "beneish_interpretation": beneish.get("interpretation"),
            "altman_score": altman.get("score"),
            "altman_zone": altman.get("zone"),
            "overall_forensic_risk": risk,
            "forensic_error": forensic.get("error"),
            "missing_evidence": missing,
            "evidence_coverage": "complete" if not missing else "partial",
            "strength_score": strength_score,
            "verdict": verdict,
        })

    results.sort(key=lambda r: (r["strength_score"] is not None, r["strength_score"] or -999), reverse=True)
    return {
        "snapshot_date": snap_date,
        "input_symbols": syms,
        "count": len(results),
        "results": results,
        "validation_rule": "Never infer missing CANSLIM, RS, fundamentals, or forensic data; missing fields are reported in missing_evidence.",
    }


def _normalise_index_query(value: str) -> str:
    text = str(value or "").upper()
    text = re.sub(r"[^A-Z0-9]+", " ", text)
    words = []
    for word in text.split():
        if word == "INDEX":
            continue
        if word == "MKT":
            word = "MARKET"
        elif word == "MKTS":
            word = "MARKETS"
        words.append(word)
    return " ".join(words).strip()


def _index_reference_aliases() -> dict[str, str]:
    """Return normalized display-name aliases from PostgreSQL ref.indices."""
    global _INDEX_REFERENCE_ALIAS_CACHE
    if _INDEX_REFERENCE_ALIAS_CACHE is not None:
        return _INDEX_REFERENCE_ALIAS_CACHE

    aliases = {
        "NIFTY": "NIFTY 50",
        "NIFTY50": "NIFTY 50",
        "BANKNIFTY": "NIFTY BANK",
        "BANK NIFTY": "NIFTY BANK",
        "MIDCPNIFTY": "NIFTY MIDCAP SELECT",
        "MIDC NIFTY": "NIFTY MIDCAP SELECT",
        "MIDCAP NIFTY": "NIFTY MIDCAP SELECT",
        "NIFTY CAPITAL MARKET": "NIFTY CAPITAL MKT",
        "NIFTY CAPITAL MARKETS": "NIFTY CAPITAL MKT",
        "NIFTY SMALLCAP 50": "NIFTY SMLCAP 50",
        "NIFTY SMALLCAP 100": "NIFTY SMLCAP 100",
        "NIFTY SMALLCAP 250": "NIFTY SMLCAP 250",
        "NIFTY MICROCAP 250": "NIFTY MICROCAP250",
        "NIFTY CONSUMER DURABLES": "NIFTY CONSR DURBL",
        "NIFTY FINANCIAL SERVICES": "NIFTY FIN SERVICE",
        "NIFTY FINANCIAL SERVICES EX BANK": "NIFTY FINSEREXBNK",
        "NIFTY FINANCIAL SERVICES 25 50": "NIFTY FINSRV25 50",
        "NIFTY INFRASTRUCTURE": "NIFTY INFRA",
        "NIFTY MIDSMALL FINANCIAL SERVICES": "NIFTY MS FIN SERV",
        "NIFTY TRANSPORTATION LOGISTICS": "NIFTY TRANS LOGIS",
        "NIFTY ALPHA QUALITY LOW VOLATILITY 30": "NIFTY AQL 30",
        "NIFTY ALPHA QUALITY VALUE LOW VOLATILITY 30": "NIFTY AQLV 30",
        "NIFTY100 EQUAL WEIGHT": "NIFTY100 EQL WGT",
        "NIFTY50 EQUAL WEIGHT": "NIFTY50 EQL WGT",
        "NIFTY500 EQUAL WEIGHT": "NIFTY500 EW",
    }
    try:
        import psycopg2

        conn = psycopg2.connect("dbname=nse_market user=nse_admin host=/tmp")
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT index_symbol, display_name FROM ref.indices")
                for index_symbol, display_name in cur.fetchall():
                    if not index_symbol:
                        continue
                    target = str(index_symbol)
                    for name in (index_symbol, display_name, f"{display_name} INDEX" if display_name else None):
                        if name:
                            aliases[_normalise_index_query(str(name))] = target
        finally:
            conn.close()
    except Exception:
        pass

    _INDEX_REFERENCE_ALIAS_CACHE = aliases
    return aliases


def _index_match_score(query_norm: str, candidate_norm: str) -> float:
    if not query_norm or not candidate_norm:
        return 0.0
    q_tokens = set(query_norm.split())
    c_tokens = set(candidate_norm.split())
    if query_norm == candidate_norm:
        return 1.0
    if q_tokens <= {"NIFTY"}:
        return 0.0
    if query_norm in candidate_norm or candidate_norm in query_norm:
        return 0.92
    if not q_tokens or not c_tokens:
        return 0.0
    overlap = q_tokens & c_tokens
    # Never match on NIFTY alone; it caused unrelated index substitutions.
    meaningful_overlap = overlap - {"NIFTY"}
    alpha_overlap = {t for t in meaningful_overlap if any(ch.isalpha() for ch in t)}
    if not alpha_overlap:
        return 0.0
    if not meaningful_overlap:
        return 0.0
    precision = len(overlap) / len(q_tokens)
    recall = len(overlap) / len(c_tokens)
    return (2 * precision * recall / (precision + recall)) if (precision + recall) else 0.0


def get_index_snapshot(index_name: str = "NIFTY 50") -> dict:
    """Get latest index OHLCV and 10-day trend from index CSV."""
    if not INDEX_CSV.exists():
        return {"error": "Index CSV not available"}

    df = pd.read_csv(
        INDEX_CSV,
        usecols=["SYMBOL", "OPEN", "HIGH", "LOW", "CLOSE", "TIMESTAMP", "HI_52_WK", "LO_52_WK"],
        low_memory=False,
    )
    df["TIMESTAMP"] = pd.to_datetime(df["TIMESTAMP"])
    for c in ["OPEN", "HIGH", "LOW", "CLOSE"]:
        df[c] = pd.to_numeric(df[c], errors="coerce")

    # Fuzzy match index name, but never fall back to a single generic token like
    # "NIFTY"; that can silently substitute the wrong index.
    names = df["SYMBOL"].unique()
    q_original = _normalise_index_query(index_name)
    q = _normalise_index_query(_index_reference_aliases().get(q_original, q_original))
    candidates = sorted(
        (
            (_index_match_score(q, _normalise_index_query(n)), n)
            for n in names
        ),
        reverse=True,
    )
    match = candidates[0][1] if candidates and candidates[0][0] >= 0.45 else None
    if not match:
        similar = [
            n for score, n in candidates
            if score > 0
        ][:10]
        return {
            "error": f"Index '{index_name}' not found",
            "available": list(names[:20]),
            "similar": similar,
        }

    idx = df[df["SYMBOL"] == match].sort_values("TIMESTAMP")
    if idx.empty:
        return {"error": f"No data for {match}"}

    latest = idx.iloc[-1]
    prev   = idx.iloc[-2] if len(idx) > 1 else latest
    cur    = float(latest["CLOSE"])
    pr     = float(prev["CLOSE"])
    chg    = round((cur / pr - 1) * 100, 2)

    trend_10 = idx.tail(10)
    closes   = trend_10["CLOSE"].tolist()
    up_days  = sum(1 for i in range(1, len(closes)) if closes[i] >= closes[i-1])
    trend_chg = round((closes[-1] / closes[0] - 1) * 100, 2) if closes[0] > 0 else 0

    return {
        "index":       match,
        "as_of":       str(latest["TIMESTAMP"].date()),
        "close":       round(cur, 2),
        "open":        round(float(latest["OPEN"]), 2),
        "high":        round(float(latest["HIGH"]), 2),
        "low":         round(float(latest["LOW"]), 2),
        "chg_pct":     chg,
        "52w_high":    float(latest["HI_52_WK"]) if pd.notna(latest.get("HI_52_WK")) else None,
        "52w_low":     float(latest["LO_52_WK"]) if pd.notna(latest.get("LO_52_WK")) else None,
        "trend_10d":   {"closes": [round(c, 2) for c in closes], "up_days": up_days,
                        "chg_pct": trend_chg},
    }


def get_market_breadth() -> dict:
    """Compute market breadth: A/D ratio, %>200MA, sector overview."""
    snap_date = _latest_snapshot_date()

    try:
        rows = _pg_fetchall(
            "SELECT symbol, change_1d_pct, change_1w_pct, relative_strength "
            "FROM scores.stage_snapshots WHERE snapshot_date=%s",
            (snap_date,),
        )
        stage_dist = dict(_pg_fetchall(
            "SELECT stage, COUNT(*) FROM scores.stage_snapshots WHERE snapshot_date=%s GROUP BY stage",
            (snap_date,),
        ))
        data_source = "PostgreSQL scores.stage_snapshots"
    except Exception:
        if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
            return {"error": "PostgreSQL scores.stage_snapshots unavailable"}
        conn = _db_conn()
        rows = conn.execute(
            "SELECT symbol, change_1d_pct, change_1w_pct, relative_strength "
            "FROM stage_snapshots WHERE snapshot_date=?", (snap_date,)
        ).fetchall()
        conn.close()
        conn2 = _db_conn()
        stage_dist = dict(conn2.execute(
            "SELECT stage, COUNT(*) FROM stage_snapshots WHERE snapshot_date=? GROUP BY stage",
            (snap_date,)
        ).fetchall())
        conn2.close()
        data_source = "SQLite stage_snapshots"

    advances = sum(1 for r in rows if (r[1] or 0) > 0)
    declines = sum(1 for r in rows if (r[1] or 0) < 0)
    unchanged = len(rows) - advances - declines
    ad_ratio  = round(advances / declines, 2) if declines > 0 else 0.0
    rs_values = [
        v for v in (normalize_relative_strength_pct(r[3]) for r in rows)
        if v is not None
    ]
    avg_rs    = round(sum(rs_values) / len(rs_values), 1) if rs_values else 0

    return {
        "snapshot_date": snap_date,
        "total_stocks":  len(rows),
        "advances":      advances,
        "declines":      declines,
        "unchanged":     unchanged,
        "ad_ratio":      ad_ratio,
        "avg_rs_pct":    avg_rs,
        "stage_distribution": stage_dist,
        "data_source":    data_source,
    }


def get_global_market_assessment() -> dict:
    """Assess global market cues from cached global index and correlation files."""
    if not GLOBAL_INDEX_CSV.exists():
        return {"error": f"Global indices file not found: {GLOBAL_INDEX_CSV}"}

    df = pd.read_csv(GLOBAL_INDEX_CSV)
    if "Date" not in df.columns or df.empty:
        return {"error": "Global indices CSV is empty or missing Date column"}

    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df = df.dropna(subset=["Date"]).sort_values("Date")
    if len(df) < 2:
        return {"error": "Global indices CSV needs at least two dated rows"}

    asset_cols = [c for c in df.columns if c != "Date"]
    # Convert all asset columns to numeric once for efficiency
    num_df = df[asset_cols].apply(pd.to_numeric, errors="coerce")

    moves: dict[str, dict[str, Any]] = {}
    for asset in asset_cols:
        col = num_df[asset].dropna()
        if len(col) < 2:
            continue
        # Use the most-recent valid row for this asset (not necessarily df.iloc[-1])
        latest_idx = col.index[-1]
        prev_idx   = col.index[-2]
        latest_val = col.iloc[-1]
        prev_val   = col.iloc[-2]
        if prev_val == 0:
            continue

        pct_change = round((latest_val / prev_val - 1) * 100, 2)
        moves[asset] = {
            "price":      round(float(latest_val), 2),
            "pct_change": pct_change,
            "as_of":      str(df.loc[latest_idx, "Date"].date()),
        }

    def _avg(names: list[str]) -> float | None:
        vals = [moves[n]["pct_change"] for n in names if n in moves]
        return round(sum(vals) / len(vals), 2) if vals else None

    def _bias(avg: float | None) -> str:
        if avg is None:
            return "unknown"
        if avg > 0.35:
            return "positive"
        if avg < -0.35:
            return "negative"
        return "mixed"

    us_avg = _avg(["S&P 500", "Nasdaq"])
    asia_avg = _avg(["Hang Seng", "Nikkei 225"])
    commodity_avg = _avg(["Gold", "Crude Oil", "Copper"])
    fx_avg = _avg(["DXY", "USDINR"])

    regions = {
        "US": {"avg_pct_change": us_avg, "bias": _bias(us_avg)},
        "Asia": {"avg_pct_change": asia_avg, "bias": _bias(asia_avg)},
        "Commodities": {"avg_pct_change": commodity_avg, "bias": _bias(commodity_avg)},
        "FX": {"avg_pct_change": fx_avg, "bias": _bias(fx_avg)},
    }

    risk_points = 0
    if us_avg is not None:
        risk_points += 2 if us_avg > 1.0 else (1 if us_avg > 0.35 else (-2 if us_avg < -1.0 else (-1 if us_avg < -0.35 else 0)))
    if asia_avg is not None:
        risk_points += 1 if asia_avg > 0.35 else (-1 if asia_avg < -0.35 else 0)
    if moves.get("DXY", {}).get("pct_change", 0) > 0.35:
        risk_points -= 1
    if moves.get("Gold", {}).get("pct_change", 0) > 1 and risk_points < 1:
        risk_points -= 1

    if risk_points >= 2:
        risk_regime = "RISK_ON"
    elif risk_points <= -2:
        risk_regime = "RISK_OFF"
    else:
        risk_regime = "MIXED"

    india_readthrough: list[str] = []
    nasdaq = moves.get("Nasdaq", {}).get("pct_change")
    crude = moves.get("Crude Oil", {}).get("pct_change")
    usdinr = moves.get("USDINR", {}).get("pct_change")
    copper = moves.get("Copper", {}).get("pct_change")
    dxy = moves.get("DXY", {}).get("pct_change")

    if nasdaq is not None and nasdaq > 0.75:
        india_readthrough.append("Nasdaq strength is supportive for IT, digital, and growth-oriented Indian equities.")
    elif nasdaq is not None and nasdaq < -0.75:
        india_readthrough.append("Nasdaq weakness is a caution flag for IT and growth-oriented Indian equities.")

    if crude is not None and crude > 1.0:
        india_readthrough.append("Crude up is a headwind for oil importers, OMCs, aviation, paints, tyres, and broad inflation sentiment.")
    elif crude is not None and crude < -1.0:
        india_readthrough.append("Crude weakness can ease pressure on import-sensitive sectors and inflation expectations.")

    if usdinr is not None and usdinr > 0.2:
        india_readthrough.append("USDINR firmness can support exporters such as IT and pharma but pressures import-heavy sectors.")
    elif usdinr is not None and usdinr < -0.2:
        india_readthrough.append("USDINR softness can help importers but reduces currency tailwind for exporters.")

    if copper is not None and copper > 1.0:
        india_readthrough.append("Copper strength supports the cyclicals and metals read-through.")
    elif copper is not None and copper < -1.0:
        india_readthrough.append("Copper weakness is a caution flag for cyclicals and metals demand.")

    if dxy is not None and dxy > 0.35:
        india_readthrough.append("DXY strength can tighten global liquidity conditions for emerging markets, including India.")

    if not india_readthrough:
        india_readthrough.append("Global cues are balanced; confirm with Nifty breadth, Bank Nifty, and FII/DII flow.")

    watch_items = [
        "Nifty gap risk versus overnight US and Asia moves",
        "Bank Nifty follow-through after the first 30 minutes",
        "FII/DII flow confirmation",
        "Crude, USDINR, and DXY continuation during Indian market hours",
    ]

    correlations: list[dict[str, Any]] = []
    if GLOBAL_CORR_CSV.exists():
        corr = pd.read_csv(GLOBAL_CORR_CSV)
        keep = ["asset", "price", "corr_30d", "corr_60d", "change", "alert"]
        if not corr.empty and all(c in corr.columns for c in keep):
            for _, row in corr[keep].head(12).iterrows():
                correlations.append(row.where(pd.notna(row), None).to_dict())

    # Determine the "as_of" date: latest date that has any data
    last_dates = {asset: m["as_of"] for asset, m in moves.items()}
    as_of = max(last_dates.values()) if last_dates else str(df.iloc[-1]["Date"].date())

    return {
        "risk_regime": risk_regime,
        "as_of": as_of,
        "asset_dates": last_dates,   # per-asset last-available date (useful when markets are closed)
        "regions": regions,
        "moves": moves,
        "india_readthrough": india_readthrough,
        "watch_items": watch_items,
        "correlations": correlations,
        "source": "Cached global market CSVs",
        "source_files": {
            "global_indices": str(GLOBAL_INDEX_CSV.relative_to(ROOT)) if GLOBAL_INDEX_CSV.is_relative_to(ROOT) else str(GLOBAL_INDEX_CSV),
            "global_correlations": str(GLOBAL_CORR_CSV.relative_to(ROOT)) if GLOBAL_CORR_CSV.is_relative_to(ROOT) else str(GLOBAL_CORR_CSV),
        },
    }


def _sqlite_table_exists(table_name: str) -> bool:
    if not DB_PATH.exists():
        return False
    conn = _db_conn()
    row = conn.execute(
        "SELECT name FROM sqlite_master WHERE type='table' AND name=?",
        (table_name,),
    ).fetchone()
    conn.close()
    return bool(row)


def _pg_table_exists(schema: str, table_name: str) -> bool:
    try:
        rows = _pg_fetchall(
            """
            SELECT 1
            FROM information_schema.tables
            WHERE table_schema=%s AND table_name=%s
            """,
            (schema, table_name),
        )
        return bool(rows)
    except Exception:
        return False


def get_intraday_source_health(max_age_minutes: int = 30) -> dict:
    """Report health of PostgreSQL intraday source tables."""
    pg_tables = {
        "quote_snapshots": "captured_at",
        "ohlcv_bars": "timestamp",
        "scan_signals": "snapshot_ts",
    }
    try:
        result: dict[str, Any] = {
            "data_mode": "intraday",
            "db_path": PG_DSN,
            "max_age_minutes": max_age_minutes,
            "tables": {},
            "source": "PostgreSQL intraday schema",
        }
        now = datetime.now()
        statuses: list[str] = []
        for table, ts_col in pg_tables.items():
            exists = _pg_table_exists("intraday", table)
            if not exists:
                result["tables"][f"intraday.{table}"] = {"exists": False, "status": "MISSING"}
                statuses.append("MISSING")
                continue
            rows = _pg_fetchall(
                f"SELECT COUNT(*), MAX({ts_col})::text FROM intraday.{table}"
            )
            row_count, raw_ts = rows[0]
            latest_ts = None
            age_minutes = None
            status = "EMPTY" if row_count == 0 else "PRESENT"
            parsed = pd.to_datetime(raw_ts, errors="coerce") if raw_ts else None
            if row_count and parsed is not None and pd.notna(parsed):
                latest_dt = parsed.to_pydatetime()
                if latest_dt.tzinfo is not None:
                    latest_dt = latest_dt.replace(tzinfo=None)
                latest_ts = latest_dt.strftime("%Y-%m-%d %H:%M:%S")
                age_minutes = round((now - latest_dt).total_seconds() / 60, 1)
                status = "FRESH" if age_minutes <= max_age_minutes else "STALE"
            result["tables"][f"intraday.{table}"] = {
                "exists": True,
                "status": status,
                "rows": row_count,
                "latest_timestamp": latest_ts,
                "age_minutes": age_minutes,
            }
            statuses.append(status)
        ohlcv_status = result["tables"].get("intraday.ohlcv_bars", {}).get("status")
        if ohlcv_status:
            result["overall_status"] = ohlcv_status
        elif "MISSING" in statuses:
            result["overall_status"] = "MISSING"
        else:
            result["overall_status"] = "UNKNOWN"
        return result
    except Exception as exc:
        result = {
            "data_mode": "intraday",
            "db_path": PG_DSN,
            "max_age_minutes": max_age_minutes,
            "tables": {},
            "source": "PostgreSQL intraday schema",
            "overall_status": "MISSING",
            "error": f"PostgreSQL intraday schema unavailable: {exc}",
        }
        result["overall_status"] = "MISSING"
        return result


def get_intraday_bars(
    symbol: str,
    timeframe: str = "15m",
    lookback: int = 120,
) -> dict:
    """Read intraday OHLCV bars from PostgreSQL, seeding PG from yfinance if needed."""
    sym = symbol.strip().upper()

    def _fmt_bar_timestamp(value: Any) -> str:
        ts = pd.to_datetime(value)
        if getattr(ts, "tzinfo", None) is not None:
            ts = ts.tz_convert(IST_TZ).tz_localize(None)
        return ts.strftime("%Y-%m-%d %H:%M:%S")

    def _bars_from_df(df: pd.DataFrame, *, source: str, persisted: dict | None = None) -> dict:
        df_out = df.copy()
        if not isinstance(df_out.index, pd.DatetimeIndex):
            df_out.index = pd.to_datetime(df_out.index, errors="coerce")
        df_out = df_out.dropna().sort_index().tail(lookback)
        bars = [
            {
                "timestamp": _fmt_bar_timestamp(idx),
                "open": round(float(row["Open"]), 2),
                "high": round(float(row["High"]), 2),
                "low": round(float(row["Low"]), 2),
                "close": round(float(row["Close"]), 2),
                "volume": int(row["Volume"]) if pd.notna(row["Volume"]) else None,
            }
            for idx, row in df_out.iterrows()
        ]
        result = {
            "symbol": sym,
            "timeframe": timeframe,
            "lookback": lookback,
            "data_mode": "intraday",
            "source": source,
            "count": len(bars),
            "latest_timestamp": bars[-1]["timestamp"] if bars else None,
            "bars": bars,
        }
        if persisted:
            result["postgres_persist"] = persisted
        return result

    try:
        df = _pg_read_df(
            """
            SELECT timestamp, open, high, low, close, volume
            FROM intraday.ohlcv_bars
            WHERE UPPER(symbol)=%s AND timeframe=%s
            ORDER BY timestamp DESC
            LIMIT %s
            """,
            (sym, timeframe, lookback),
        )
        if not df.empty:
            df = df.sort_values("timestamp")
            bars = [
                {
                    "timestamp": _fmt_bar_timestamp(row["timestamp"]),
                    "open": round(float(row["open"]), 2),
                    "high": round(float(row["high"]), 2),
                    "low": round(float(row["low"]), 2),
                    "close": round(float(row["close"]), 2),
                    "volume": int(row["volume"]) if pd.notna(row["volume"]) else None,
                }
                for _, row in df.iterrows()
            ]
            return {
                "symbol": sym,
                "timeframe": timeframe,
                "lookback": lookback,
                "data_mode": "intraday",
                "source": "PostgreSQL intraday.ohlcv_bars",
                "count": len(bars),
                "latest_timestamp": bars[-1]["timestamp"] if bars else None,
                "bars": bars,
            }
    except Exception:
        pass

    try:
        yf_df = get_intraday_candles(sym, timeframe)
        if not yf_df.empty:
            seed_bars = [
                {
                    "timestamp": _fmt_bar_timestamp(idx),
                    "open": row["Open"],
                    "high": row["High"],
                    "low": row["Low"],
                    "close": row["Close"],
                    "volume": row["Volume"],
                }
                for idx, row in yf_df.iterrows()
            ]
            persisted = persist_intraday_bars(
                sym,
                seed_bars,
                timeframe=timeframe,
                source="Yahoo Finance (yfinance)",
            )
            return _bars_from_df(
                yf_df,
                source="PostgreSQL intraday.ohlcv_bars seeded from Yahoo Finance (yfinance)",
                persisted=persisted,
            )
    except Exception as exc:
        seed_error = str(exc)
    else:
        seed_error = "No yfinance candles available to seed PostgreSQL intraday.ohlcv_bars"

    return {
        "symbol": sym,
        "timeframe": timeframe,
        "data_mode": "intraday",
        "source": "PostgreSQL intraday.ohlcv_bars",
        "error": f"No PostgreSQL intraday.ohlcv_bars for {sym} at {timeframe}; {seed_error}",
    }


def _bars_to_intraday_df(bars: list[dict]) -> pd.DataFrame:
    df = pd.DataFrame(bars)
    if df.empty:
        return df
    df["timestamp"] = pd.to_datetime(df["timestamp"], errors="coerce")
    df = df.dropna(subset=["timestamp"]).set_index("timestamp").sort_index()
    return df.rename(
        columns={
            "open": "Open",
            "high": "High",
            "low": "Low",
            "close": "Close",
            "volume": "Volume",
        }
    )[["Open", "High", "Low", "Close", "Volume"]]


def get_intraday_levels(symbol: str, timeframe: str = "15m") -> dict:
    """Compute support/resistance levels from PostgreSQL intraday OHLCV bars."""
    sym = symbol.strip().upper()
    bars_result = get_intraday_bars(sym, timeframe=timeframe, lookback=240)
    if bars_result.get("error"):
        return bars_result

    df = _bars_to_intraday_df(bars_result.get("bars", []))
    if df.empty or len(df) < 10:
        return {
            "symbol": sym,
            "timeframe": timeframe,
            "data_mode": "intraday",
            "source": bars_result.get("source", "intraday bars"),
            "error": "Insufficient PostgreSQL intraday bars for level analysis",
        }

    df_ind = _compute_intraday_all(df)
    levels = _intraday_key_levels(df_ind)
    latest_close = float(df_ind["Close"].iloc[-1])
    pivot_levels = {
        str(k): _safe_float(v)
        for k, v in (levels.get("pivot_levels") or {}).items()
    }
    return {
        "symbol": sym,
        "timeframe": timeframe,
        "data_mode": "intraday",
        "source": bars_result.get("source", "intraday bars"),
        "latest_timestamp": bars_result.get("latest_timestamp"),
        "latest_close": round(latest_close, 2),
        "pivot": _safe_float(levels.get("pivot")),
        "supports": [_safe_float(v) for v in levels.get("supports", []) if _safe_float(v) is not None],
        "resistances": [_safe_float(v) for v in levels.get("resistances", []) if _safe_float(v) is not None],
        "ema_levels": {
            "ema9": _safe_float(levels.get("ema9")),
            "ema21": _safe_float(levels.get("ema21")),
            "ema50": _safe_float(levels.get("ema50")),
            "ema200": _safe_float(levels.get("ema200")),
        },
        "pivot_levels": pivot_levels,
        "copy_rule": "Technical levels only. Not investment advice or a trade recommendation.",
    }


def _normalise_signal_direction(direction: str | None) -> str:
    d = (direction or "").upper()
    if d == "BUY":
        return "LONG_SETUP"
    if d == "SELL":
        return "SHORT_SETUP"
    if d == "WATCH":
        return "WATCH"
    return "AVOID"


def compute_intraday_indicators(symbol: str, timeframe: str = "15m") -> dict:
    """Compute latest intraday indicators from PostgreSQL intraday OHLCV bars."""
    sym = symbol.strip().upper()
    bars_result = get_intraday_bars(sym, timeframe=timeframe, lookback=240)
    if bars_result.get("error"):
        return bars_result

    df = _bars_to_intraday_df(bars_result.get("bars", []))
    if df.empty or len(df) < 30:
        return {
            "symbol": sym,
            "timeframe": timeframe,
            "data_mode": "intraday",
            "source": bars_result.get("source", "intraday bars"),
            "error": "Insufficient PostgreSQL intraday bars for indicator calculation",
        }

    df_ind = _compute_intraday_all(df)
    last = df_ind.iloc[-1]
    first_close = float(df_ind["Close"].iloc[0])
    latest_close = float(last["Close"])
    volume_avg_20 = float(df_ind["Volume"].tail(20).mean()) if len(df_ind) >= 20 else None
    volume_last = float(last["Volume"]) if pd.notna(last["Volume"]) else None
    volume_ratio = round(volume_last / volume_avg_20, 2) if volume_avg_20 and volume_last else None
    roc = round((latest_close / first_close - 1) * 100, 2) if first_close else None
    macd_hist = float(last["MACD_hist"]) if pd.notna(last["MACD_hist"]) else 0.0
    supertrend_dir = int(last["Supertrend_dir"]) if pd.notna(last["Supertrend_dir"]) else 0

    momentum_score = 0
    momentum_score += 20 if latest_close > float(last["EMA21"]) else -10
    momentum_score += 20 if latest_close > float(last["EMA50"]) else -10
    momentum_score += 15 if macd_hist > 0 else -10
    rsi = float(last["RSI"]) if pd.notna(last["RSI"]) else 50.0
    momentum_score += 15 if 50 <= rsi <= 75 else (-10 if rsi > 80 or rsi < 35 else 0)
    momentum_score += 15 if supertrend_dir == 1 else -10
    momentum_score += 15 if (volume_ratio or 0) >= 1.2 else 0
    score = max(0, min(100, 50 + momentum_score))

    return {
        "symbol": sym,
        "timeframe": timeframe,
        "data_mode": "intraday",
        "source": bars_result.get("source", "intraday bars"),
        "latest_timestamp": bars_result.get("latest_timestamp"),
        "latest_close": round(latest_close, 2),
        "pct_change": roc,
        "bars": len(df_ind),
        "score": round(score, 1),
        "indicators": {
            "rsi": _safe_float(last["RSI"], 1),
            "macd": _safe_float(last["MACD"], 4),
            "macd_signal": _safe_float(last["MACD_signal"], 4),
            "macd_hist": _safe_float(last["MACD_hist"], 4),
            "supertrend": _safe_float(last["Supertrend"]),
            "supertrend_dir": supertrend_dir,
            "ema9": _safe_float(last["EMA9"]),
            "ema21": _safe_float(last["EMA21"]),
            "ema50": _safe_float(last["EMA50"]),
            "ema200": _safe_float(last["EMA200"]),
            "bb_pct": _safe_float(last["BB_pct"]),
            "bb_width": _safe_float(last["BB_width"], 4),
            "atr": _safe_float(last["ATR"]),
            "volume_ratio": volume_ratio,
        },
    }


def _compute_intraday_position_sizing(
    symbol: str,
    entry_price: float,
    stoploss_price: float,
    risk_per_trade: float = 5000.0,
) -> dict:
    """Position sizing for intraday setups based on a fixed risk budget."""
    risk_per_share = abs(entry_price - stoploss_price)
    if risk_per_share < 0.01:
        return {"error": "Risk per share is effectively zero; cannot size position"}

    shares = int(risk_per_trade / risk_per_share)
    capital = round(shares * entry_price, 2)

    lot_size = get_lot_size(symbol)
    futures = None
    if lot_size and lot_size > 0:
        risk_per_lot = round(risk_per_share * lot_size, 2)
        lots = max(1, int(risk_per_trade / risk_per_lot)) if risk_per_lot > 0 else 1
        margin_approx = round(entry_price * lot_size * 0.12, 2)
        futures = {
            "lots": lots,
            "lot_size": lot_size,
            "units": lots * lot_size,
            "risk_per_lot": risk_per_lot,
            "approx_margin_per_lot": margin_approx,
        }
        options_note = (
            f"With lot size {lot_size}, {lots} lot(s) keeps risk near "
            f"\u20b9{risk_per_trade:,.0f}. Actual premium-based sizing "
            f"requires live option chain data."
        )
    else:
        options_note = "Not an F&O symbol; options/futures not available."

    return {
        "risk_per_trade": risk_per_trade,
        "risk_per_share": round(risk_per_share, 2),
        "lot_size": lot_size,
        "cash": {"shares": shares, "capital_required": capital},
        "futures": futures,
        "options_note": options_note,
    }


def _build_trade_plan(
    setup_label: str,
    entry_price: float | None,
    invalidation: float | None,
    target_zones: list[float],
    indicators: dict,
) -> dict:
    """Build a direction-aware trade plan from setup data."""
    if setup_label not in ("LONG_SETUP", "SHORT_SETUP"):
        return {}
    if not isinstance(entry_price, (int, float)) or not isinstance(invalidation, (int, float)):
        return {}

    is_long = setup_label == "LONG_SETUP"
    direction = "LONG" if is_long else "SHORT"

    confirmations: list[str] = []
    if is_long:
        confirmations.append(f"Price holds above \u20b9{invalidation:,.2f} on retest")
    else:
        confirmations.append(f"Price stays below \u20b9{invalidation:,.2f} on retest")

    st_dir = indicators.get("supertrend_dir")
    if st_dir == (1 if is_long else -1):
        confirmations.append("Supertrend confirms " + ("bullish" if is_long else "bearish"))

    rsi = indicators.get("rsi")
    if isinstance(rsi, (int, float)):
        if is_long and rsi > 50:
            confirmations.append(f"RSI {rsi:.0f} supports momentum (above 50)")
        elif not is_long and rsi < 50:
            confirmations.append(f"RSI {rsi:.0f} confirms weakness (below 50)")

    macd_hist = indicators.get("macd_hist")
    if isinstance(macd_hist, (int, float)):
        if (is_long and macd_hist > 0) or (not is_long and macd_hist < 0):
            confirmations.append("MACD histogram aligned with direction")

    scale_out: list[str] = []
    t1 = target_zones[0] if len(target_zones) > 0 else None
    t2 = target_zones[1] if len(target_zones) > 1 else None
    if t1:
        scale_out.append(f"Book 50% at T1 (\u20b9{t1:,.2f}), trail SL to entry for remainder")
    if t2:
        scale_out.append(f"Book remaining at T2 (\u20b9{t2:,.2f}) or trail with Supertrend")
    elif t1:
        scale_out.append("Trail remainder with Supertrend or EMA21")

    word = "below" if is_long else "above"
    inv_action = (
        f"Exit full position if price closes {word} \u20b9{invalidation:,.2f}. "
        f"Do not average into a losing setup."
    )

    return {
        "direction": direction,
        "entry_confirmations": confirmations,
        "scale_out": scale_out,
        "invalidation_action": inv_action,
    }


def explain_intraday_setup(symbol: str, timeframe: str = "15m") -> dict:
    """Explain an intraday setup from PostgreSQL intraday bars with research-only labels."""
    sym = symbol.strip().upper()
    ind = compute_intraday_indicators(sym, timeframe=timeframe)
    if ind.get("error"):
        return ind

    bars_result = get_intraday_bars(sym, timeframe=timeframe, lookback=240)
    df = _bars_to_intraday_df(bars_result.get("bars", []))
    df_ind = _compute_intraday_all(df)
    raw_signals = _run_intraday_all_signals(df_ind)
    levels = get_intraday_levels(sym, timeframe=timeframe)
    indicators = ind.get("indicators", {})

    long_evidence: list[str] = []
    short_evidence: list[str] = []
    watch_evidence: list[str] = []

    if (indicators.get("supertrend_dir") or 0) == 1:
        long_evidence.append("Supertrend is bullish")
    elif (indicators.get("supertrend_dir") or 0) == -1:
        short_evidence.append("Supertrend is bearish")

    if (indicators.get("macd_hist") or 0) > 0:
        long_evidence.append("MACD histogram is positive")
    elif (indicators.get("macd_hist") or 0) < 0:
        short_evidence.append("MACD histogram is negative")

    rsi = indicators.get("rsi")
    if isinstance(rsi, (int, float)):
        if 50 <= rsi <= 75:
            long_evidence.append(f"RSI {rsi} supports momentum without extreme overextension")
        elif rsi > 80:
            watch_evidence.append(f"RSI {rsi} is extended")
        elif rsi < 40:
            short_evidence.append(f"RSI {rsi} shows weak momentum")

    latest_close = ind.get("latest_close")
    ema21 = indicators.get("ema21")
    ema50 = indicators.get("ema50")
    if isinstance(latest_close, (int, float)) and isinstance(ema21, (int, float)):
        if latest_close > ema21:
            long_evidence.append("Price is above EMA21")
        else:
            short_evidence.append("Price is below EMA21")
    if isinstance(latest_close, (int, float)) and isinstance(ema50, (int, float)):
        if latest_close > ema50:
            long_evidence.append("Price is above EMA50")
        else:
            short_evidence.append("Price is below EMA50")

    signal_labels = [_normalise_signal_direction(s.get("direction")) for s in raw_signals]
    if "LONG_SETUP" in signal_labels:
        long_evidence.append("At least one indicator strategy triggered a long-side setup")
    if "SHORT_SETUP" in signal_labels:
        short_evidence.append("At least one indicator strategy triggered a short-side setup")
    if "WATCH" in signal_labels:
        watch_evidence.append("At least one volatility or pattern alert is in watch state")

    long_score = len(long_evidence) * 12
    short_score = len(short_evidence) * 12
    score = ind.get("score", 0)
    if long_score >= 36 and long_score >= short_score:
        setup_label = "LONG_SETUP"
    elif short_score >= 36 and short_score > long_score:
        setup_label = "SHORT_SETUP"
    elif watch_evidence or score >= 55:
        setup_label = "WATCH"
    else:
        setup_label = "AVOID"

    supports = levels.get("supports") or []
    resistances = levels.get("resistances") or []
    atr = indicators.get("atr") or 0
    invalidation = None
    target_zones: list[float] = []
    if setup_label == "LONG_SETUP":
        invalidation = supports[0] if supports else (round(latest_close - atr, 2) if latest_close and atr else None)
        target_zones = resistances[:2] or ([round(latest_close + atr, 2)] if latest_close and atr else [])
    elif setup_label == "SHORT_SETUP":
        invalidation = resistances[0] if resistances else (round(latest_close + atr, 2) if latest_close and atr else None)
        target_zones = supports[:2] or ([round(latest_close - atr, 2)] if latest_close and atr else [])

    supports = [v for v in (_safe_float(v) for v in supports) if v is not None]
    resistances = [v for v in (_safe_float(v) for v in resistances) if v is not None]
    invalidation = _safe_float(invalidation)
    target_zones = [v for v in (_safe_float(v) for v in target_zones) if v is not None]

    analyzers = {
        "TrendAgent": long_evidence[:2] if long_evidence else short_evidence[:2],
        "MomentumAgent": [e for e in long_evidence + short_evidence + watch_evidence if "RSI" in e or "MACD" in e],
        "LevelsAgent": {
            "support": supports[:2],
            "resistance": resistances[:2],
            "invalidation_level": invalidation,
            "technical_target_zones": target_zones,
        },
        "RiskAgent": [
            "Technical setup only; no order placement or recommendation",
            "Confirm freshness, liquidity, spreads, and market regime before using this for research",
        ],
    }

    position_sizing: dict = {}
    rr_frame: dict = {}
    trade_plan: dict = {}
    if (
        isinstance(latest_close, (int, float))
        and isinstance(invalidation, (int, float))
        and latest_close > 0
    ):
        position_sizing = _compute_intraday_position_sizing(
            sym, latest_close, invalidation
        )
        if not position_sizing.get("error"):
            rps = position_sizing["risk_per_share"]
            shares = position_sizing["cash"]["shares"]
            rupee_risk = round(rps * shares, 2)
            t1_val = target_zones[0] if target_zones else None
            t2_val = target_zones[1] if len(target_zones) > 1 else None
            rr_frame = {"risk_per_share": rps, "rupee_risk": rupee_risk}
            if isinstance(t1_val, (int, float)):
                t1_reward_per_share = abs(t1_val - latest_close)
                rr_frame["t1_rr"] = round(t1_reward_per_share / rps, 2) if rps > 0 else 0
                rr_frame["t1_rupee_reward"] = round(t1_reward_per_share * shares, 2)
            if isinstance(t2_val, (int, float)):
                t2_reward_per_share = abs(t2_val - latest_close)
                rr_frame["t2_rr"] = round(t2_reward_per_share / rps, 2) if rps > 0 else 0
                rr_frame["t2_rupee_reward"] = round(t2_reward_per_share * shares, 2)

        trade_plan = _build_trade_plan(
            setup_label, latest_close, invalidation, target_zones, indicators
        )

    return {
        "symbol": sym,
        "timeframe": timeframe,
        "data_mode": "intraday",
        "source": ind.get("source", "PostgreSQL intraday.ohlcv_bars"),
        "latest_timestamp": ind.get("latest_timestamp"),
        "latest_close": latest_close,
        "setup_label": setup_label,
        "setup_side": "long" if setup_label == "LONG_SETUP" else ("short" if setup_label == "SHORT_SETUP" else "neutral"),
        "score": score,
        "indicators": indicators,
        "signals": [
            {
                **s,
                "setup_label": _normalise_signal_direction(s.get("direction")),
                "copy_rule": "Research setup only; not a buy/sell recommendation.",
            }
            for s in raw_signals
        ],
        "analyzers": analyzers,
        "levels": levels,
        "invalidation_level": invalidation,
        "technical_target_zones": target_zones,
        "position_sizing": position_sizing,
        "risk_reward_frame": rr_frame,
        "trade_plan": trade_plan,
        "disclaimer": (
            "Research and learning only. Not investment advice. Not a recommendation to buy, "
            "sell, trade, copy, or replicate. Users are responsible for their own risk, "
            "compliance, and legal obligations. Agent Adda is not SEBI registered."
        ),
    }


def run_intraday_screener(
    screen_type: str = "momentum",
    timeframe: str = "15m",
    min_score: float = 55,
    top_n: int = 10,
    symbols: list[str] | None = None,
) -> dict:
    """Run an intraday screener (PostgreSQL-backed or live yfinance fallback).

    Screen types:
      Original : momentum, breakouts, vcp, supertrend, levels, all
      New      : opening_range_breakout, gap_and_go, macd_crossover,
                 rsi_divergence, bb_squeeze, vwap_reclaim
    """
    screen_key = screen_type.lower().strip()
    supported = {
        # original
        "momentum", "breakouts", "vcp", "supertrend", "levels", "all",
        # new
        "opening_range_breakout", "gap_and_go", "macd_crossover",
        "rsi_divergence", "bb_squeeze", "vwap_reclaim",
    }
    if screen_key not in supported:
        return {"error": f"Unknown intraday screener: {screen_type}", "supported": sorted(supported)}

    pg_has_intraday = False
    if _pg_table_exists("intraday", "ohlcv_bars"):
        try:
            pg_rows = _pg_fetchall(
                "SELECT COUNT(*) FROM intraday.ohlcv_bars WHERE timeframe=%s",
                (timeframe,),
            )
            pg_has_intraday = bool(pg_rows and pg_rows[0][0])
        except Exception:
            pg_has_intraday = False
        if symbols is not None:
            # Explicit symbol scans should use the per-symbol intraday path even
            # when the table is currently empty. explain_intraday_setup can
            # seed/report each requested symbol; the broad yfinance fallback is
            # only for universe scans with no explicit symbol list.
            pg_has_intraday = True
    # ── Intraday store unavailable → live yfinance fallback ─────────────────
    if not pg_has_intraday:
        strategy_map = {
            # original
            "breakouts":              ["ema", "volume", "macd"],
            "momentum":               ["macd", "rsi", "supertrend"],
            "vcp":                    ["vcp", "volume"],
            "supertrend":             ["supertrend"],
            "levels":                 ["ema", "bollinger"],
            "all":                    None,
            # new intraday screeners
            "opening_range_breakout": ["ema", "volume"],   # ORB = EMA crossover + volume
            "gap_and_go":             ["volume", "macd"],  # gap continuation via vol + MACD
            "macd_crossover":         ["macd"],            # pure MACD signal line cross
            "rsi_divergence":         ["rsi", "bollinger"],# RSI extreme + Bollinger
            "bb_squeeze":             ["bollinger", "volume"],  # BB bandwidth squeeze
            "vwap_reclaim":           ["ema", "rsi"],      # VWAP proxy via short EMA + RSI
        }

        _descriptions = {
            "opening_range_breakout": "Opening Range Breakout — price breaks above/below first 15-30min range with volume",
            "gap_and_go":             "Gap & Go — stocks gapping up/down continuing with volume + MACD momentum",
            "macd_crossover":         "MACD Crossover — fresh MACD signal line cross (bullish or bearish)",
            "rsi_divergence":         "RSI Divergence — RSI extreme reversal setup with Bollinger mean-reversion",
            "bb_squeeze":             "Bollinger Squeeze — low-volatility band squeeze followed by expansion breakout",
            "vwap_reclaim":           "VWAP Reclaim — price reclaiming short EMA (VWAP proxy) from below (bullish) or losing it (bearish)",
            "momentum":               "MACD + RSI + Supertrend momentum alignment",
            "breakouts":              "EMA crossover + volume + MACD breakout",
            "vcp":                    "Volatility Contraction Pattern — tight consolidation before expansion",
            "supertrend":             "Supertrend state flip or continuation",
            "levels":                 "Support/Resistance levels with EMA + Bollinger context",
            "all":                    "All strategies combined",
        }

        strategies = strategy_map.get(screen_key)
        result = scan_intraday_market(
            index="NIFTY 500",
            interval=timeframe if timeframe in ("5m","15m","30m","1h") else "15m",
            strategies=strategies,
            direction_filter="buy",
            min_rr=1.3,
            top_n=top_n,
        )
        result["screen_type"]    = screen_key
        result["description"]    = _descriptions.get(screen_key, "")
        result["data_mode"]      = "live-yfinance-fallback"
        result["data_source"]    = "NSE website constituents + yfinance candles"
        result["source_priority"] = ["PostgreSQL intraday.ohlcv_bars", "NSE website live constituents", "yfinance candles"]
        result["fallback_note"]  = (
            "PostgreSQL intraday bars not available; fetched NSE website constituents first, "
            "then used yfinance only for intraday candle history."
        )
        return result

    if symbols is None:
        if pg_has_intraday:
            try:
                rows = _pg_fetchall(
                    "SELECT DISTINCT UPPER(symbol) FROM intraday.ohlcv_bars WHERE timeframe=%s ORDER BY UPPER(symbol)",
                    (timeframe,),
                )
                symbols = [r[0] for r in rows]
            except Exception:
                symbols = []

    results: list[dict[str, Any]] = []
    errors: list[dict[str, str]] = []
    for sym in symbols:
        setup = explain_intraday_setup(sym, timeframe=timeframe)
        if setup.get("error"):
            errors.append({"symbol": sym, "error": setup["error"]})
            continue
        label = setup.get("setup_label")
        indicators = setup.get("indicators", {})
        include = setup.get("score", 0) >= min_score
        if screen_key == "breakouts":
            include = include and label in ("LONG_SETUP", "WATCH") and bool(setup.get("technical_target_zones"))
        elif screen_key == "vcp":
            include = any(s.get("strategy_key") == "vcp" for s in setup.get("signals", [])) or label == "WATCH"
        elif screen_key == "supertrend":
            include = include and indicators.get("supertrend_dir") in (1, -1)
        elif screen_key == "levels":
            include = True

        if include:
            levels = setup.get("levels", {})
            results.append({
                "symbol": setup["symbol"],
                "timeframe": timeframe,
                "setup_label": label,
                "setup_side": setup.get("setup_side"),
                "score": setup.get("score"),
                "price": setup.get("latest_close"),
                "rsi": indicators.get("rsi"),
                "macd_hist": indicators.get("macd_hist"),
                "supertrend_dir": indicators.get("supertrend_dir"),
                "support": (levels.get("supports") or [None])[0],
                "resistance": (levels.get("resistances") or [None])[0],
                "invalidation_level": setup.get("invalidation_level"),
                "technical_target_zones": setup.get("technical_target_zones"),
                "freshness": setup.get("latest_timestamp"),
            })

    results.sort(key=lambda r: (r.get("score") or 0), reverse=True)
    return {
        "screen_type": screen_key,
        "timeframe": timeframe,
        "data_mode": "intraday",
        "source": "PostgreSQL intraday.ohlcv_bars",
        "min_score": min_score,
        "scanned": len(symbols),
        "count": len(results[:top_n]),
        "results": results[:top_n],
        "errors": errors,
        "disclaimer": "Research and learning only. Not investment advice or a trade recommendation.",
    }


def get_data_health() -> dict:
    """Check freshness of all data sources."""
    today = date.today()

    def _days_old(path: Path, date_col: str | None = None) -> dict:
        if not path.exists():
            return {"exists": False}
        stat = path.stat()
        mtime = datetime.fromtimestamp(stat.st_mtime).date()
        return {"exists": True, "file_mtime": str(mtime), "mtime_days_old": (today - mtime).days,
                "size_mb": round(stat.st_size / 1e6, 2)}

    stock_info = _days_old(STOCK_CSV)
    index_info = _days_old(INDEX_CSV)
    db_info    = _days_old(DB_PATH)

    snap_date = _latest_snapshot_date()
    db_age = (today - date.fromisoformat(snap_date)).days if snap_date != "N/A" else -1

    return {
        "as_of":              str(today),
        "stock_csv":          stock_info,
        "index_csv":          index_info,
        "tracker_db":         {**db_info, "latest_snapshot": snap_date, "snapshot_days_old": db_age},
        "overall_status":     "FRESH" if db_age <= 3 else ("STALE" if db_age <= 7 else "OLD"),
    }


def find_latest_report(report_type: str = "any") -> dict:
    """List available generated reports."""
    report_dirs = [
        REPORTS / "latest",
        REPORTS / "generated",
        REPORTS / "generated_csv",
        REPORTS / "strategy_council",
    ]
    files: list[dict] = []
    for d in report_dirs:
        if not d.exists():
            continue
        for f in sorted(d.iterdir(), reverse=True):
            if f.is_file() and f.suffix in (".html", ".csv", ".json", ".md", ".pdf"):
                keyword = report_type.lower()
                if keyword == "any" or keyword in f.name.lower():
                    files.append({
                        "name": f.name,
                        "path": str(f.relative_to(ROOT)),
                        "size_kb": round(f.stat().st_size / 1024, 1),
                        "modified": str(datetime.fromtimestamp(f.stat().st_mtime).date()),
                    })
    return {"report_type": report_type, "count": len(files), "files": files[:10]}


def search_latest_catalysts(symbol: str, max_results: int = 5) -> dict:
    """Search for recent news/catalysts for a symbol via DuckDuckGo Lite."""
    try:
        import requests
        from html.parser import HTMLParser

        class _ResultParser(HTMLParser):
            def __init__(self):
                super().__init__()
                self.results: list[dict] = []
                self._in_result = False
                self._cur: dict = {}
                self._tag_stack: list[str] = []

            def handle_starttag(self, tag, attrs):
                attrs_d = dict(attrs)
                self._tag_stack.append(tag)
                if tag == "a" and attrs_d.get("class") == "result__a":
                    self._cur = {"url": attrs_d.get("href", ""), "title": ""}
                    self._in_result = True
                elif tag == "td" and "result__snippet" in attrs_d.get("class", ""):
                    self._in_result = True

            def handle_data(self, data):
                if self._in_result and data.strip():
                    if "title" in self._cur and not self._cur["title"]:
                        self._cur["title"] = data.strip()
                    elif "snippet" not in self._cur:
                        self._cur["snippet"] = data.strip()

            def handle_endtag(self, tag):
                if self._tag_stack:
                    self._tag_stack.pop()
                if tag == "a" and self._in_result and self._cur.get("title"):
                    self.results.append(dict(self._cur))
                    self._cur = {}
                    self._in_result = False

        # Get company name for better search
        sym_data = get_symbol_snapshot(symbol)
        company  = sym_data.get("company_name") or symbol
        query    = f"{company} NSE India news 2026"

        import urllib.parse
        url  = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
        resp = requests.get(url, headers={"User-Agent": "Mozilla/5.0"}, timeout=8)

        parser = _ResultParser()
        parser.feed(resp.text)

        def _decode_url(raw: str) -> str:
            """Extract real URL from DuckDuckGo redirect (/l/?uddg=<encoded>).
            Returns '' for DDG ad/tracking URLs (y.js) so they get filtered out."""
            import urllib.parse
            if not raw:
                return ""
            # Add scheme if missing
            if raw.startswith("//"):
                raw = "https:" + raw
            # DDG ad tracking URLs — opaque Bing redirect chains, not useful results
            if "duckduckgo.com/y.js" in raw:
                return ""
            parsed = urllib.parse.urlparse(raw)
            qs = urllib.parse.parse_qs(parsed.query)
            if "uddg" in qs:
                return qs["uddg"][0]
            return raw

        results = []
        for r in parser.results[:max_results]:
            if r.get("title") and len(r["title"]) > 5:
                decoded_url = _decode_url(r.get("url", ""))
                if not decoded_url:  # skip DDG ad/tracking URLs
                    continue
                results.append({
                    "title":   r.get("title", ""),
                    "url":     decoded_url,
                    "snippet": r.get("snippet", ""),
                })

        # Fetch article content for top results — gives LLM real text to summarize
        for item in results[:3]:
            url = item.get("url", "")
            if not url:
                continue
            try:
                art_resp = requests.get(
                    url,
                    headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"},
                    timeout=6,
                    allow_redirects=True,
                )
                if art_resp.status_code == 200 and "text/html" in art_resp.headers.get("Content-Type", ""):
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(art_resp.text, "lxml")
                    # Remove noise: scripts, styles, nav, ads, footers
                    for tag in soup(["script", "style", "nav", "footer", "header",
                                     "aside", "iframe", "form", "noscript"]):
                        tag.decompose()
                    # Extract article body — try common selectors first
                    body = (
                        soup.select_one("article") or
                        soup.select_one('[class*="article"]') or
                        soup.select_one('[class*="story"]') or
                        soup.select_one('[class*="content"]') or
                        soup.select_one("main") or
                        soup.body
                    )
                    if body:
                        text = body.get_text(separator="\n", strip=True)
                        # Clean up: collapse whitespace, take first ~2000 chars
                        import re as _re
                        text = _re.sub(r'\n{3,}', '\n\n', text)
                        text = _re.sub(r'[ \t]{2,}', ' ', text)
                        item["article_text"] = text[:2000]
            except Exception:
                pass  # Fetch failed — snippet is still available

        return {
            "symbol":  symbol.upper(),
            "company": company,
            "query":   query,
            "results": results,
            "source":  "DuckDuckGo",
            "disclaimer": "Web search results — verify before acting",
        }
    except Exception as e:
        return {"symbol": symbol.upper(), "error": str(e), "results": []}


def fetch_article_content(url: str, max_chars: int = 3000) -> dict:
    """Fetch and extract readable text from a news article URL.

    Strips navigation, ads, footers — returns clean article body text.
    Use this to read full articles found via search tools when you need
    deeper context for summary and opinion.
    """
    import requests
    try:
        resp = requests.get(
            url,
            headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)"},
            timeout=8,
            allow_redirects=True,
        )
        if resp.status_code != 200:
            return {"url": url, "error": f"HTTP {resp.status_code}"}
        if "text/html" not in resp.headers.get("Content-Type", ""):
            return {"url": url, "error": "Not an HTML page"}

        from bs4 import BeautifulSoup
        import re as _re
        soup = BeautifulSoup(resp.text, "lxml")
        # Remove noise
        for tag in soup(["script", "style", "nav", "footer", "header",
                         "aside", "iframe", "form", "noscript"]):
            tag.decompose()
        # Try common article selectors, fall back to body
        body = (
            soup.select_one("article") or
            soup.select_one('[class*="article"]') or
            soup.select_one('[class*="story"]') or
            soup.select_one('[class*="content"]') or
            soup.select_one("main") or
            soup.body
        )
        if not body:
            return {"url": url, "error": "No readable content found"}

        text = body.get_text(separator="\n", strip=True)
        text = _re.sub(r'\n{3,}', '\n\n', text)
        text = _re.sub(r'[ \t]{2,}', ' ', text)

        # Extract title
        title = ""
        title_tag = soup.find("title")
        if title_tag:
            title = title_tag.get_text(strip=True)

        return {
            "url": url,
            "title": title,
            "text": text[:max_chars],
            "truncated": len(text) > max_chars,
            "total_chars": len(text),
        }
    except Exception as e:
        return {"url": url, "error": str(e)}


def search_market_knowledge(query: str, sources: list[str] | None = None) -> dict:
    """Source-backed market education from Wikipedia and Investopedia."""
    import html as _html
    import urllib.parse
    from html.parser import HTMLParser

    import requests

    requested = [s.lower() for s in (sources or ["wikipedia", "investopedia"])]
    clean_query = (query or "").strip()
    if not clean_query:
        return {
            "query": query,
            "source_count": 0,
            "sources": [],
            "answer_markdown": "No market education query was provided.",
        }
    search_query = _market_knowledge_search_query(clean_query)

    headers = {
        "User-Agent": "AgentAdda/1.0 market-education research (Mozilla/5.0)",
        "Accept": "text/html,application/json",
    }

    def _compact_text(text: str, max_chars: int = 900) -> str:
        text = _html.unescape(re.sub(r"\s+", " ", text or "")).strip()
        return text[:max_chars].rstrip()

    def _decode_ddg_url(raw: str) -> str:
        if not raw:
            return ""
        if raw.startswith("//"):
            raw = "https:" + raw
        parsed = urllib.parse.urlparse(raw)
        params = urllib.parse.parse_qs(parsed.query)
        if "uddg" in params:
            return params["uddg"][0]
        return raw

    class _DDGParser(HTMLParser):
        def __init__(self):
            super().__init__()
            self.results: list[dict[str, str]] = []
            self._capture = False
            self._current: dict[str, str] = {}

        def handle_starttag(self, tag, attrs):
            attrs_d = dict(attrs)
            if tag == "a" and "result__a" in attrs_d.get("class", ""):
                self._capture = True
                self._current = {"url": attrs_d.get("href", ""), "title": ""}

        def handle_data(self, data):
            if self._capture and data.strip():
                self._current["title"] = (self._current.get("title", "") + " " + data.strip()).strip()

        def handle_endtag(self, tag):
            if tag == "a" and self._capture:
                if self._current.get("title") and self._current.get("url"):
                    self.results.append(dict(self._current))
                self._current = {}
                self._capture = False

    def _readable_html_text(html_text: str, max_chars: int = 1400) -> str:
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_text, "lxml")
            for tag in soup(["script", "style", "nav", "footer", "header", "aside", "iframe", "form", "noscript"]):
                tag.decompose()
            body = (
                soup.select_one("article")
                or soup.select_one("main")
                or soup.select_one('[class*="article"]')
                or soup.select_one('[class*="content"]')
                or soup.body
            )
            return _compact_text(body.get_text(" ", strip=True) if body else "", max_chars=max_chars)
        except Exception:
            return _compact_text(re.sub(r"<[^>]+>", " ", html_text), max_chars=max_chars)

    source_records: list[dict[str, str]] = []
    errors: list[str] = []

    if "wikipedia" in requested:
        try:
            exact_titles = _market_knowledge_wikipedia_titles(clean_query)
            if exact_titles:
                for exact_title in exact_titles:
                    summary_resp = requests.get(
                        f"https://en.wikipedia.org/api/rest_v1/page/summary/{urllib.parse.quote(exact_title)}",
                        headers=headers,
                        timeout=8,
                    )
                    summary_resp.raise_for_status()
                    summary = summary_resp.json()
                    url = (
                        summary.get("content_urls", {})
                        .get("desktop", {})
                        .get("page", f"https://en.wikipedia.org/wiki/{urllib.parse.quote(exact_title.replace(' ', '_'))}")
                    )
                    extract = _compact_text(summary.get("extract") or "")
                    if extract:
                        source_records.append({
                            "source": "Wikipedia",
                            "title": summary.get("title") or exact_title,
                            "url": url,
                            "snippet": extract,
                        })
            else:
                search_resp = requests.get(
                    "https://en.wikipedia.org/w/api.php",
                    params={
                        "action": "query",
                        "list": "search",
                        "srsearch": search_query,
                        "format": "json",
                        "srlimit": 1,
                    },
                    headers=headers,
                    timeout=8,
                )
                search_resp.raise_for_status()
                hits = search_resp.json().get("query", {}).get("search", [])
                if hits:
                    title = hits[0].get("title", "")
                    summary_resp = requests.get(
                        f"https://en.wikipedia.org/api/rest_v1/page/summary/{urllib.parse.quote(title)}",
                        headers=headers,
                        timeout=8,
                    )
                    summary_resp.raise_for_status()
                    summary = summary_resp.json()
                    url = (
                        summary.get("content_urls", {})
                        .get("desktop", {})
                        .get("page", f"https://en.wikipedia.org/wiki/{urllib.parse.quote(title.replace(' ', '_'))}")
                    )
                    extract = _compact_text(summary.get("extract") or hits[0].get("snippet", ""))
                    if extract:
                        source_records.append({
                            "source": "Wikipedia",
                            "title": summary.get("title") or title,
                            "url": url,
                            "snippet": extract,
                        })
        except Exception as exc:
            errors.append(f"Wikipedia: {exc}")

    if "investopedia" in requested:
        try:
            found_investopedia = False
            ddg_url = "https://html.duckduckgo.com/html/?q=" + urllib.parse.quote(
                f"site:investopedia.com {search_query} finance investing"
            )
            ddg_resp = requests.get(ddg_url, headers=headers, timeout=8)
            ddg_resp.raise_for_status()
            parser = _DDGParser()
            parser.feed(ddg_resp.text)
            for result in parser.results:
                article_url = _decode_ddg_url(result.get("url", ""))
                if "investopedia.com" not in article_url:
                    continue
                article_resp = requests.get(article_url, headers=headers, timeout=8, allow_redirects=True)
                article_resp.raise_for_status()
                snippet = _readable_html_text(article_resp.text)
                if snippet:
                    source_records.append({
                        "source": "Investopedia",
                        "title": _compact_text(result.get("title", "Investopedia article"), max_chars=180),
                        "url": article_url,
                        "snippet": snippet,
                    })
                    found_investopedia = True
                break
            if not found_investopedia:
                direct_url = _market_knowledge_investopedia_url(clean_query)
                if direct_url:
                    article_resp = requests.get(direct_url, headers=headers, timeout=8, allow_redirects=True)
                    if article_resp.status_code == 200:
                        snippet = _readable_html_text(article_resp.text)
                    else:
                        snippet = (
                            f"Investopedia source URL found, but automated text extraction returned "
                            f"HTTP {article_resp.status_code}; no unsourced detail was inferred from it."
                        )
                    if snippet:
                        source_records.append({
                            "source": "Investopedia",
                            "title": "Investopedia concept article",
                            "url": direct_url,
                            "snippet": snippet,
                        })
        except Exception as exc:
            errors.append(f"Investopedia: {exc}")

    concept_note = _market_knowledge_concept_note(clean_query)
    if not source_records:
        return {
            "query": clean_query,
            "sources_requested": requested,
            "source_count": 0,
            "sources": [],
            "errors": errors,
            "answer_markdown": (
                f"No reliable Investopedia or Wikipedia source was found for {clean_query}. "
                "I will not infer the answer without source evidence."
            ),
        }

    answer_lines = [
        f"Source-backed market education: {clean_query}",
        "",
        "What the sources say:",
    ]
    for src in source_records:
        answer_lines.append(f"- {src['source']}: {src['snippet']}")
    if concept_note:
        answer_lines.extend(["", "Concept guide:", concept_note])
    answer_lines.extend(["", "Sources:"])
    for src in source_records:
        answer_lines.append(f"- {src['source']} - {src['title']}: {src['url']}")
    answer_lines.append("")
    answer_lines.append("Use this as market education, not a buy/sell recommendation.")

    return {
        "query": clean_query,
        "sources_requested": requested,
        "source_count": len(source_records),
        "sources": source_records,
        "errors": errors,
        "answer_markdown": "\n".join(answer_lines),
    }


def _market_knowledge_concept_note(query: str) -> str:
    """Small deterministic guide for common finance education concepts."""
    q = query.lower()
    if any(term in q for term in ["p/e", "pe ratio", "p e ratio", "price earnings", "price-to-earnings"]) or re.search(r"\bpe\b", q):
        return (
            "P/E = market price per share divided by earnings per share. "
            "Use it against peers, growth rate, balance-sheet quality, and cycle position; "
            "a low P/E is not automatically cheap and a high P/E is not automatically expensive."
        )
    if "roce" in q and "roe" in q:
        return (
            "ROE focuses on profit generated on shareholders' equity. "
            "ROCE focuses on profit generated on total capital employed, so it is often better "
            "for comparing capital-intensive companies with different debt levels."
        )
    if "roce" in q:
        return (
            "ROCE usually compares operating profit with capital employed. "
            "It helps assess whether the business earns attractive returns on the debt plus equity "
            "capital used in operations."
        )
    if "roe" in q:
        return (
            "ROE compares profit with shareholders' equity. "
            "Check leverage and one-off profits before treating a high ROE as high business quality."
        )
    if "minervini" in q or "vcp" in q or "volatility contraction" in q:
        return (
            "For Minervini-style analysis, separate the concept from a trade call: trend template, "
            "relative strength, earnings/sales growth, constructive base behavior, tight volatility "
            "contraction, and disciplined risk control all need evidence before acting."
        )
    return ""


def _market_knowledge_search_query(query: str) -> str:
    """Normalize ambiguous finance abbreviations before source search."""
    q = query.lower()
    if re.search(r"\bpe\b", q) or "p/e" in q:
        return "price earnings ratio finance investing"
    if "roce" in q and "roe" in q:
        return "return on capital employed return on equity comparison finance"
    if "roce" in q:
        return "return on capital employed finance"
    if "roe" in q:
        return "return on equity finance"
    if "rsi" in q:
        return "relative strength index technical analysis"
    if "minervini" in q:
        return "Mark Minervini trading strategy volatility contraction pattern"
    return query


def _market_knowledge_wikipedia_titles(query: str) -> list[str]:
    """Return exact Wikipedia titles for ambiguous finance abbreviations."""
    q = query.lower()
    if re.search(r"\bpe\b", q) or "p/e" in q or "price earnings" in q:
        return ["Price\u2013earnings ratio"]
    if "roce" in q and "roe" in q:
        return ["Return on capital employed", "Return on equity"]
    if "roce" in q and "roe" not in q:
        return ["Return on capital employed"]
    if "roe" in q and "roce" not in q:
        return ["Return on equity"]
    if "rsi" in q:
        return ["Relative strength index"]
    return []


def _market_knowledge_investopedia_url(query: str) -> str:
    """Return direct Investopedia URLs for stable, common concept pages."""
    q = query.lower()
    if re.search(r"\bpe\b", q) or "p/e" in q or "price earnings" in q:
        return "https://www.investopedia.com/terms/p/price-earningsratio.asp"
    if "roce" in q:
        return "https://www.investopedia.com/terms/r/roce.asp"
    if "roe" in q:
        return "https://www.investopedia.com/terms/r/returnonequity.asp"
    if "ebitda" in q:
        return "https://www.investopedia.com/terms/e/ebitda.asp"
    if "rsi" in q:
        return "https://www.investopedia.com/terms/r/rsi.asp"
    if "macd" in q:
        return "https://www.investopedia.com/terms/m/macd.asp"
    if "cagr" in q:
        return "https://www.investopedia.com/terms/c/cagr.asp"
    return ""


def _vision_transcribe_page(page, page_no: int, dpi: int = 180) -> str:
    """Render a PDF page to PNG and ask the OpenAI vision model to transcribe it.

    Used as a fallback for image-heavy / scanned pages where PyMuPDF text
    extraction yields little or no usable text. Returns the transcribed text
    or an empty string if the call fails or the API key is missing.
    """
    import os, io, base64
    api_key = os.environ.get("OPENAI_API_KEY")
    if not api_key:
        return ""
    try:
        from openai import OpenAI
        zoom = dpi / 72.0
        import fitz  # PyMuPDF
        pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        png_bytes = pix.tobytes("png")
        b64 = base64.b64encode(png_bytes).decode("ascii")
        client = OpenAI(api_key=api_key, timeout=120.0)
        model = os.environ.get("OPENAI_VISION_MODEL") or os.environ.get("OPENAI_MODEL") or "gpt-4o"
        resp = client.chat.completions.create(
            model=model,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": (
                        f"Transcribe page {page_no} VERBATIM. Preserve every number, "
                        f"row label and column header. Render tables as pipe-separated "
                        f"Markdown (| col | col |). Do NOT summarise, do NOT add commentary. "
                        f"If the page contains a Profit & Loss, Balance Sheet, Cash Flow or "
                        f"notes table, reproduce ALL rows and columns exactly as printed."
                    )},
                    {"type": "image_url",
                     "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"}},
                ],
            }],
            max_tokens=4096,
            temperature=0,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception:
        return ""


def fetch_pdf_text(url: str, max_pages: int = 15, vision_fallback: bool = True,
                   vision_threshold: int = 200) -> dict:
    """Download a PDF from any URL and extract its text content using PyMuPDF.

    Handles BSE results PDFs, NSE circulars, concall transcripts, annual reports.
    Returns structured dict with per-page text, page count, and metadata.

    If ``vision_fallback`` is True, any page whose extracted text is shorter
    than ``vision_threshold`` characters is re-extracted via the OpenAI vision
    model (image → text), which captures scanned pages and image-only tables
    that PyMuPDF cannot read.
    """
    import io
    import requests
    original_url = url
    url = _normalise_http_url(url)
    try:
        resolved_url = _resolve_embedded_pdf_url(url) or url
        resp = requests.get(
            resolved_url,
            headers={
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7)",
                "Accept": "application/pdf,*/*",
            },
            timeout=20,
            allow_redirects=True,
        )
        if resp.status_code != 200:
            result = {"url": original_url, "error": f"HTTP {resp.status_code}"}
            if resolved_url != url:
                result["resolved_url"] = resolved_url
            return result

        content_type = resp.headers.get("Content-Type", "")
        if "pdf" not in content_type.lower() and not resolved_url.lower().endswith(".pdf"):
            # Try to proceed anyway — some BSE URLs don't set content-type correctly
            pass

        import fitz  # PyMuPDF
        doc = fitz.open(stream=io.BytesIO(resp.content), filetype="pdf")
        total_pages = len(doc)
        pages_to_read = min(total_pages, max_pages)

        pages_text: list[dict] = []
        full_text_parts: list[str] = []

        for i in range(pages_to_read):
            page = doc[i]
            text = page.get_text("text").strip()
            method = "text"
            if vision_fallback and len(text) < vision_threshold:
                vision_text = _vision_transcribe_page(page, i + 1)
                if vision_text and len(vision_text) > len(text):
                    text = vision_text
                    method = "vision"
            if text:
                pages_text.append({"page": i + 1, "text": text, "extraction_method": method})
                full_text_parts.append(f"--- Page {i + 1} ({method}) ---\n{text}")

        doc.close()

        full_text = "\n\n".join(full_text_parts)
        result = {
            "url":          original_url,
            "source_type":  "pdf",
            "total_pages":  total_pages,
            "pages_read":   pages_to_read,
            "truncated":    total_pages > max_pages,
            "text":         full_text,
            "pages":        pages_text,
        }
        if resolved_url != url:
            result["resolved_url"] = resolved_url
        return result
    except Exception as e:
        result = {"url": original_url, "error": str(e)}
        resolved_url = _resolve_embedded_pdf_url(url) or url
        if resolved_url != url:
            result["resolved_url"] = resolved_url
        return result


def _normalise_http_url(url: str) -> str:
    """Remove paste/wrap whitespace from HTTP URLs."""
    text = (url or "").strip()
    if text.lower().startswith(("http://", "https://")):
        return re.sub(r"\s+", "", text)
    return text


def _resolve_embedded_pdf_url(url: str) -> str | None:
    """Resolve corporate PDF-viewer URLs that carry the real PDF in query params."""
    from urllib.parse import parse_qs, unquote, urljoin, urlparse

    url = _normalise_http_url(url)
    parsed = urlparse(url or "")
    if parsed.scheme not in {"http", "https"}:
        return None

    for key, values in parse_qs(parsed.query, keep_blank_values=True).items():
        if key.lower() not in {"src", "file", "url", "pdf", "path", "document"}:
            continue
        for value in values:
            candidate = unquote(value or "").strip()
            candidate = re.sub(r"\s+", "", candidate)
            if not candidate or ".pdf" not in candidate.lower():
                continue
            return urljoin(url, candidate)
    return None


# ─────────────────────────────────────────────────────────────────────────────
# Portfolio tools
# ─────────────────────────────────────────────────────────────────────────────

HOLDINGS_CSV = ROOT / "portfolio-analyzer" / "output" / "holdings.csv"


def get_portfolio_exposure(sector: str = None) -> dict:
    """Return portfolio holdings summary, optionally filtered by sector.

    Returns total_stocks, sector_counts, top_holdings.  If sector is given,
    also returns filtered symbol list for that sector.
    """
    if not HOLDINGS_CSV.exists():
        return {"error": f"Holdings file not found: {HOLDINGS_CSV}"}
    try:
        df = pd.read_csv(HOLDINGS_CSV)
        df.columns = df.columns.str.lower()
        sym_col = "symbol" if "symbol" in df.columns else df.columns[0]
        val_col = next(
            (c for c in ["value_rs", "value", "current_value"] if c in df.columns), None
        )
        symbols = df[sym_col].str.upper().tolist()

        # Sector mapping from PostgreSQL stage snapshot.
        sector_counts: dict[str, int] = {}
        sym_sector: dict[str, str] = {}
        try:
            placeholders = ",".join(["%s"] * len(symbols))
            rows = _pg_fetchall(
                f"""
                SELECT symbol, sector
                FROM scores.stage_snapshots
                WHERE UPPER(symbol) IN ({placeholders})
                  AND snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)
                """,
                symbols,
            )
            sym_sector = {r[0]: (r[1] or "Unknown") for r in rows}
        except Exception:
            pass
        if not sym_sector and _legacy_sqlite_fallbacks_enabled() and DB_PATH.exists():
            conn = _db_conn()
            placeholders = ",".join("?" * len(symbols))
            rows = conn.execute(
                f"SELECT symbol, sector FROM stage_snapshots "
                f"WHERE symbol IN ({placeholders}) "
                f"GROUP BY symbol ORDER BY snapshot_date DESC",
                symbols,
            ).fetchall()
            conn.close()
            sym_sector = {r[0]: (r[1] or "Unknown") for r in rows}
        for sym in symbols:
            sec = sym_sector.get(sym, "Unknown")
            sector_counts[sec] = sector_counts.get(sec, 0) + 1

        # Top 10 by value
        top_holdings: list[dict] = []
        if val_col and val_col in df.columns:
            df[val_col] = pd.to_numeric(df[val_col], errors="coerce").fillna(0)
            for _, row in df.nlargest(10, val_col).iterrows():
                top_holdings.append({
                    "symbol": row[sym_col],
                    "value":  round(float(row[val_col]), 2),
                    "sector": sym_sector.get(str(row[sym_col]).upper(), "—"),
                })

        result: dict[str, Any] = {
            "total_stocks":  len(df),
            "sector_counts": sector_counts,
            "top_holdings":  top_holdings,
        }
        if sector:
            result["filtered_by_sector"] = sector
            result["sector_symbols"] = [
                sym for sym in symbols
                if sector.lower() in sym_sector.get(sym, "").lower()
            ]
        return result
    except Exception as e:
        return {"error": str(e)}


def find_portfolio_overlap(screener: str = "stage2") -> dict:
    """Find portfolio holdings that also appear in a DB screener result.

    screener: 'stage2' | 'supertrend_buy' | 'all'
    Returns overlap_count, overlapping_symbols with stage/signal data.
    """
    if not HOLDINGS_CSV.exists():
        return {"error": f"Holdings file not found: {HOLDINGS_CSV}"}
    try:
        df = pd.read_csv(HOLDINGS_CSV)
        df.columns = df.columns.str.lower()
        sym_col = "symbol" if "symbol" in df.columns else df.columns[0]
        portfolio_symbols = set(df[sym_col].str.upper().tolist())

        try:
            if screener == "stage2":
                rows = _pg_fetchall(
                    "SELECT symbol, stage, investment_score, trading_signal FROM scores.stage_snapshots "
                    "WHERE stage='STAGE_2' AND snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)"
                )
            elif screener == "supertrend_buy":
                rows = _pg_fetchall(
                    "SELECT symbol, stage, investment_score, trading_signal FROM scores.stage_snapshots "
                    "WHERE supertrend_state IN ('BUY','BULLISH') AND snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)"
                )
            else:
                rows = _pg_fetchall(
                    "SELECT symbol, stage, investment_score, trading_signal FROM scores.stage_snapshots "
                    "WHERE snapshot_date=(SELECT MAX(snapshot_date) FROM scores.stage_snapshots)"
                )
        except Exception:
            if not _legacy_sqlite_fallbacks_enabled() or not DB_PATH.exists():
                return {"error": "PostgreSQL stage snapshots unavailable"}
            conn = _db_conn()
            if screener == "stage2":
                rows = conn.execute(
                    "SELECT symbol, stage, investment_score, trading_signal FROM stage_snapshots "
                    "WHERE stage='STAGE_2' "
                    "AND snapshot_date=(SELECT MAX(snapshot_date) FROM stage_snapshots)"
                ).fetchall()
            elif screener == "supertrend_buy":
                rows = conn.execute(
                    "SELECT symbol, stage, investment_score, trading_signal FROM stage_snapshots "
                    "WHERE supertrend_state='BUY' "
                    "AND snapshot_date=(SELECT MAX(snapshot_date) FROM stage_snapshots)"
                ).fetchall()
            else:
                rows = conn.execute(
                    "SELECT symbol, stage, investment_score, trading_signal FROM stage_snapshots "
                    "WHERE snapshot_date=(SELECT MAX(snapshot_date) FROM stage_snapshots)"
                ).fetchall()
            conn.close()

        screener_map = {r[0]: r for r in rows}
        overlap = [
            {
                "symbol":          sym,
                "stage":           screener_map[sym][1],
                "investment_score": screener_map[sym][2],
                "trading_signal":  screener_map[sym][3],
            }
            for sym in portfolio_symbols
            if sym in screener_map
        ]
        return {
            "screener":          screener,
            "portfolio_count":   len(portfolio_symbols),
            "screener_count":    len(screener_map),
            "overlap_count":     len(overlap),
            "overlapping_symbols": overlap,
        }
    except Exception as e:
        return {"error": str(e)}


# ─────────────────────────────────────────────────────────────────────────────
# Live NSE quote tool
# ─────────────────────────────────────────────────────────────────────────────

_NSE_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    ),
    "Accept":          "application/json, text/plain, */*",
    "Accept-Language": "en-US,en;q=0.9",
    "Referer":         "https://www.nseindia.com/",
}

_live_session: Any = None
_live_session_ts: float = 0.0


def _get_live_session():
    import requests, time as _time
    global _live_session, _live_session_ts
    if _live_session is None or (_time.time() - _live_session_ts) > 180:
        s = requests.Session()
        s.headers.update(_NSE_HEADERS)
        try:
            s.get("https://www.nseindia.com/", timeout=8)
            # Second warmup is required for gated NSE endpoints (cookie/csrf).
            s.get("https://www.nseindia.com/market-data/live-equity-market?symbol=NIFTY+50", timeout=8)
        except Exception:
            pass
        _live_session = s
        _live_session_ts = _time.time()
    return _live_session


def _force_refresh_live_session():
    """Drop the cached NSE session so the next :func:`_get_live_session`
    call performs a fresh warmup. Used as a recovery step when an NSE
    endpoint returns a non-JSON body (cookies expired, rate-limited, or
    the API returned the marketing splash HTML).
    """
    global _live_session, _live_session_ts
    _live_session = None
    _live_session_ts = 0.0


def _nse_get_json(url: str, *, timeout: int = 10) -> dict:
    """Fetch a JSON payload from an NSE endpoint with one cookie-refresh retry.

    Raises ``RuntimeError`` with a human-readable message when the
    response cannot be parsed as JSON on both the initial attempt and
    after a forced session refresh. The legacy callers wrap this in
    ``try/except`` so the error message ends up in the ``ERROR:`` slot
    of the rendered ``SOURCE TRAIL`` instead of an opaque
    ``Expecting value: line 1 column 1 (char 0)``.
    """
    import json as _json

    def _attempt(session) -> dict:
        resp = session.get(url, timeout=timeout)
        status = getattr(resp, "status_code", 0)
        body = resp.text if hasattr(resp, "text") else ""
        body_preview = (body or "").strip()[:160]
        if status >= 400:
            raise RuntimeError(
                f"NSE returned HTTP {status} for {url}; body preview: {body_preview!r}"
            )
        if not body_preview:
            raise RuntimeError(f"NSE returned empty body for {url}")
        # NSE serves a splash/landing HTML page when cookies are stale.
        stripped = body_preview.lstrip().lower()
        if stripped.startswith("<!doctype") or stripped.startswith("<html"):
            raise RuntimeError(
                f"NSE returned HTML (likely stale cookies) for {url}; preview: {body_preview!r}"
            )
        try:
            return resp.json()
        except (_json.JSONDecodeError, ValueError) as exc:
            raise RuntimeError(
                f"NSE returned non-JSON body for {url}: {exc}; preview: {body_preview!r}"
            ) from exc

    try:
        return _attempt(_get_live_session())
    except RuntimeError:
        # Retry once with a freshly-warmed session — covers the common
        # case where the cached session's cookies have just expired.
        _force_refresh_live_session()
        return _attempt(_get_live_session())


def get_live_quote(symbol: str) -> dict:
    """Fetch live intraday quote for a single NSE symbol from the NSE API.

    Returns current price, VWAP, day OHLC, volume, % change, 52-week range,
    circuit limits, sector P/E, and last-update timestamp — all from NSE live.
    """
    import requests
    sym = symbol.strip().upper()
    try:
        s   = _get_live_session()
        url = f"https://www.nseindia.com/api/quote-equity?symbol={requests.utils.quote(sym)}"
        r   = s.get(url, timeout=10)
        r.raise_for_status()
        d   = r.json()

        info  = d.get("info", {})
        meta  = d.get("metadata", {})
        price = d.get("priceInfo", {})
        week  = price.get("weekHighLow", {})
        idhl  = price.get("intraDayHighLow", {})

        last   = price.get("lastPrice",     idhl.get("value"))
        open_  = price.get("open",          None)
        high   = idhl.get("max",            None)
        low    = idhl.get("min",            None)
        prev   = price.get("previousClose", None)
        chg    = price.get("change",        None)
        pchg   = price.get("pChange",       None)
        vwap   = price.get("vwap",          None)

        if last is None:
            return {"symbol": sym, "error": "No price data returned"}

        # Fetch trade info (volume, value, market cap) via separate section call
        vol_lakh = None; val_cr = None; mkt_cap_cr = None
        try:
            r2 = s.get(
                f"https://www.nseindia.com/api/quote-equity?symbol={requests.utils.quote(sym)}&section=trade_info",
                timeout=8,
            )
            ti = r2.json().get("marketDeptOrderBook", {}).get("tradeInfo", {})
            vol_lakh   = ti.get("totalTradedVolume")   # in lakh shares
            val_cr     = ti.get("totalTradedValue")    # in crores
            mkt_cap_cr = ti.get("totalMarketCap")      # in crores
        except Exception:
            pass

        return {
            "symbol":          sym,
            "name":            info.get("companyName", meta.get("symbol", sym)),
            "series":          meta.get("series", "EQ"),
            "last_price":      last,
            "open":            open_,
            "day_high":        high,
            "day_low":         low,
            "vwap":            vwap,
            "prev_close":      prev,
            "change":          round(chg,  2) if chg  is not None else None,
            "pct_change":      round(pchg, 2) if pchg is not None else None,
            "volume":          vol_lakh,        # in lakh shares
            "volume_shares":   round(vol_lakh * 1e5) if vol_lakh else None,
            "traded_value_cr": val_cr,
            "market_cap_cr":   mkt_cap_cr,
            "52w_high":        week.get("max"),
            "52w_low":         week.get("min"),
            "52w_high_date":   week.get("maxDate"),
            "52w_low_date":    week.get("minDate"),
            "lower_circuit":   price.get("lowerCP"),
            "upper_circuit":   price.get("upperCP"),
            "sector":          meta.get("industry"),
            "sector_pe":       meta.get("pdSectorPe"),
            "stock_pe":        meta.get("pdSymbolPe"),
            "indices":         meta.get("pdSectorIndAll", [])[:5],
            "as_of":           meta.get("lastUpdateTime",
                                        datetime.now().strftime("%d-%b-%Y %H:%M:%S")),
            "source":          "NSE live API (real-time)",
        }
    except Exception as e:
        return {"symbol": sym, "error": str(e)}


def get_nse_quotes(symbols: list[str]) -> dict:
    """Fetch live NSE prices for multiple symbols in parallel.

    Returns a dict keyed by symbol with price, % change, VWAP, volume.
    Faster than calling get_live_quote() sequentially.
    Use for: 'prices of RELIANCE, TCS, INFY', 'watchlist prices',
    'how are these stocks doing: X, Y, Z'.

    Args:
        symbols: List of NSE tickers, up to 20.
    """
    import requests
    from concurrent.futures import ThreadPoolExecutor, as_completed

    syms    = [s.strip().upper() for s in symbols[:20] if s.strip()]
    session = _get_live_session()

    def _fetch(sym: str) -> tuple[str, dict]:
        try:
            url = f"https://www.nseindia.com/api/quote-equity?symbol={requests.utils.quote(sym)}"
            r   = session.get(url, timeout=10)
            r.raise_for_status()
            d     = r.json()
            meta  = d.get("metadata", {})
            price = d.get("priceInfo", {})
            idhl  = price.get("intraDayHighLow", {})
            trade = d.get("marketDeptOrderBook", {}).get("tradeInfo", {})
            last  = price.get("lastPrice", idhl.get("value"))
            return sym, {
                "symbol":     sym,
                "name":       d.get("info", {}).get("companyName", sym),
                "last_price": last,
                "change":     round(price.get("change", 0), 2),
                "pct_change": round(price.get("pChange", 0), 2),
                "day_high":   idhl.get("max"),
                "day_low":    idhl.get("min"),
                "vwap":       price.get("vwap"),
                "volume":     trade.get("totalTradedVolume"),
                "prev_close": price.get("previousClose"),
                "lower_circuit": price.get("lowerCP"),
                "upper_circuit": price.get("upperCP"),
                "sector":     meta.get("industry"),
                "sector_pe":  meta.get("pdSectorPe"),
                "stock_pe":   meta.get("pdSymbolPe"),
                "as_of":      meta.get("lastUpdateTime"),
            }
        except Exception as e:
            return sym, {"symbol": sym, "error": str(e)}

    results: dict = {}
    with ThreadPoolExecutor(max_workers=6) as ex:
        futures = {ex.submit(_fetch, s): s for s in syms}
        for fut in as_completed(futures):
            sym, data = fut.result()
            results[sym] = data

    return {
        "quotes":     results,
        "count":      len(results),
        "as_of":      datetime.now().strftime("%d-%b-%Y %H:%M:%S"),
        "source":     "NSE live API (real-time, parallel fetch)",
    }


def nse_search(query: str, top_n: int = 5) -> dict:
    """Search NSE by company name or partial symbol → returns matching stocks with live price.

    Useful for: 'search for Tata companies', 'find steel stocks on NSE',
    'what is the symbol for Adani Ports', 'price of Larsen and Toubro'.

    Args:
        query: Company name, partial name, or keyword (e.g. 'Tata Steel', 'HDFC').
        top_n: Max results to return.
    """
    import requests
    try:
        s   = _get_live_session()
        url = f"https://www.nseindia.com/api/search?q={requests.utils.quote(query)}&type=equity"
        r   = s.get(url, timeout=10)
        r.raise_for_status()
        results = r.json().get("results", [])[:top_n]
        if not results:
            return {"query": query, "results": [], "error": "No matches found on NSE"}

        # Enrich with live prices for top results
        syms   = [x["symbol"] for x in results if x.get("symbol")]
        prices = {}
        if syms:
            batch = get_nse_quotes(syms[:5])
            prices = batch.get("quotes", {})

        enriched = []
        for item in results:
            sym  = item.get("symbol", "")
            info = item.get("symbol_info", "")
            q    = prices.get(sym, {})
            enriched.append({
                "symbol":     sym,
                "name":       info,
                "last_price": q.get("last_price"),
                "pct_change": q.get("pct_change"),
                "sector":     q.get("sector"),
                "series":     item.get("activeSeries", ["EQ"])[0],
            })

        return {
            "query":   query,
            "results": enriched,
            "source":  "NSE search API + live prices",
        }
    except Exception as e:
        return {"query": query, "error": str(e)}


def get_live_market_overview() -> dict:
    """Fetch live NSE index values (broad-market + sectoral) plus advances/declines.

    Returns every NIFTY-prefixed broad-market and sectoral index that the NSE
    /allIndices endpoint exposes, grouped by category. This powers the agent's
    "Latest Sector and Indices View" so the user sees the full picture, not
    just the headline 5.
    """
    import requests
    # PG-indices: full whitelist so we never silently drop a sector.
    # Anything matching these prefixes is included; we sort by category for the UI.
    BROAD_MARKET = {
        "NIFTY 50", "NIFTY NEXT 50", "NIFTY 100", "NIFTY 200", "NIFTY 500",
        "NIFTY MIDCAP 50", "NIFTY MIDCAP 100", "NIFTY MIDCAP 150",
        "NIFTY SMALLCAP 50", "NIFTY SMALLCAP 100", "NIFTY SMALLCAP 250",
        "NIFTY MIDSMALLCAP 400", "NIFTY LARGEMIDCAP 250", "NIFTY TOTAL MARKET",
        "NIFTY MICROCAP 250", "INDIA VIX",
    }
    SECTORAL = {
        "NIFTY BANK", "NIFTY AUTO", "NIFTY FIN SERVICE", "NIFTY FINANCIAL SERVICES",
        "NIFTY FMCG", "NIFTY IT", "NIFTY MEDIA", "NIFTY METAL", "NIFTY PHARMA",
        "NIFTY PSU BANK", "NIFTY PVT BANK", "NIFTY REALTY", "NIFTY HEALTHCARE INDEX",
        "NIFTY CONSUMER DURABLES", "NIFTY OIL & GAS", "NIFTY ENERGY",
        "NIFTY INFRA", "NIFTY COMMODITIES", "NIFTY ENERGY", "NIFTY SERVICES SECTOR",
        "NIFTY MNC", "NIFTY CPSE", "NIFTY PSE", "NIFTY INDIA DEFENCE",
        "NIFTY INDIA MANUFACTURING", "NIFTY INDIA CONSUMPTION", "NIFTY INDIA TOURISM",
        "NIFTY INDIA RAILWAYS PSU", "NIFTY CORE HOUSING", "NIFTY HOUSING",
        "NIFTY MOBILITY", "NIFTY EV & NEW AGE AUTOMOTIVE", "NIFTY TRANSPORTATION & LOGISTICS",
        "NIFTY MIDSMALL HEALTHCARE", "NIFTY MIDSMALL FINANCIAL SERVICES",
        "NIFTY MIDSMALL IT & TELECOM", "NIFTY DIVIDEND OPPORTUNITIES 50",
        "NIFTY GROWTH SECTORS 15", "NIFTY HIGH BETA 50", "NIFTY LOW VOLATILITY 50",
        "NIFTY ALPHA 50", "NIFTY QUALITY 30", "NIFTY VALUE 20",
    }
    try:
        s   = _get_live_session()
        url = "https://www.nseindia.com/api/allIndices"
        r   = s.get(url, timeout=10)
        r.raise_for_status()
        data = r.json().get("data", [])
        broad: dict = {}
        sector: dict = {}
        other: dict = {}
        for item in data:
            nm = (item.get("index") or "").strip().upper()
            if not nm:
                continue
            last  = item.get("last",          item.get("lastPrice", 0))
            prev  = item.get("previousClose", 0)
            chg   = item.get("variation")
            if chg is None:
                chg = round(last - prev, 2) if last and prev else 0
            pchg  = item.get("percentChange")
            if pchg is None:
                pchg = round(chg / prev * 100, 2) if prev else 0
            entry = {
                "last":       last,
                "change":     chg,
                "pct_change": pchg,
                "day_high":   item.get("dayHigh", item.get("high")),
                "day_low":    item.get("dayLow", item.get("low")),
            }
            if nm in BROAD_MARKET:
                broad[nm] = entry
            elif nm in SECTORAL or nm.startswith("NIFTY "):
                # Catch-all: any NIFTY * index we didn't explicitly enumerate
                # still goes into the sectoral bucket so nothing is dropped.
                sector[nm] = entry
            else:
                other[nm] = entry

        # PG-indices: keep `indices` flat for backward compat with existing callers.
        indices = {**broad, **sector}

        # Top movers among equity sector/thematic indices for quick
        # "winners/losers" framing. NSE allIndices also includes bond/debt
        # indices under NIFTY names; exclude them from sector-strength answers.
        non_equity_markers = ("BOND", "G-SEC", "SDL", "T-BILL", "LIQUID", "OVERNIGHT")
        equity_sector = {
            name: row
            for name, row in sector.items()
            if name in SECTORAL and not any(marker in name for marker in non_equity_markers)
        }
        sector_sorted = sorted(equity_sector.items(), key=lambda kv: kv[1]["pct_change"], reverse=True)
        top_sectors = [{"name": k, **v} for k, v in sector_sorted[:5]]
        bot_sectors = [{"name": k, **v} for k, v in sector_sorted[-5:][::-1]]

        adv_dec = {}
        try:
            # Migrated from deprecated `equity-stockIndices?index=NIFTY%20500`
            # to `live-analysis-variations` (the documented replacement).
            # The variations payload exposes advance/decline counts both at
            # the bucket level and via per-row `pChange` — derive defensively.
            url2 = "https://www.nseindia.com/api/live-analysis-variations?index=gainers"
            r2 = s.get(url2, timeout=10)
            payload2 = r2.json() or {}
            bucket = payload2.get("allSec") or payload2.get("NIFTY") or {}
            items = bucket.get("data") or []
            adv = bucket.get("advances")
            dec = bucket.get("declines")
            unc = bucket.get("unchanged")
            if adv is None or dec is None:
                # Fall back to row-level tally when the bucket summary is absent.
                adv = sum(1 for x in items if float(x.get("perChange", x.get("pChange", 0)) or 0) > 0)
                dec = sum(1 for x in items if float(x.get("perChange", x.get("pChange", 0)) or 0) < 0)
                unc = len(items) - adv - dec
            adv_dec = {"advances": int(adv), "declines": int(dec),
                       "unchanged": int(unc) if unc is not None else 0}
        except Exception:
            pass
        return {
            "indices":           indices,        # flat dict (back-compat)
            "broad_market":      broad,
            "sectoral":          sector,
            "other_indices":     other,
            "top_sectors":       top_sectors,
            "bottom_sectors":    bot_sectors,
            "adv_dec":           adv_dec,
            "as_of":             datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "source":            "NSE live API",
        }
    except Exception as e:
        return {"error": str(e)}


def _intraday_market_overview_from_pg(preferred: list[str], live_error: str | None = None) -> dict:
    """Build a market-overview-shaped payload from the latest stored index tape."""
    try:
        import psycopg2
        from psycopg2.extras import RealDictCursor

        conn = psycopg2.connect("dbname=nse_market user=nse_admin host=/tmp")
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                cur.execute(
                    """
                    WITH ranked AS (
                        SELECT symbol, as_of, captured_at, last_price, change, pct_change,
                               day_high, day_low,
                               row_number() OVER (PARTITION BY symbol ORDER BY captured_at DESC) AS rn
                        FROM   intraday.quote_snapshots
                        WHERE  symbol = ANY(%s)
                          AND  last_price IS NOT NULL
                    )
                    SELECT symbol, as_of, captured_at, last_price, change, pct_change, day_high, day_low
                    FROM   ranked
                    WHERE  rn = 1
                    """,
                    (preferred,),
                )
                rows = [dict(row) for row in cur.fetchall()]
        finally:
            conn.close()
    except Exception as exc:
        return {"error": f"PG intraday quote fallback unavailable: {exc}"}

    if not rows:
        return {"error": "PG intraday quote fallback has no stored index tape"}

    indices = {}
    latest_captured = None
    for row in rows:
        symbol = str(row.get("symbol") or "").upper()
        indices[symbol] = {
            "last": float(row["last_price"]) if row.get("last_price") is not None else None,
            "change": float(row["change"]) if row.get("change") is not None else None,
            "pct_change": float(row["pct_change"]) if row.get("pct_change") is not None else None,
            "day_high": float(row["day_high"]) if row.get("day_high") is not None else None,
            "day_low": float(row["day_low"]) if row.get("day_low") is not None else None,
        }
        captured_at = row.get("captured_at")
        if captured_at and (latest_captured is None or captured_at > latest_captured):
            latest_captured = captured_at

    return {
        "indices": indices,
        "broad_market": {k: v for k, v in indices.items() if k in {"NIFTY 50", "NIFTY NEXT 50"}},
        "sectoral": {k: v for k, v in indices.items() if k not in {"NIFTY 50", "NIFTY NEXT 50"}},
        "adv_dec": {},
        "as_of": str(latest_captured) if latest_captured else datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "source": "PG intraday.quote_snapshots fallback",
        "degraded": True,
        "fallback_reason": f"NSE live overview unavailable: {live_error}" if live_error else "NSE live overview unavailable.",
    }


def get_intraday_market_recap(minutes: int = 15) -> dict:
    """Summarize what changed in the broad market over the last N minutes."""
    minutes = max(1, min(int(minutes or 15), 120))
    preferred = [
        "NIFTY 50",
        "NIFTY BANK",
        "NIFTY MIDCAP SELECT",
        "NIFTY MIDCAP 50",
        "NIFTY MIDCAP 100",
    ]
    overview = get_live_market_overview()
    if overview.get("error"):
        live_error = str(overview.get("error"))
        overview = _intraday_market_overview_from_pg(preferred, live_error=live_error)
        if overview.get("error"):
            return {
                "error": f"Live market overview unavailable: {live_error}; {overview['error']}",
                "minutes": minutes,
            }

    indices = overview.get("indices") or {}
    current_rows: list[dict[str, Any]] = []
    for name in preferred:
        row = indices.get(name)
        if not row:
            continue
        current_rows.append({
            "symbol": name,
            "source": overview.get("source", "NSE live API"),
            "as_of": overview.get("as_of"),
            "name": name,
            "last_price": row.get("last"),
            "change": row.get("change"),
            "pct_change": row.get("pct_change"),
            "day_high": row.get("day_high"),
            "day_low": row.get("day_low"),
        })

    # Persist the current live tape so later "last N minutes" requests have an anchor.
    for row in current_rows:
        try:
            persist_intraday_snapshot(row)
        except Exception:
            pass

    prior_map: dict[str, dict] = {}
    try:
        import psycopg2
        from psycopg2.extras import RealDictCursor

        conn = psycopg2.connect("dbname=nse_market user=nse_admin host=/tmp")
        try:
            with conn.cursor(cursor_factory=RealDictCursor) as cur:
                # PG-recap-fix: previously this query required snapshots inside a
                # narrow [minutes, minutes+10] slice, which was almost always empty
                # because intraday.quote_snapshots is populated sporadically (only
                # when /briefing or recap is invoked). The new logic takes the
                # most recent snapshot at or before now() - minutes; if none exists
                # we fall back to the second-most-recent snapshot per symbol so the
                # user still sees a "since last refresh" delta instead of an empty
                # narrative.
                symbols = [row["symbol"] for row in current_rows]
                cur.execute(
                    """
                    WITH bound AS (
                        SELECT  symbol, MAX(captured_at) AS captured_at
                        FROM    intraday.quote_snapshots
                        WHERE   symbol = ANY(%s)
                          AND   captured_at <= now() - (%s::text || ' minutes')::interval
                        GROUP   BY symbol
                    )
                    SELECT  q.symbol, q.as_of, q.captured_at, q.last_price, q.pct_change
                    FROM    intraday.quote_snapshots q
                    JOIN    bound b
                           ON b.symbol = q.symbol AND b.captured_at = q.captured_at
                    """,
                    (symbols, str(minutes)),
                )
                prior_map = {str(row["symbol"]): dict(row) for row in cur.fetchall()}

                # Fallback: any symbol still missing → use the snapshot just
                # before the most recent one (i.e. the previous tape).
                missing = [s for s in symbols if s not in prior_map]
                if missing:
                    cur.execute(
                        """
                        WITH ranked AS (
                            SELECT symbol, as_of, captured_at, last_price, pct_change,
                                   row_number() OVER (PARTITION BY symbol
                                                      ORDER BY captured_at DESC) AS rn
                            FROM   intraday.quote_snapshots
                            WHERE  symbol = ANY(%s)
                        )
                        SELECT symbol, as_of, captured_at, last_price, pct_change
                        FROM   ranked WHERE rn = 2
                        """,
                        (missing,),
                    )
                    for row in cur.fetchall():
                        prior_map[str(row["symbol"])] = dict(row)
        finally:
            conn.close()
    except Exception:
        prior_map = {}

    rows = []
    for row in current_rows:
        prior = prior_map.get(row["symbol"])
        last = row.get("last_price")
        prior_last = prior.get("last_price") if prior else None
        try:
            point_change = float(last) - float(prior_last) if last is not None and prior_last is not None else None
            interval_pct = (point_change / float(prior_last) * 100) if prior_last else None
        except (TypeError, ValueError, ZeroDivisionError):
            point_change = None
            interval_pct = None
        rows.append({
            "symbol": row["symbol"],
            "current": last,
            "current_pct_change": row.get("pct_change"),
            "prior": prior_last,
            "point_change": round(point_change, 2) if point_change is not None else None,
            "interval_pct_change": round(interval_pct, 3) if interval_pct is not None else None,
            "prior_as_of": str(prior.get("as_of")) if prior else None,
            "prior_captured_at": str(prior.get("captured_at")) if prior else None,
        })

    adv_dec = overview.get("adv_dec") or {}
    changes = [r["interval_pct_change"] for r in rows if r.get("interval_pct_change") is not None]
    avg_change = round(sum(changes) / len(changes), 3) if changes else None
    if avg_change is None:
        tone = "snapshot_only"
        narrative = "No earlier stored tape was available for a clean interval comparison; showing current live market state."
    elif avg_change > 0.05:
        tone = "improved"
        narrative = f"Market improved over the last {minutes} minutes, led by positive index drift."
    elif avg_change < -0.05:
        tone = "weakened"
        narrative = f"Market weakened over the last {minutes} minutes, with headline indices drifting lower."
    else:
        tone = "flat"
        narrative = f"Market was largely range-bound over the last {minutes} minutes."

    return {
        "minutes": minutes,
        "as_of": overview.get("as_of"),
        "source": overview.get("source", "NSE live API"),
        "tone": tone,
        "narrative": narrative,
        "adv_dec": adv_dec,
        "rows": rows,
        "note": "Interval deltas use the nearest stored intraday.quote_snapshots row at or before the requested lookback.",
        "degraded": bool(overview.get("degraded")),
        "fallback_reason": overview.get("fallback_reason"),
    }


_NSE_LIVE_INDEX_ALIASES: dict[str, str] = {
    "NIFTY": "NIFTY 50",
    "NIFTY50": "NIFTY 50",
    "NIFTY 50": "NIFTY 50",
    "BANKNIFTY": "NIFTY BANK",
    "BANK NIFTY": "NIFTY BANK",
    "NIFTY BANK": "NIFTY BANK",
    "FINNIFTY": "NIFTY FIN SERVICE",
    "NIFTY FIN": "NIFTY FIN SERVICE",
    "NIFTY FINANCIAL": "NIFTY FIN SERVICE",
    "NIFTY IT": "NIFTY IT",
    "NIFTY NEXT 50": "NIFTY NEXT 50",
    "NIFTYNXT50": "NIFTY NEXT 50",
    "MIDCPNIFTY": "NIFTY MIDCAP SELECT",
    "MIDC NIFTY": "NIFTY MIDCAP SELECT",
    "MIDCAP NIFTY": "NIFTY MIDCAP SELECT",
    "NIFTY MIDCAP SELECT": "NIFTY MIDCAP SELECT",
    "NIFTY MID SELECT": "NIFTY MIDCAP SELECT",
    "INDIA VIX": "INDIA VIX",
    "VIX": "INDIA VIX",
    "SENSEX": "SENSEX",
}


def _canonical_nse_live_index(symbol: str) -> str | None:
    raw = symbol.strip().upper().replace("_", " ")
    raw = re.sub(r"\s+", " ", raw)
    return _NSE_LIVE_INDEX_ALIASES.get(raw)


def _with_intraday_pg_persistence(snapshot: dict) -> dict:
    """Attach best-effort PostgreSQL persistence metadata to a live snapshot."""
    if not snapshot or snapshot.get("error"):
        return snapshot
    out = dict(snapshot)
    try:
        persisted = persist_intraday_snapshot(out)
        out["postgres_persist"] = {
            "ok": bool(persisted.get("ok")),
            "schema": persisted.get("schema", "intraday"),
            "table": persisted.get("table", "quote_snapshots"),
            "rows_inserted": persisted.get("rows_inserted", 0),
        }
    except Exception as exc:
        out["postgres_persist"] = {
            "ok": False,
            "schema": "intraday",
            "table": "quote_snapshots",
            "error": str(exc),
        }
    return out


def _nse_quote_payload_to_snapshot(symbol: str, data: dict, *, source: str) -> dict:
    """Normalize NSE quote-equity JSON into the agent's live snapshot shape."""
    sym = symbol.strip().upper()
    info = data.get("info", {}) if isinstance(data, dict) else {}
    meta = data.get("metadata", {}) if isinstance(data, dict) else {}
    price = data.get("priceInfo", {}) if isinstance(data, dict) else {}
    week = price.get("weekHighLow", {}) if isinstance(price, dict) else {}
    idhl = price.get("intraDayHighLow", {}) if isinstance(price, dict) else {}
    last = price.get("lastPrice", idhl.get("value"))
    if last is None:
        return {"symbol": sym, "error": "No price data returned from NSE browser quote page"}
    return {
        "symbol": sym,
        "name": info.get("companyName", meta.get("symbol", sym)),
        "series": meta.get("series", "EQ"),
        "last_price": last,
        "open": price.get("open"),
        "day_high": idhl.get("max"),
        "day_low": idhl.get("min"),
        "vwap": price.get("vwap"),
        "prev_close": price.get("previousClose"),
        "change": round(price.get("change"), 2) if price.get("change") is not None else None,
        "pct_change": round(price.get("pChange"), 2) if price.get("pChange") is not None else None,
        "52w_high": week.get("max"),
        "52w_low": week.get("min"),
        "52w_high_date": week.get("maxDate"),
        "52w_low_date": week.get("minDate"),
        "lower_circuit": price.get("lowerCP"),
        "upper_circuit": price.get("upperCP"),
        "sector": meta.get("industry"),
        "sector_pe": meta.get("pdSectorPe"),
        "stock_pe": meta.get("pdSymbolPe"),
        "indices": meta.get("pdSectorIndAll", [])[:5] if isinstance(meta.get("pdSectorIndAll"), list) else [],
        "as_of": meta.get("lastUpdateTime", datetime.now().strftime("%d-%b-%Y %H:%M:%S")),
        "source": source,
    }


def _playwright_nse_quote_snapshot(symbol: str) -> dict:
    """Try NSE's rendered quote page in a real browser and capture quote JSON."""
    sym = symbol.strip().upper()
    try:
        from playwright.sync_api import sync_playwright
    except Exception as exc:
        return {"symbol": sym, "error": f"Playwright unavailable: {exc}"}

    try:
        from urllib.parse import quote

        quote_url = f"https://www.nseindia.com/get-quotes/equity?symbol={quote(sym)}"
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(
                headless=True,
                args=["--disable-blink-features=AutomationControlled", "--no-sandbox"],
            )
            try:
                context = browser.new_context(
                    viewport={"width": 1440, "height": 900},
                    locale="en-US",
                    timezone_id="Asia/Kolkata",
                    user_agent=(
                        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) "
                        "Chrome/124.0.0.0 Safari/537.36"
                    ),
                    extra_http_headers={"Accept-Language": "en-US,en;q=0.9"},
                )
                try:
                    page = context.new_page()
                    payloads: list[dict] = []

                    def _capture_quote_response(response) -> None:
                        if "quote-equity" not in response.url or response.status >= 400:
                            return
                        try:
                            data = response.json()
                        except Exception:
                            return
                        if isinstance(data, dict):
                            payloads.append(data)

                    page.on("response", _capture_quote_response)
                    page.goto(quote_url, wait_until="domcontentloaded", timeout=12000)
                    try:
                        page.wait_for_response(
                            lambda response: "quote-equity" in response.url and response.status < 400,
                            timeout=8000,
                        )
                    except Exception:
                        pass
                    page.wait_for_timeout(1500)
                    if payloads:
                        return _nse_quote_payload_to_snapshot(sym, payloads[-1], source="NSE browser quote page")
                    return {"symbol": sym, "error": "NSE browser quote page did not expose quote-equity JSON"}
                finally:
                    context.close()
            finally:
                browser.close()
    except Exception as exc:
        return {"symbol": sym, "error": f"NSE browser quote page unavailable: {exc}"}


def _yfinance_snapshot_from_intraday_candles(symbol: str, fallback_reason: str) -> dict:
    """Build a degraded live-snapshot shape from yfinance intraday candles."""
    sym = symbol.strip().upper()
    try:
        df = get_intraday_candles(sym, "15m")
    except Exception as exc:
        return {
            "symbol": sym,
            "error": f"{fallback_reason}; yfinance fallback unavailable: {exc}",
            "source_priority": ["NSE website live quote", "yfinance candles fallback"],
        }
    if df.empty:
        return {
            "symbol": sym,
            "error": f"{fallback_reason}; yfinance fallback returned no candles",
            "source_priority": ["NSE website live quote", "yfinance candles fallback"],
        }

    df = df.copy().sort_index()
    latest = df.iloc[-1]
    first = df.iloc[0]
    last_price = float(latest["Close"])
    open_price = float(first["Open"])
    change = last_price - open_price
    try:
        as_of = pd.to_datetime(df.index[-1])
        if getattr(as_of, "tzinfo", None) is not None:
            as_of = as_of.tz_convert(IST_TZ).tz_localize(None)
        as_of_text = as_of.strftime("%Y-%m-%d %H:%M:%S")
    except Exception:
        as_of_text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    return {
        "symbol": sym,
        "name": sym,
        "last_price": round(last_price, 2),
        "open": round(open_price, 2),
        "day_high": round(float(df["High"].max()), 2),
        "day_low": round(float(df["Low"].min()), 2),
        "prev_close": None,
        "change": round(change, 2),
        "pct_change": round(change / open_price * 100, 2) if open_price else None,
        "volume": int(df["Volume"].sum()) if "Volume" in df.columns else None,
        "volume_shares": int(df["Volume"].sum()) if "Volume" in df.columns else None,
        "as_of": as_of_text,
        "source": "Yahoo Finance (yfinance) fallback",
        "source_priority": ["NSE website live quote", "yfinance candles fallback"],
        "degraded": True,
        "fallback_reason": fallback_reason,
        "note": "NSE quote-equity was unavailable; snapshot uses latest yfinance intraday candle.",
    }


def get_nse_intraday_snapshot(symbol: str) -> dict:
    """Fetch the NSE website live snapshot before any yfinance intraday fallback.

    NSE's public website APIs provide current quote/index snapshots, not a complete
    intraday OHLCV candle history. Agent Adda therefore tries NSE first, then uses
    yfinance when NSE blocks stock quotes or when candle history is needed.
    """
    sym = symbol.strip().upper()
    index_name = _canonical_nse_live_index(sym)
    source_priority = ["NSE website live quote", "yfinance candles fallback"]

    if index_name:
        overview = get_live_market_overview()
        if overview.get("error"):
            return {
                "symbol": index_name,
                "source_priority": source_priority,
                "error": f"NSE index snapshot unavailable: {overview['error']}",
            }
        index_row = (overview.get("indices") or {}).get(index_name)
        if not index_row:
            return {
                "symbol": index_name,
                "source": overview.get("source", "NSE live API"),
                "source_priority": source_priority,
                "error": f"NSE index snapshot not found for {index_name}",
            }
        return _with_intraday_pg_persistence({
            "symbol": index_name,
            "name": index_name,
            "last_price": index_row.get("last"),
            "change": index_row.get("change"),
            "pct_change": index_row.get("pct_change"),
            "day_high": index_row.get("day_high"),
            "day_low": index_row.get("day_low"),
            "as_of": overview.get("as_of"),
            "source": overview.get("source", "NSE live API"),
            "source_priority": source_priority,
            "note": "NSE website snapshot. yfinance is used only if NSE quote is unavailable or candle history is needed.",
        })

    quote = get_live_quote(sym)
    if quote.get("error"):
        browser_quote = _playwright_nse_quote_snapshot(sym)
        if not browser_quote.get("error"):
            browser_quote = dict(browser_quote)
            browser_quote["source_priority"] = source_priority
            browser_quote["note"] = "NSE quote fetched through browser-rendered quote page."
            return _with_intraday_pg_persistence(browser_quote)
        fallback_reason = f"NSE live quote unavailable: {quote['error']}; {browser_quote.get('error')}"
        fallback = _yfinance_snapshot_from_intraday_candles(
            sym,
            fallback_reason,
        )
        return _with_intraday_pg_persistence(fallback)
    quote = dict(quote)
    quote["source_priority"] = source_priority
    quote["note"] = "NSE website snapshot. yfinance is used only if NSE quote is unavailable or candle history is needed."
    return _with_intraday_pg_persistence(quote)


_VARIATIONS_BUCKET_FOR_INDEX: dict[str, str] = {
    "NIFTY":              "NIFTY",
    "NIFTY 50":           "NIFTY",
    "NIFTY50":            "NIFTY",
    "BANKNIFTY":          "BANKNIFTY",
    "BANK NIFTY":         "BANKNIFTY",
    "NIFTY BANK":         "BANKNIFTY",
    "NIFTYNEXT50":        "NIFTYNEXT50",
    "NIFTY NEXT 50":      "NIFTYNEXT50",
    "FNO":                "FOSec",
    "F&O":                "FOSec",
    "FOSEC":              "FOSec",
    "NIFTY F&O":          "FOSec",
    "SHOCKERS":           "SecGtr20",
    "GTR20":              "SecGtr20",
}


def _variations_bucket_key(index: str) -> str:
    """Map a user-facing index name to a `live-analysis-variations` bucket.

    NSE replaced the per-index `equity-stockIndices` constituent endpoint
    with a fixed set of buckets in `live-analysis-variations`. Anything we
    don't recognise (NIFTY 500, NIFTY MIDCAP 100, etc.) falls back to the
    `allSec` bucket which covers the broader market.
    """
    key = (index or "").strip().upper()
    key = re.sub(r"\s+", " ", key)
    return _VARIATIONS_BUCKET_FOR_INDEX.get(key, "allSec")


def _fmt_variation_row(x: dict) -> dict:
    return {
        "symbol":     x.get("symbol"),
        "last_price": x.get("ltp"),
        "change":     round(float(x.get("net_price", 0) or 0), 2),
        "pct_change": round(float(x.get("perChange", 0) or 0), 2),
        "volume":     x.get("trade_quantity"),
        "turnover":   x.get("turnover"),
        "day_high":   x.get("high_price"),
        "day_low":    x.get("low_price"),
        "open_price": x.get("open_price"),
        "prev_price": x.get("prev_price"),
    }


def get_top_gainers_losers(
    index: str = "NIFTY 500",
    top_n: int = 10,
    direction: str = "both",
) -> dict:
    """Return top gaining and/or losing stocks from NSE right now.

    Backed by NSE's `live-analysis-variations` API (the same data that
    powers nseindia.com's "Top Gainers / Losers" page). The legacy
    `equity-stockIndices?index=...` endpoint that this used to call was
    deprecated by NSE and now returns 404 for every index.

    Args:
        index: Universe to scan. Recognised: 'NIFTY 50', 'NIFTY BANK',
               'NIFTY NEXT 50', 'F&O' (F&O securities), 'SHOCKERS'
               (>20% movers). Anything else (incl. 'NIFTY 500',
               'NIFTY MIDCAP 100') falls back to all securities.
        top_n: Number of stocks to return in each list (default 10).
        direction: 'gainers', 'losers', or 'both' (default 'both').
    """
    try:
        bucket = _variations_bucket_key(index)
        result: dict = {
            "index":  index,
            "bucket": bucket,
            "as_of":  datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "source": "NSE live-analysis-variations",
        }

        if direction in ("gainers", "both"):
            payload = _nse_get_json(
                "https://www.nseindia.com/api/live-analysis-variations?index=gainers",
                timeout=10,
            )
            rows = ((payload.get(bucket) or {}).get("data") or [])
            rows_sorted = sorted(
                rows,
                key=lambda x: float(x.get("perChange", 0) or 0),
                reverse=True,
            )[:top_n]
            result["gainers"] = [_fmt_variation_row(x) for x in rows_sorted]

        if direction in ("losers", "both"):
            # NSE's endpoint spells the param `loosers` (sic) — using
            # the correct spelling returns "Missing index or key.".
            payload = _nse_get_json(
                "https://www.nseindia.com/api/live-analysis-variations?index=loosers",
                timeout=10,
            )
            rows = ((payload.get(bucket) or {}).get("data") or [])
            rows_sorted = sorted(
                rows,
                key=lambda x: float(x.get("perChange", 0) or 0),
            )[:top_n]
            result["losers"] = [_fmt_variation_row(x) for x in rows_sorted]

        return result
    except Exception as e:
        return {"error": str(e), "index": index}


def get_eod_top_movers(
    index: str = "NIFTY 500",
    top_n: int = 10,
    direction: str = "both",
) -> dict:
    """Return top EOD gainers and/or losers from the latest snapshot.

    Sourced from `scores.stage_snapshots` (the same daily EOD snapshot
    that powers the rest of the historical-mode toolkit), ordered by
    `change_1d_pct`. Use this when the agent is in historical/EOD mode
    or when the user explicitly asks for end-of-day top movers.

    Args:
        index: Universe filter. Recognised tokens map to known indices
               in the snapshot ('NIFTY 50', 'NIFTY BANK', 'NIFTY NEXT 50',
               'NIFTY 500'). Anything else returns the full snapshot
               universe.
        top_n: Number of stocks per list (default 10).
        direction: 'gainers', 'losers', or 'both' (default 'both').
    """
    try:
        snap_date = _latest_snapshot_date()
        if not snap_date:
            return {"error": "no EOD snapshot available", "index": index}

        cols = (
            "symbol, company_name, price, change_1d_pct, "
            "stage, investment_score, relative_strength, rsi, "
            "trading_signal, sector"
        )

        result: dict = {
            "index": index,
            "as_of": snap_date,
            "source": "PostgreSQL scores.stage_snapshots",
        }

        def _rows(order: str) -> list[dict]:
            rows = _pg_fetchall(
                f"SELECT {cols} FROM scores.stage_snapshots "
                f"WHERE snapshot_date=%s AND change_1d_pct IS NOT NULL "
                f"ORDER BY change_1d_pct {order} NULLS LAST LIMIT %s",
                (snap_date, max(int(top_n), 1)),
            )
            keys = (
                "symbol", "company_name", "price", "change_1d_pct",
                "stage", "investment_score", "relative_strength", "rsi",
                "trading_signal", "sector",
            )
            out: list[dict] = []
            for r in rows:
                d = dict(zip(keys, r))
                # Map to the live-API shape so renderer.render_gainers_losers
                # can be reused without a new renderer.
                d["last_price"] = d.get("price")
                d["pct_change"] = d.get("change_1d_pct")
                d["change"] = None
                d["volume"] = None
                d["year_high"] = None
                d["year_low"] = None
                out.append(d)
            return out

        if direction in ("gainers", "both"):
            result["gainers"] = _rows("DESC")
        if direction in ("losers", "both"):
            result["losers"] = _rows("ASC")
        return result
    except Exception as e:
        return {"error": str(e), "index": index}


def get_most_active_stocks(
    by: str = "value",
    index: str = "NIFTY 500",
    top_n: int = 10,
) -> dict:
    """Return most actively traded stocks by volume or traded value.

    Backed by NSE's `live-analysis-most-active-securities` endpoint
    (replaces the deprecated `equity-stockIndices` aggregation). The
    `index` parameter is preserved for backwards compatibility but is
    no longer honoured — NSE's endpoint returns a single market-wide
    list — and is echoed back in the result for trace fidelity.

    Args:
        by: 'volume' or 'value' (default 'value').
        index: Echoed back in the response (no longer scoped by NSE).
        top_n: Number of results (default 10).
    """
    try:
        sort_param = "volume" if by == "volume" else "value"
        payload = _nse_get_json(
            "https://www.nseindia.com/api/live-analysis-most-active-securities"
            f"?index={sort_param}&limit={max(int(top_n), 1)}",
            timeout=10,
        )
        rows = payload.get("data", []) or []
        return {
            "by":     by,
            "index":  index,
            "as_of":  datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "source": "NSE live-analysis-most-active-securities",
            "stocks": [
                {
                    "symbol":       x.get("symbol"),
                    "last_price":   x.get("lastPrice"),
                    "pct_change":   round(float(x.get("pChange", 0) or 0), 2),
                    "volume":       x.get("quantityTraded") or x.get("totalTradedVolume"),
                    "traded_value": x.get("turnoverInLakhs") or x.get("totalTradedValue"),
                }
                for x in rows[:top_n]
            ],
        }
    except Exception as e:
        return {"error": str(e)}


def get_52week_extremes(
    direction: str = "high",
    index: str = "NIFTY 500",
    top_n: int = 15,
) -> dict:
    """Return stocks at their 52-week high or low.

    Backed by NSE's `live-analysis-data-52weekhighstock` /
    `live-analysis-data-52weeklowstock` endpoints (replace the
    deprecated `equity-stockIndices` scan). The `index` parameter is
    preserved for backwards compatibility but is no longer used — NSE's
    endpoint returns a single market-wide list — and is echoed back in
    the result for trace fidelity.

    Args:
        direction: 'high' (at 52w high) or 'low' (at 52w low).
        index:     Echoed back; not used to filter (NSE returns all).
        top_n:     Number of stocks to return (default 15).
    """
    try:
        url = (
            "https://www.nseindia.com/api/live-analysis-data-52weekhighstock"
            if direction == "high"
            else "https://www.nseindia.com/api/live-analysis-data-52weeklowstock"
        )
        payload = _nse_get_json(url, timeout=10)
        rows = payload.get("data", []) or []

        results = []
        for x in rows[:top_n]:
            ltp = float(x.get("ltp") or 0)
            new_extreme = float(x.get("new52WHL") or 0)
            prev_extreme = float(x.get("prev52WHL") or 0)
            row = {
                "symbol":         x.get("symbol"),
                "company":        x.get("comapnyName") or x.get("companyName"),
                "last_price":     ltp,
                "prev_52w_extreme":  prev_extreme or None,
                "prev_extreme_date": x.get("prevHLDate"),
                "pct_change_day": round(float(x.get("pChange") or 0), 2),
                "change_day":     round(float(x.get("change") or 0), 2),
            }
            if direction == "high":
                row["52w_high"]      = new_extreme or None
                row["year_high"]     = new_extreme or None  # backward-compat for renderer
                row["pct_from_high"] = round((ltp - new_extreme) / new_extreme * 100, 2) \
                                       if new_extreme > 0 else None
            else:
                row["52w_low"]      = new_extreme or None
                row["year_low"]     = new_extreme or None   # backward-compat for renderer
                row["pct_from_low"] = round((ltp - new_extreme) / new_extreme * 100, 2) \
                                      if new_extreme > 0 else None
            results.append(row)

        return {
            "direction": direction,
            "index":     index,
            "as_of":     datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "source":    "NSE live-analysis-data-52week",
            "stocks":    results,
        }
    except Exception as e:
        return {"error": str(e)}


def get_fii_dii_activity() -> dict:
    """Fetch today's FII (Foreign Institutional Investors) and DII (Domestic)
    buy/sell activity from NSE. Returns net values in crores.
    """
    try:
        s = _get_live_session()
        r = s.get("https://www.nseindia.com/api/fiidiiTradeReact", timeout=10)
        raw = r.json()
        result: dict = {"as_of": datetime.now().strftime("%Y-%m-%d"), "data": []}
        for entry in raw:
            buy  = float(entry.get("buyValue",  0) or 0)
            sell = float(entry.get("sellValue", 0) or 0)
            net  = float(entry.get("netValue",  buy - sell) or 0)
            result["data"].append({
                "category":   entry.get("category"),
                "date":       entry.get("date"),
                "buy_crore":  round(buy,  2),
                "sell_crore": round(sell, 2),
                "net_crore":  round(net,  2),
                "sentiment":  "BUYING" if net > 0 else "SELLING",
            })
        return result
    except Exception as e:
        return {"error": str(e)}


def get_bulk_block_deals(top_n: int = 20) -> dict:
    """Fetch today's bulk deals and block deals from NSE.

    Bulk deals (> 0.5% of total shares) and block deals (≥ 5 lakh shares
    or ≥ ₹5 crore in value) indicate institutional activity.

    Args:
        top_n: Max number of deals to return (default 20).
    """
    try:
        import time as _time
        s = _get_live_session()
        # Extra warm-up hit needed for this endpoint to return session-gated data
        s.get("https://www.nseindia.com/market-data/bulk-deals", timeout=8)
        _time.sleep(0.5)
        r = s.get(
            "https://www.nseindia.com/api/snapshot-capital-market-largedeal",
            timeout=10,
        )
        d = r.json()
        bulk  = d.get("BULK_DEALS_DATA",  [])[:top_n]
        block = d.get("BLOCK_DEALS_DATA", [])[:top_n]   # block deals

        def _fmt(x: dict) -> dict:
            return {
                "date":        x.get("date"),
                "symbol":      x.get("symbol"),
                "company":     x.get("name"),
                "client":      x.get("clientName"),
                "buy_sell":    x.get("buySell"),
                "qty":         x.get("qty"),
                "price":       x.get("watp"),   # weighted average trade price
                "remarks":     x.get("remarks"),
            }

        return {
            "as_of":       d.get("as_on_date", datetime.now().strftime("%Y-%m-%d")),
            "source":      "NSE live API",
            "bulk_deals":  [_fmt(x) for x in bulk],
            "block_deals": [_fmt(x) for x in block],
        }
    except Exception as e:
        return {"error": str(e)}


def _ratio_pb(ratios: dict) -> str | None:
    """Derive P/B from Current Price and Book Value."""
    price = ratios.get("Current Price", "").replace(",", "")
    bv    = ratios.get("Book Value",    "").replace(",", "")
    try:
        return str(round(float(price) / float(bv), 1)) if price and bv else None
    except (ValueError, ZeroDivisionError):
        return None


# ── E4: Event-Driven Alert Engine ────────────────────────────────────────────

def get_upcoming_events(
    symbols: list[str] | None = None,
    index: str = "NIFTY 50",
    days_ahead: int = 30,
    event_types: list[str] | None = None,
) -> dict:
    """
    E4 Event-Driven Alert Engine: fetch upcoming corporate action events for
    a list of symbols or an entire index.

    Event types tracked:
      • Dividend  — ex-dividend dates, record dates, amount
      • Bonus Issue — bonus ratio (e.g. 1:1, 2:1)
      • Stock Split  — split ratio and record date
      • Rights Issue — entitlement and pricing
      • AGM / EGM   — Annual / Extraordinary General Meeting dates
      • Board Meeting — typically announcing results or dividends
      • Results Calendar — Q1/Q2/Q3/Q4 earnings announcement dates

    Args:
        symbols:     Explicit list of NSE tickers (e.g. ['TCS', 'INFY']).
                     If None, fetches top symbols from the Nifty 50 DB snapshot.
        index:       Index name for auto-symbol lookup when symbols=None.
        days_ahead:  Window in calendar days to consider events "upcoming".
        event_types: Filter to specific types e.g. ['Dividend', 'Results'].
                     None means return all types.

    Returns:
        {
          "events_by_date":  {date_str: [events]},
          "events_by_type":  {type_str: [events]},
          "upcoming":        [...sorted by ex_date...],
          "total":           int,
          "next_7_days":     [...],
          "next_30_days":    [...],
        }
    """
    import urllib.parse
    from datetime import date, timedelta

    today = date.today()
    cutoff = today + timedelta(days=days_ahead)

    # Get symbols list
    if not symbols:
        symbols = _get_index_symbols(index)[:50]  # top 50 from index

    all_events: list[dict[str, Any]] = []

    # Batch NSE corporate actions API calls
    sess = _get_live_session()

    def _fetch_actions(sym: str) -> list[dict]:
        url = (
            f"https://www.nseindia.com/api/corporates-corporateActions"
            f"?index=equities&symbol={urllib.parse.quote(sym)}&issuer="
        )
        try:
            r = sess.get(url, timeout=8)
            if not r.ok:
                return []
            items = r.json()
            if isinstance(items, dict):
                items = items.get("data", [])
            return items or []
        except Exception:
            return []

    from concurrent.futures import ThreadPoolExecutor, as_completed as _as_completed
    with ThreadPoolExecutor(max_workers=5) as pool:
        futures = {pool.submit(_fetch_actions, sym): sym for sym in (symbols or [])}
        for fut in _as_completed(futures):
            sym = futures[fut]
            try:
                items = fut.result()
                for item in items:
                    subject = str(item.get("subject", "")).strip()
                    ex_date_str = item.get("exDate", "") or item.get("ex_date", "")
                    rec_date_str = item.get("recDate", "") or item.get("rec_date", "")

                    # Parse date
                    ex_date = None
                    for fmt in ("%d-%b-%Y", "%Y-%m-%d", "%d/%m/%Y", "%b %d %Y"):
                        try:
                            ex_date = datetime.strptime(ex_date_str, fmt).date()
                            break
                        except (ValueError, TypeError):
                            continue

                    # Classify event type
                    subj_lower = subject.lower()
                    if "dividend" in subj_lower:
                        etype = "Dividend"
                    elif "bonus" in subj_lower:
                        etype = "Bonus"
                    elif "split" in subj_lower or "sub-division" in subj_lower:
                        etype = "Split"
                    elif "rights" in subj_lower:
                        etype = "Rights"
                    elif "agm" in subj_lower or "annual general" in subj_lower:
                        etype = "AGM"
                    elif "egm" in subj_lower or "extra" in subj_lower:
                        etype = "EGM"
                    elif "board meeting" in subj_lower:
                        etype = "Board Meeting"
                    elif "result" in subj_lower:
                        etype = "Results"
                    else:
                        etype = "Corporate Action"

                    if event_types and etype not in event_types:
                        continue

                    days_until = (ex_date - today).days if ex_date else None

                    all_events.append({
                        "symbol":      sym,
                        "type":        etype,
                        "subject":     subject,
                        "ex_date":     ex_date_str,
                        "ex_date_obj": ex_date,
                        "record_date": rec_date_str,
                        "face_value":  item.get("faceVal"),
                        "series":      item.get("series", "EQ"),
                        "days_until":  days_until,
                        "is_upcoming": (
                            ex_date is not None and today <= ex_date <= cutoff
                        ),
                    })
            except Exception:
                pass

    # Filter to upcoming + sort by date
    upcoming = sorted(
        [e for e in all_events if e.get("is_upcoming")],
        key=lambda x: x["ex_date_obj"] or date.max,
    )

    # Clean up non-serializable date objects
    for e in all_events:
        e.pop("ex_date_obj", None)
    for e in upcoming:
        e.pop("ex_date_obj", None)

    # Group by date
    events_by_date: dict[str, list] = {}
    for ev in upcoming:
        d = ev["ex_date"]
        events_by_date.setdefault(d, []).append(ev)

    # Group by type
    events_by_type: dict[str, list] = {}
    for ev in upcoming:
        events_by_type.setdefault(ev["type"], []).append(ev)

    # Next 7 days
    week_cutoff = today + timedelta(days=7)
    next_7 = [
        e for e in upcoming
        if e.get("days_until") is not None and 0 <= e["days_until"] <= 7
    ]

    return {
        "as_of":          today.isoformat(),
        "days_window":    days_ahead,
        "symbols_scanned": len(symbols or []),
        "total":          len(upcoming),
        "events_by_date": events_by_date,
        "events_by_type": events_by_type,
        "upcoming":       upcoming,
        "next_7_days":    next_7,
        "types_found":    list(events_by_type.keys()),
        "source":         "NSE Corporate Actions API (live)",
    }


# ─── Market-wide latest quarterly results feed ─────────────────────────────
_RESULTS_FEED_CACHE: dict[str, Any] = {"ts": 0.0, "data": None}
_RESULTS_FEED_TTL = 1800  # 30 minutes


def _fetch_screener_latest_results() -> list[dict]:
    """Screener.in /results/latest/ fallback parser. Best-effort only."""
    try:
        import requests as _req
        from bs4 import BeautifulSoup  # type: ignore
        r = _req.get(
            "https://www.screener.in/results/latest/",
            timeout=15,
            headers={"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 13_0) "
                                     "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121"},
        )
        if not r.ok:
            return []
        soup = BeautifulSoup(r.text, "html.parser")
        rows: list[dict] = []
        for a in soup.select("a[href^='/company/']")[:200]:
            href = a.get("href", "")
            sym = href.rstrip("/").split("/")[-1].upper()
            name = a.get_text(strip=True)
            if not sym or not name or len(sym) > 12:
                continue
            rows.append({
                "symbol":   sym,
                "company":  name,
                "industry": "",
                "period":   "",
                "filing_date": "",
                "audited":  "",
                "consolidated": "",
                "xbrl_url": "",
                "isin":     "",
                "source":   "screener.in/results/latest",
            })
        return rows
    except Exception:
        return []


def get_latest_results_feed(days_back: int = 7, limit: int = 50) -> dict:
    """
    Market-wide feed of companies that have filed quarterly financial results.

    Primary source: NSE `corporates-financial-results` API (cached 30 min).
    Fallback: screener.in `/results/latest/` HTML scrape.

    When no filings exist within `days_back`, returns the most recent rows
    available with an explanatory `window_note`.

    Args:
        days_back: calendar days back to include (default 7).
        limit:     max rows to return (default 50).
    """
    import requests as _req
    import datetime as _dt
    import time as _time

    out: dict[str, Any] = {
        "results": [], "days_back": days_back, "limit": limit,
        "total_in_window": 0, "total_available": 0,
        "source": None, "window_note": "",
    }

    def _parse_dt(s: str) -> _dt.datetime | None:
        for fmt in ("%d-%b-%Y %H:%M", "%d-%b-%Y %H:%M:%S", "%d-%b-%Y"):
            try:
                return _dt.datetime.strptime(s, fmt)
            except Exception:
                continue
        return None

    # Use cached payload when fresh
    now_ts = _time.time()
    cached = _RESULTS_FEED_CACHE.get("data")
    if cached and (now_ts - float(_RESULTS_FEED_CACHE.get("ts") or 0)) < _RESULTS_FEED_TTL:
        rows = cached
    else:
        rows = None
        try:
            s = _get_live_session()
            url = ("https://www.nseindia.com/api/corporates-financial-results"
                   "?index=equities&period=Quarterly")
            r = s.get(url, timeout=15)
            if r.ok:
                payload = r.json()
                rows = (payload if isinstance(payload, list)
                        else payload.get("data") or payload.get("results") or [])
                _RESULTS_FEED_CACHE["data"] = rows
                _RESULTS_FEED_CACHE["ts"] = now_ts
        except Exception as exc:
            out["nse_error"] = str(exc)

    if not rows:
        screener_rows = _fetch_screener_latest_results()
        if screener_rows:
            out["results"] = screener_rows[: int(limit)]
            out["total_available"] = len(screener_rows)
            out["source"] = "screener.in/results/latest"
            out["window_note"] = "NSE feed unavailable; using screener.in fallback."
        else:
            out["source"] = "none"
            out["window_note"] = "Both NSE and screener.in fallbacks failed."
        return out

    # Decorate and sort by filing date
    decorated: list[dict] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        fd = _parse_dt(str(row.get("filingDate") or row.get("broadCastDate") or ""))
        if not fd:
            continue
        xbrl = row.get("xbrl") or ""
        if isinstance(xbrl, str) and xbrl.endswith("/-"):
            xbrl = ""
        decorated.append({
            "symbol":         row.get("symbol", ""),
            "company":        row.get("companyName", ""),
            "industry":       row.get("industry", "") if row.get("industry") != "-" else "",
            "period":         row.get("relatingTo", ""),
            "from_date":      row.get("fromDate", ""),
            "to_date":        row.get("toDate", ""),
            "financial_year": row.get("financialYear", ""),
            "filing_date":    row.get("filingDate", ""),
            "audited":        row.get("audited", ""),
            "consolidated":   row.get("consolidated", ""),
            "xbrl_url":       xbrl,
            "isin":           row.get("isin", ""),
            "_dt":            fd,
        })
    decorated.sort(key=lambda x: x["_dt"], reverse=True)

    cutoff = _dt.datetime.now() - _dt.timedelta(days=int(days_back))
    in_window = [r_ for r_ in decorated if r_["_dt"] >= cutoff]
    if in_window:
        filtered = in_window
        out["window_note"] = f"Showing filings from last {days_back} day(s)."
    else:
        filtered = decorated[: int(limit)]
        if decorated:
            age_days = (_dt.datetime.now() - decorated[0]["_dt"]).days
            out["window_note"] = (
                f"No filings in last {days_back} day(s); "
                f"showing most recent available (latest filed {age_days} days ago)."
            )
    for f in filtered:
        f.pop("_dt", None)
    out["results"] = filtered[: int(limit)]
    out["total_in_window"] = len(in_window)
    out["total_available"] = len(decorated)
    out["source"] = "nseindia.com/api/corporates-financial-results"
    return out


_FORTHCOMING_RESULTS_CACHE: dict = {"ts": 0.0, "rows": None}


def _fetch_nse_event_calendar() -> list[dict]:
    """Pull NSE event-calendar; 30-min in-process cache."""
    import time as _time
    now = _time.time()
    if _FORTHCOMING_RESULTS_CACHE["rows"] is not None and (now - _FORTHCOMING_RESULTS_CACHE["ts"]) < 1800:
        return _FORTHCOMING_RESULTS_CACHE["rows"]
    try:
        sess = _get_live_session()
        r = sess.get("https://www.nseindia.com/api/event-calendar?index=equities", timeout=15)
        if not r.ok:
            return []
        rows = r.json() if isinstance(r.json(), list) else []
    except Exception:
        return []
    _FORTHCOMING_RESULTS_CACHE["ts"] = now
    _FORTHCOMING_RESULTS_CACHE["rows"] = rows
    return rows


def get_forthcoming_results(days_ahead: int = 14, limit: int = 50) -> dict:
    """Forthcoming board meetings/events that will declare financial results.

    Sources NSE `/api/event-calendar?index=equities` (cached 30 min) and filters
    purposes containing 'Financial Results'. Returns a chronologically sorted
    list of upcoming results events with company-level details.

    Args:
        days_ahead: Calendar days ahead to include (default 14).
        limit:      Max rows to return (default 50).

    Returns:
        dict with `results` (list), `total_in_window`, `total_available`,
        `window_note`, `source`, `days_ahead`.
    """
    import datetime as _dt
    out: dict = {"days_ahead": int(days_ahead), "results": [], "source": ""}
    raw = _fetch_nse_event_calendar()
    if not raw:
        out["error"] = "Could not fetch NSE event-calendar"
        out["source"] = "nseindia.com/api/event-calendar"
        return out

    today = _dt.datetime.now().date()
    horizon = today + _dt.timedelta(days=int(days_ahead))
    decorated: list[dict] = []
    for row in raw:
        purpose = (row.get("purpose") or "").strip()
        if "financial results" not in purpose.lower():
            continue
        date_s = (row.get("date") or "").strip()
        try:
            evt_date = _dt.datetime.strptime(date_s, "%d-%b-%Y").date()
        except Exception:
            continue
        decorated.append({
            "symbol":      (row.get("symbol") or "").strip(),
            "company":     (row.get("company") or "").strip(),
            "purpose":     purpose,
            "description": (row.get("bm_desc") or "").strip(),
            "date":        date_s,
            "_d":          evt_date,
        })
    decorated.sort(key=lambda r: r["_d"])
    in_window = [r for r in decorated if today <= r["_d"] <= horizon]

    if in_window:
        chosen = in_window[: int(limit)]
        out["window_note"] = f"Showing results events in next {days_ahead} day(s)."
    else:
        chosen = decorated[: int(limit)]
        out["window_note"] = (
            f"No results events scheduled within next {days_ahead} day(s); "
            f"showing earliest upcoming entries from feed."
        )
    for r in chosen:
        r.pop("_d", None)
    out["results"] = chosen
    out["total_in_window"] = len(in_window)
    out["total_available"] = len(decorated)
    out["source"] = "nseindia.com/api/event-calendar"
    return out


def get_event_calendar_summary(
    index: str = "NIFTY 50",
    days_ahead: int = 14,
) -> dict:
    """
    Quick event calendar overview for an index — upcoming dividends, splits,
    bonuses, results, and board meetings in the next N days.

    Args:
        index:      Index name (e.g. 'NIFTY 50', 'NIFTY NEXT 50', 'NIFTY 500').
        days_ahead: Calendar days to look ahead (default 14).

    Returns compact event summary suitable for terminal sidebar display.
    """
    full = get_upcoming_events(index=index, days_ahead=days_ahead)
    if full.get("error"):
        return full

    # Compact view for each upcoming event
    compact = []
    for ev in full.get("upcoming", []):
        compact.append({
            "symbol":    ev["symbol"],
            "type":      ev["type"],
            "ex_date":   ev["ex_date"],
            "days_away": ev.get("days_until"),
            "detail":    ev["subject"][:80],
        })

    type_counts = {t: len(v) for t, v in full.get("events_by_type", {}).items()}

    return {
        "index":         index,
        "days_ahead":    days_ahead,
        "total_events":  full.get("total", 0),
        "event_counts":  type_counts,
        "events":        compact,
        "next_7_days":   full.get("next_7_days", []),
        "source":        full.get("source"),
    }


def _get_index_symbols(index: str = "NIFTY 50", top_n: int = 50) -> list[str]:
    """Get symbols for an index from the DB snapshot."""
    syms: list[str] = []
    try:
        rows = _pg_fetchall(
            "SELECT DISTINCT symbol FROM scores.stage_snapshots "
            "WHERE COALESCE(sector,'') != '' "
            "AND snapshot_date = (SELECT MAX(snapshot_date) FROM scores.stage_snapshots) "
            "ORDER BY symbol LIMIT %s",
            (top_n,),
        )
        syms = [r[0] for r in rows]
    except Exception:
        try:
            if _legacy_sqlite_fallbacks_enabled() and DB_PATH.exists():
                conn = _db_conn()
                rows = conn.execute(
                    "SELECT DISTINCT symbol FROM stage_snapshots "
                    "WHERE LOWER(sector) != '' "
                    "AND snapshot_date = (SELECT MAX(snapshot_date) FROM stage_snapshots) "
                    "ORDER BY investment_score DESC NULLS LAST LIMIT ?",
                    (top_n,),
                ).fetchall()
                conn.close()
                syms = [r[0] for r in rows]
        except Exception:
            pass
    # Fallback to a small hardcoded Nifty 50 subset
    if not syms:
        syms = [
            "RELIANCE", "TCS", "HDFCBANK", "INFY", "ICICIBANK", "HINDUNILVR",
            "ITC", "SBIN", "BHARTIARTL", "BAJFINANCE", "KOTAKBANK", "LT",
            "AXISBANK", "ASIANPAINT", "MARUTI", "TITAN", "SUNPHARMA", "WIPRO",
            "ULTRACEMCO", "HCLTECH", "NESTLEIND", "POWERGRID", "NTPC",
            "TATAMOTORS", "TATASTEEL", "M&M", "TECHM", "ONGC", "DRREDDY",
            "DIVISLAB", "ADANIPORTS", "BAJAJFINSV", "CIPLA", "HEROMOTOCO",
            "EICHERMOT", "GRASIM", "BRITANNIA", "COALINDIA", "INDUSINDBK",
            "JSWSTEEL", "BPCL", "APOLLOHOSP", "TATACONSUM", "LTIM",
            "HINDALCO", "UPL", "SBILIFE", "HDFCLIFE", "BAJAJ-AUTO", "VEDL",
        ][:top_n]
    return syms


# ── B3: Sectoral Heat Calendar ────────────────────────────────────────────────

_SECTOR_INDEX_MAP = {
    "NIFTY IT":                 "IT",
    "NIFTY BANK":               "Banking",
    "NIFTY PHARMA":             "Pharma",
    "NIFTY AUTO":               "Auto",
    "NIFTY FMCG":               "FMCG",
    "NIFTY METAL":              "Metals",
    "NIFTY REALTY":             "Realty",
    "NIFTY ENERGY":             "Energy",
    "NIFTY FINANCIAL SERVICES": "Financials",
    "NIFTY INFRA":              "Infra",
}

# yfinance tickers for NSE sector indices (when local CSV lacks them)
_SECTOR_YF_MAP = {
    "NIFTY IT":                 "^CNXIT",
    "NIFTY BANK":               "^NSEBANK",
    "NIFTY PHARMA":             "^CNXPHARMA",
    "NIFTY AUTO":               "^CNXAUTO",
    "NIFTY FMCG":               "^CNXFMCG",
    "NIFTY METAL":              "^CNXMETAL",
    "NIFTY REALTY":             "^CNXREALTY",
    "NIFTY ENERGY":             "^CNXENERGY",
    "NIFTY FINANCIAL SERVICES": "^CNXFIN",
    "NIFTY INFRA":              "^CNXINFRA",
}


def _fetch_sector_monthly_returns_yf(lookback_years: int = 7) -> pd.DataFrame:
    """Fetch monthly returns for sector indices via yfinance (fallback path)."""
    import yfinance as yf

    records = []
    for index_name, sector in _SECTOR_INDEX_MAP.items():
        ticker = _SECTOR_YF_MAP.get(index_name)
        if not ticker:
            continue
        try:
            df = yf.download(ticker, period=f"{lookback_years}y", interval="1mo",
                             progress=False, auto_adjust=True)
            if df.empty:
                continue
            closes = df["Close"].squeeze().dropna()
            monthly_ret = closes.pct_change().dropna() * 100
            for period, ret in monthly_ret.items():
                records.append({
                    "sector":    sector,
                    "month_num": period.month,
                    "return_pct": float(ret),
                })
        except Exception:
            continue

    if not records:
        return pd.DataFrame(columns=["sector", "month_num", "return_pct"])
    return pd.DataFrame(records)


def get_sector_heat_calendar(month: int | None = None) -> dict:
    """
    B3 Sectoral Heat Calendar — monthly seasonal return heatmap for NSE sectors.

    Builds a 12-month × 10-sector matrix of average historical monthly returns
    and classifies each sector-month as TAILWIND (>+2%), HEADWIND (<-1%), or NEUTRAL.

    Args:
        month: Target month number 1-12 (default: current month).

    Returns:
        {
          "current_month": {"IT": "TAILWIND", "Banking": "HEADWIND", ...},
          "heatmap": {sector: {month_name: avg_pct, ...}, ...},
          "current_month_name": "May",
          "tailwinds": [...],
          "headwinds": [...],
          "neutral": [...],
          "note": str,
        }
    """
    target_month = month or datetime.now().month
    month_name = ["Jan","Feb","Mar","Apr","May","Jun",
                  "Jul","Aug","Sep","Oct","Nov","Dec"][target_month - 1]

    # Try local index CSV first, fall back to yfinance
    try:
        matrix, heat = build_seasonal_heat_calendar(_SECTOR_INDEX_MAP)
        if heat.empty:
            raise ValueError("Local index CSV lacks sector data — using yfinance")
        use_yf = False
    except Exception:
        heat = None
        use_yf = True

    if use_yf:
        raw = _fetch_sector_monthly_returns_yf()
        if raw.empty:
            return {"error": "Unable to fetch sector index data for heat calendar"}
        heat = (
            raw.groupby(["sector", "month_num"])["return_pct"]
            .agg(avg="mean", std="std", n="count")
            .reset_index()
            .rename(columns={"avg": "avg", "std": "std", "n": "n"})
        )
        heat.columns = ["sector", "month_num", "avg", "std", "n"]

    # Current-month signals
    _TAILWIND = 2.0
    _HEADWIND = -1.0
    _MIN_OBS = 5
    signals = {}
    for _, row in heat[heat["month_num"] == target_month].iterrows():
        if row["n"] < _MIN_OBS:
            signals[row["sector"]] = "NEUTRAL"
        elif row["avg"] > _TAILWIND:
            signals[row["sector"]] = "TAILWIND"
        elif row["avg"] < _HEADWIND:
            signals[row["sector"]] = "HEADWIND"
        else:
            signals[row["sector"]] = "NEUTRAL"

    # Full heatmap
    month_names = ["Jan","Feb","Mar","Apr","May","Jun",
                   "Jul","Aug","Sep","Oct","Nov","Dec"]
    heatmap: dict[str, dict] = {}
    for _, row in heat.iterrows():
        sec = row["sector"]
        mn  = int(row["month_num"])
        if sec not in heatmap:
            heatmap[sec] = {}
        heatmap[sec][month_names[mn - 1]] = round(float(row["avg"]), 2)

    tailwinds = [s for s, sig in signals.items() if sig == "TAILWIND"]
    headwinds = [s for s, sig in signals.items() if sig == "HEADWIND"]
    neutral   = [s for s, sig in signals.items() if sig == "NEUTRAL"]

    return {
        "current_month":      month_name,
        "current_month_num":  target_month,
        "current_month_signals": signals,
        "tailwinds":          tailwinds,
        "headwinds":          headwinds,
        "neutral":            neutral,
        "heatmap":            heatmap,
        "source":             "yfinance NSE sector indices"
                              if use_yf
                              else "PostgreSQL market.index_eod / seasonal cache (CSV fallback)",
        "note": (
            f"Seasonal analysis based on 7yr history. "
            f"{len(tailwinds)} sector(s) show historical tailwind in {month_name}."
        ),
    }


# ── B5: Economic Cycle Tracker ────────────────────────────────────────────────

def get_economic_cycle_assessment() -> dict:
    """
    B5 Economic Cycle Tracker — detect current macro cycle phase and sector positioning.

    Reads macro proxy signals (USD/INR, Brent, Copper, US10Y, VIX, CPI, PMI, IIP, GST)
    from the local data store and classifies the market into one of four phases:
      • EARLY_EXPANSION   — growth improving, rates/inflation contained
      • LATE_EXPANSION    — growth firm but inflation/commodities/rates rising
      • SLOWDOWN          — trend weakens, volatility/inflation rise
      • RECOVERY          — risk eases, trend starting to improve

    Returns:
        {
          "cycle_phase":        str,
          "confidence":         float (0-1),
          "definition":         str,
          "preferred_sectors":  list[str],
          "avoid_sectors":      list[str],
          "alignment":          {indicator: score, ...},
          "macro_snapshot":     {indicator: {signal, value, direction}, ...},
          "note":               str,
        }
    """
    macro_csv = ROOT / "data" / "macro_proxy_signals.csv"
    if not macro_csv.exists():
        return {"error": "macro_proxy_signals.csv not found — run fetch_macro_proxies.py first"}

    try:
        df = pd.read_csv(str(macro_csv))
    except Exception as e:
        return {"error": f"Failed to read macro signals: {e}"}

    try:
        result = detect_economic_cycle_phase(df)
    except Exception as e:
        return {"error": f"Cycle detection failed: {e}"}

    # Build macro snapshot for display
    macro_snapshot = {}
    for _, row in df.iterrows():
        ind = str(row.get("indicator", "")).strip()
        if not ind:
            continue
        macro_snapshot[ind] = {
            "signal":    str(row.get("signal",    "")),
            "value":     str(row.get("value",     "")),
            "direction": str(row.get("direction", "")),
        }

    return {
        "cycle_phase":       result.get("cycle_phase"),
        "confidence":        result.get("confidence"),
        "definition":        result.get("definition", ""),
        "preferred_sectors": result.get("preferred_sectors", []),
        "avoid_sectors":     result.get("avoid_sectors", []),
        "alignment":         result.get("alignment", {}),
        "macro_snapshot":    macro_snapshot,
        "updated":           str(pd.Timestamp.now().date()),
        "note": (
            f"Cycle phase: {result.get('cycle_phase')} "
            f"(confidence {result.get('confidence', 0):.0%}). "
            f"Prefer: {', '.join(result.get('preferred_sectors', [])[:3])}."
        ),
    }


# ── D4: Concall NLP Sentiment Engine ─────────────────────────────────────────

def analyze_concall_sentiment(symbol: str) -> dict:
    """
    D4 Concall NLP — extract management tone, key themes, and sentiment from
    the most recent earnings call / investor day transcript for a stock.

    Fetches transcript content via the deep search engine, then uses the
    OpenAI LLM to extract:
      • Overall sentiment (Bullish / Cautious / Bearish / Neutral)
      • Tone score (-1.0 to +1.0)
      • Top 3-5 key themes from management (growth drivers, concerns, guidance)
      • Risk flags (revenue miss, margin pressure, write-offs, promoter concerns)
      • Key quotes (direct management statements worth tracking)

    Args:
        symbol: NSE ticker symbol (e.g. 'TCS', 'RELIANCE').

    Returns:
        {
          "symbol":     str,
          "sentiment":  "Bullish" | "Cautious" | "Bearish" | "Neutral",
          "tone_score": float,
          "themes":     list[str],
          "risk_flags": list[str],
          "key_quotes": list[str],
          "guidance":   str,
          "transcript_source": str,
          "note":       str,
        }
    """
    import openai, os, json as _json

    sym = symbol.upper().strip()

    # Step 1: Fetch transcript content
    concall_data = search_concall_transcripts(sym)
    screener_data = {}
    try:
        from terminal.web_research import scrape_screener_in
        screener_data = scrape_screener_in(sym)
    except Exception:
        pass

    # Gather concall text
    transcript_text = ""
    transcript_source = "N/A"
    for item in concall_data.get("results", []):
        snip = item.get("snippet", "") or item.get("body", "") or ""
        transcript_text += "\n" + snip
        if not transcript_source or transcript_source == "N/A":
            transcript_source = item.get("url", "N/A")

    # Also pull from screener concalls list if available
    concalls = screener_data.get("concalls", [])
    if concalls:
        for c in concalls[:3]:
            transcript_text += f"\n[Transcript: {c.get('date','')} {c.get('title','')}]"
            if transcript_source == "N/A" and c.get("ppt_url"):
                transcript_source = c["ppt_url"]

    if not transcript_text.strip():
        return {
            "symbol": sym,
            "error": "No concall transcript content found. Try /search SYMBOL concall for manual lookup.",
        }

    # Step 2: LLM extraction
    api_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        return {"symbol": sym, "error": "OPENAI_API_KEY not set — cannot run NLP extraction."}

    client = openai.OpenAI(api_key=api_key) if os.environ.get("OPENAI_API_KEY") else None
    if not client:
        return {"symbol": sym, "error": "OpenAI client unavailable."}

    prompt = f"""You are a buy-side analyst specializing in earnings call analysis.

Analyze the following concall excerpts for {sym} and return a JSON object with:
- "sentiment": "Bullish" | "Cautious" | "Bearish" | "Neutral"
- "tone_score": float from -1.0 (very bearish) to +1.0 (very bullish)
- "themes": list of 3-5 key themes management discussed (growth drivers, product launches, capex, guidance)
- "risk_flags": list of specific risks mentioned (margin pressure, write-offs, slower demand, pricing pressure)
- "key_quotes": list of 2-4 direct or paraphrased management quotes worth tracking
- "guidance": one-line summary of forward guidance given

Return ONLY the JSON object, no markdown fences.

TRANSCRIPT EXCERPTS:
{transcript_text[:3000]}
"""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=600,
        )
        raw = resp.choices[0].message.content.strip()
        parsed = _json.loads(raw)
    except Exception as e:
        return {"symbol": sym, "error": f"LLM extraction failed: {e}", "raw_snippets": transcript_text[:500]}

    return {
        "symbol":            sym,
        "sentiment":         parsed.get("sentiment", "Neutral"),
        "tone_score":        parsed.get("tone_score", 0.0),
        "themes":            parsed.get("themes", []),
        "risk_flags":        parsed.get("risk_flags", []),
        "key_quotes":        parsed.get("key_quotes", []),
        "guidance":          parsed.get("guidance", ""),
        "transcript_source": transcript_source,
        "note": f"NLP analysis of {sym} concall via LLM. Verify with original transcript.",
    }


# ── P2-2: Scenario Analysis Engine ───────────────────────────────────────────

def run_scenario_analysis(
    symbol: str,
    price_scenarios: list[float] | None = None,
    scenario_labels: list[str] | None = None,
) -> dict:
    """
    P2-2 Scenario Engine — what-if price analysis for a held stock.

    For each hypothetical price level, calculates:
      • % change from current price
      • New implied RSI (interpolated from historical distribution)
      • New RS rank vs NIFTY (estimated)
      • Stage implication (Stage 1/2/3/4 boundary crossings)
      • Support/resistance proximity
      • Risk/reward vs current technical levels

    Args:
        symbol:          NSE ticker (e.g. 'TCS').
        price_scenarios: List of price levels to evaluate.
                         Defaults to ±5%, ±10%, ±20% from current price.
        scenario_labels: Optional labels for each scenario (e.g. ['Base', 'Bull', 'Bear']).

    Returns:
        {
          "symbol":          str,
          "current_price":   float,
          "scenarios":       [{label, price, pct_chg, stage_implication, rsi_est, note}, ...],
          "key_levels":      {support, resistance, ma50, ma200},
          "current_rsi":     float,
          "current_stage":   str,
        }
    """
    sym = symbol.upper().strip()

    # Get current snapshot
    snap = get_symbol_snapshot(sym)
    if snap.get("error"):
        return {"error": snap["error"], "symbol": sym}

    current_price = snap.get("last_price") or snap.get("price") or snap.get("close")
    if not current_price:
        return {"error": "No current price available", "symbol": sym}

    current_price = float(current_price)
    current_rsi   = float(snap.get("rsi") or 50.0)
    current_stage = str(snap.get("stage") or "Unknown")
    ma50  = float(snap.get("ma50")  or current_price)
    ma200 = float(snap.get("ma200") or current_price)

    # Load price history for support/resistance
    hist = _load_price_history(sym, days=200)
    support    = float(hist["LOW"].quantile(0.15))  if not hist.empty else current_price * 0.90
    resistance = float(hist["HIGH"].quantile(0.85)) if not hist.empty else current_price * 1.10

    # Build price scenarios
    if not price_scenarios:
        price_scenarios = [
            round(current_price * m, 2)
            for m in [0.80, 0.90, 0.95, 1.00, 1.05, 1.10, 1.20]
        ]
        scenario_labels = ["Bear -20%", "Bear -10%", "Dip -5%", "Current",
                           "Bull +5%", "Bull +10%", "Bull +20%"]

    if not scenario_labels:
        scenario_labels = [f"S{i+1}" for i in range(len(price_scenarios))]

    scenarios = []
    for label, price in zip(scenario_labels, price_scenarios):
        price = float(price)
        pct   = (price - current_price) / current_price * 100

        # Estimate RSI at scenario price (linear approximation from current distribution)
        rsi_delta = pct * 0.8  # rough: ±10% price → ±8 RSI points
        rsi_est = min(100, max(0, current_rsi + rsi_delta))

        # Stage implication
        if price > max(ma50, ma200) * 1.02:
            stage_impl = "Stage 2 (uptrend) — above both MAs"
        elif price > ma200 and price < ma50:
            stage_impl = "Stage 1/2 transition — above 200MA, below 50MA"
        elif price < ma200 and price > ma200 * 0.95:
            stage_impl = "Stage 3 warning — near/at 200MA"
        elif price < ma200 * 0.95:
            stage_impl = "Stage 4 (downtrend) — below 200MA"
        else:
            stage_impl = "Stage transition zone"

        # Key observations
        notes = []
        if abs(price - support) / current_price < 0.03:
            notes.append(f"Near support ₹{support:.0f}")
        if abs(price - resistance) / current_price < 0.03:
            notes.append(f"Near resistance ₹{resistance:.0f}")
        if abs(price - ma50) / current_price < 0.02:
            notes.append(f"Near 50-DMA ₹{ma50:.0f}")
        if abs(price - ma200) / current_price < 0.02:
            notes.append(f"At 200-DMA ₹{ma200:.0f} — critical level")

        scenarios.append({
            "label":            label,
            "price":            round(price, 2),
            "pct_change":       round(pct, 1),
            "rsi_estimate":     round(rsi_est, 1),
            "stage_implication": stage_impl,
            "notes":            "; ".join(notes) if notes else "No key level proximity",
        })

    return {
        "symbol":        sym,
        "current_price": current_price,
        "current_rsi":   current_rsi,
        "current_stage": current_stage,
        "key_levels": {
            "support":    round(support, 2),
            "resistance": round(resistance, 2),
            "ma50":       round(ma50, 2),
            "ma200":      round(ma200, 2),
        },
        "scenarios": scenarios,
        "note": f"What-if scenario analysis for {sym}. RSI estimates are approximate linear projections.",
    }


# ── P2-4: Portfolio Narrative Generator ──────────────────────────────────────

def generate_portfolio_narratives(
    symbols: list[str] | None = None,
    top_n: int = 5,
) -> dict:
    """
    P2-4 Portfolio Narrative Engine — per-stock LLM commentary combining technical,
    fundamental, and market context into a concise investment narrative.

    For each stock in the portfolio (or provided list), synthesises:
      • Current stage + RSI + key signals
      • Fundamental snapshot (P/E, ROE, ROCE from screener.in)
      • Recent news catalyst
      • 2-3 sentence investment thesis (bull) and 1-2 sentence bear case

    Args:
        symbols: Explicit list of symbols. If None, uses top portfolio holdings by score.
        top_n:   Number of stocks to narrate (default 5).

    Returns:
        {
          "narratives": [{symbol, thesis, bear_case, signals, action_hint}, ...],
          "generated_at": str,
        }
    """
    # Get symbols from portfolio if not provided
    if not symbols:
        port = get_portfolio_exposure()
        if port.get("error"):
            return {"error": "No symbols provided and portfolio is empty."}
        holdings = port.get("holdings", [])
        symbols = [h["symbol"] for h in holdings[:top_n] if h.get("symbol")]

    if not symbols:
        return {"error": "No symbols to narrate"}

    narratives = []
    for sym in symbols[:top_n]:
        sym = sym.upper().strip()
        snap = get_symbol_snapshot(sym)
        if snap.get("error"):
            narratives.append({"symbol": sym, "error": snap["error"]})
            continue

        stage   = snap.get("stage", "Unknown")
        rsi     = snap.get("rsi")
        rs_pct  = snap.get("rs_pct")
        signal  = snap.get("trading_signal", "")
        score   = snap.get("investment_score")
        price   = snap.get("last_price") or snap.get("price") or snap.get("close")
        ma50    = snap.get("ma50")
        ma200   = snap.get("ma200")

        # Compute MA50/MA200 from price history if not in snapshot
        if not ma50 or not ma200:
            try:
                _hist = _load_price_history(sym, days=220)
                if not _hist.empty:
                    _c = _hist["CLOSE"].dropna()
                    if not ma50 and len(_c) >= 50:
                        ma50 = round(float(_c.iloc[-50:].mean()), 2)
                    if not ma200 and len(_c) >= 200:
                        ma200 = round(float(_c.iloc[-200:].mean()), 2)
            except Exception:
                pass
        ma50_str  = f"₹{ma50:,.0f}" if ma50 else "50-DMA"
        ma200_str = f"₹{ma200:,.0f}" if ma200 else "200-DMA"

        # Build brief context for LLM (no API call to screener.in to keep this fast)
        context_lines = [
            f"Stock: {sym}",
            f"Stage: {stage}  |  RSI: {rsi}  |  RS%: {rs_pct}",
            f"Signal: {signal}  |  Investment Score: {score}",
            f"Price: {price}  |  50-DMA: {ma50}  |  200-DMA: {ma200}",
        ]

        # Signals
        signals_list = []
        if rsi and float(str(rsi).replace(",", "") or 0) > 70:
            signals_list.append("RSI overbought")
        elif rsi and float(str(rsi).replace(",", "") or 0) < 35:
            signals_list.append("RSI oversold")
        if "Stage 2" in str(stage):
            signals_list.append("Stage 2 uptrend")
        elif "Stage 4" in str(stage):
            signals_list.append("Stage 4 downtrend — caution")
        if rs_pct and float(str(rs_pct).replace(",", "") or 0) > 0:
            signals_list.append(f"Outperforming market (+{rs_pct}% RS)")

        # Craft narrative using rules-based logic (no API call needed for basic output)
        stage_upper = str(stage).upper()
        if ("STAGE_2" in stage_upper or "STAGE 2" in stage_upper) and signals_list:
            thesis = (
                f"{sym} is in a confirmed Stage 2 uptrend with RSI at {rsi}, "
                f"outperforming the broader market by {rs_pct}%. "
                f"Current signal: {signal}. Trend intact while above {ma50_str}."
            )
            bear_case = f"Break below {ma50_str} or Stage 2→3 transition would invalidate the setup."
            action = "Hold / Add on dips above 50-DMA"
        elif "STAGE_4" in stage_upper or "STAGE 4" in stage_upper:
            thesis = (
                f"{sym} is in Stage 4 downtrend. RSI: {rsi}. "
                f"Price below both key MAs. Avoid new positions."
            )
            bear_case = f"Further downside likely unless reclaims {ma200_str}."
            action = "Avoid / Exit on bounces"
        elif "STAGE_1" in stage_upper or "STAGE 1" in stage_upper:
            thesis = (
                f"{sym} is basing (Stage 1). RSI: {rsi}. "
                f"Watching for Stage 2 breakout above {ma50_str} with volume."
            )
            bear_case = "No catalyst yet. Monitor for breakout confirmation."
            action = "Watch / Small starter if breadth improves"
        elif "STAGE_3" in stage_upper or "STAGE 3" in stage_upper:
            thesis = (
                f"{sym} is in Stage 3 topping phase. RSI: {rsi}. "
                f"Distribution or early breakdown — caution warranted."
            )
            bear_case = f"Potential Stage 3→4 transition. Reduce or hedge below {ma50_str}."
            action = "Reduce / Tighten stops"
        else:
            thesis = f"{sym} — Stage: {stage}, RSI: {rsi}, Signal: {signal}. Monitor closely."
            bear_case = "Unclear trend — wait for direction."
            action = "Neutral / Wait for clarity"

        narratives.append({
            "symbol":      sym,
            "stage":       stage,
            "rsi":         rsi,
            "rs_pct":      rs_pct,
            "price":       price,
            "signals":     signals_list,
            "thesis":      thesis,
            "bear_case":   bear_case,
            "action_hint": action,
        })

    return {
        "narratives":    narratives,
        "total_stocks":  len(narratives),
        "generated_at":  str(pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")),
        "note": f"Portfolio narratives for {len(narratives)} stock(s) based on EOD snapshot + stage analysis.",
    }


# ── P3-2: Voice Briefing Generator ───────────────────────────────────────────

def generate_voice_briefing(
    text: str | None = None,
    voice: str = "cedar",
    save_path: str | None = None,
) -> dict:
    """
    P3-2 Daily Voice Briefing — convert market summary to audio using OpenAI GPT TTS.

    If no text is provided, auto-generates a 60-second market briefing using
    the current day's live market overview, top movers, and portfolio status.

    Args:
        text:      Custom text to convert to speech. If None, generates auto-briefing.
        voice:     OpenAI GPT TTS voice. Defaults to 'cedar' for a grounded assistant tone.
        save_path: Where to save the MP3 file. Defaults to data/voice_briefing.mp3.

    Returns:
        {
          "audio_file": str path,
          "text_used":  str,
          "duration_est": str,
          "voice":      str,
          "note":       str,
        }
    """
    import os
    from pathlib import Path as _Path

    api_key = os.environ.get("OPENAI_API_KEY") or os.environ.get("SHUNYAAI_OPENAI_API_KEY", "")
    # Auto-generate briefing text if not provided
    if not text:
        try:
            overview = get_live_market_overview()
            mkt_mood = overview.get("market_mood", "mixed")
            nifty_chg = overview.get("nifty_change_pct", 0)
            top_gain = overview.get("top_gainers", [{}])[:2]
            top_loss = overview.get("top_losers",  [{}])[:2]
            gainers_str = ", ".join(g.get("symbol","") for g in top_gain if g.get("symbol"))
            losers_str  = ", ".join(g.get("symbol","") for g in top_loss if g.get("symbol"))
        except Exception:
            overview = {}
            mkt_mood = "mixed"
            nifty_chg = 0
            gainers_str = ""
            losers_str  = ""

        sign = "up" if float(nifty_chg or 0) >= 0 else "down"
        text = (
            f"Good morning. Here is your Agent Adda market briefing. "
            f"NIFTY 50 is {sign} {abs(float(nifty_chg or 0)):.2f} percent. "
            f"Market mood is {mkt_mood}. "
            + (f"Top gainers today include {gainers_str}. " if gainers_str else "")
            + (f"Key stocks under pressure: {losers_str}. " if losers_str else "")
            + "Check Agent Adda for full analysis. Have a great trading day."
        )

    out_path = _Path(save_path) if save_path else ROOT / "data" / "voice_briefing.mp3"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    model = os.environ.get("OPENAI_TTS_MODEL", "gpt-4o-mini-tts")
    provider = "openai_gpt_tts"
    if api_key:
        try:
            from openai import OpenAI

            client = OpenAI(api_key=api_key)
            with client.audio.speech.with_streaming_response.create(
                model=model,
                voice=voice,
                input=text,
                instructions=(
                    "Speak like a calm senior Indian-market operator. "
                    "Be concise, risk-first, evidence-aware, and avoid hype. "
                    "This is an AI-generated research briefing, not investment advice."
                ),
            ) as response:
                response.stream_to_file(str(out_path))
        except Exception as e:
            return {"error": f"TTS generation failed: {e}", "text_used": text, "model": model}
    else:
        import subprocess

        provider = "macos_say"
        out_path = out_path.with_suffix(".aiff")
        try:
            subprocess.run(["say", "-o", str(out_path), text], check=True, capture_output=True)
        except Exception as e:
            return {
                "error": f"OPENAI_API_KEY not set and macOS say fallback failed: {e}",
                "text_used": text,
                "model": model,
            }

    word_count = len(text.split())
    duration_secs = word_count / 2.5  # approx 150 wpm
    duration_str = f"~{duration_secs:.0f}s"

    return {
        "audio_file":    str(out_path),
        "text_used":     text,
        "duration_est":  duration_str,
        "voice":         voice,
        "model":         model,
        "provider":      provider,
        "note": f"Voice briefing saved to {out_path.name}. Play with: open '{out_path}'",
    }


def compare_stocks(
    symbols: list[str],
    aspects: list[str] | None = None,
) -> dict:
    """Compare multiple NSE stocks side-by-side on technical AND fundamental metrics.

    Combines EOD DB snapshot (stage, RS, RSI, signals, scores) with screener.in ratios
    (P/E, ROE, ROCE, market cap, P/B, dividend yield) into one unified comparison table.

    Args:
        symbols: List of NSE ticker symbols, e.g. ['TCS', 'INFY', 'WIPRO'].
        aspects: Optional list; 'technical', 'fundamental', or 'both' (default).
    """
    if not symbols:
        return {"error": "No symbols provided"}

    fetch_tech = True
    fetch_fund = True
    if aspects:
        fetch_tech = any(a in ("technical", "both") for a in aspects)
        fetch_fund = any(a in ("fundamental", "both") for a in aspects)

    rows: list[dict] = []

    for raw in symbols:
        input_sym = raw.strip().upper()
        resolution = resolve_symbol(input_sym)
        sym = str(resolution.get("symbol") or input_sym).strip().upper()
        row: dict = {
            "symbol": sym,
            "input_symbol": input_sym,
            "resolved_symbol": resolution.get("symbol"),
            "symbol_resolution_confidence": resolution.get("confidence"),
            "missing_evidence": [],
        }
        if not resolution.get("symbol"):
            row["resolution_error"] = resolution.get("error") or f"No exact NSE symbol found for '{input_sym}'"
            row["missing_evidence"].append("symbol_resolution")
            row["evidence_coverage"] = "missing"
            rows.append(row)
            continue

        if fetch_tech:
            try:
                snap = get_symbol_snapshot(sym)
                if not snap.get("error"):
                    row.update({
                        "company":          snap.get("company_name", sym),
                        "stage":            snap.get("stage"),
                        "rsi":              snap.get("rsi"),
                        "rs_pct":           snap.get("rs_pct"),
                        "technical_score":  snap.get("technical_score"),
                        "investment_score": snap.get("investment_score"),
                        "trading_signal":   snap.get("trading_signal"),
                        "trend_signal":     snap.get("trend_signal"),
                        "supertrend":       snap.get("supertrend_state"),
                        "sector":           snap.get("sector"),
                        "change_1d_pct":    snap.get("change_1d_pct"),
                        "change_1w_pct":    snap.get("change_1w_pct"),
                        "change_1m_pct":    snap.get("change_1m_pct"),
                        "db_price":         snap.get("price"),
                        "snapshot_date":    snap.get("snapshot_date"),
                        "stance":           snap.get("stance"),
                        "narrative":        snap.get("narrative"),
                    })
                    row["missing_evidence"].extend(snap.get("missing_evidence") or [])
                else:
                    row["tech_error"] = snap["error"]
                    row["missing_evidence"].extend(snap.get("missing_evidence") or ["technical_snapshot"])
            except Exception as e:
                row["tech_error"] = str(e)
                row["missing_evidence"].append("technical_snapshot")

        if fetch_fund:
            try:
                sr = scrape_screener_in(sym)
                if not sr.get("error"):
                    ratios = sr.get("ratios", {})
                    row.update({
                        "pe":            ratios.get("Stock P/E"),
                        "pb":            _ratio_pb(ratios),
                        "roe":           ratios.get("ROE"),
                        "roce":          ratios.get("ROCE"),
                        "div_yield":     ratios.get("Dividend Yield"),
                        "market_cap_cr": ratios.get("Market Cap"),
                        "book_value":    ratios.get("Book Value"),
                        "current_price": ratios.get("Current Price"),
                        "high_low_52w":  ratios.get("High / Low"),
                        "screener_url":  sr.get("source_url"),
                        "pros":          sr.get("pros", [])[:3],
                        "cons":          sr.get("cons", [])[:2],
                    })
                else:
                    row["fund_error"] = sr["error"]
                    row["missing_evidence"].append("screener_fundamentals")
            except Exception as e:
                row["fund_error"] = str(e)
                row["missing_evidence"].append("screener_fundamentals")

        row["missing_evidence"] = list(dict.fromkeys(row.get("missing_evidence") or []))
        row["evidence_coverage"] = "complete" if not row["missing_evidence"] else "partial"

        rows.append(row)

    tech_cols = ["stage", "rsi", "rs_pct", "technical_score", "investment_score",
                 "trading_signal", "trend_signal", "supertrend",
                 "change_1d_pct", "change_1w_pct", "change_1m_pct"]
    fund_cols = ["pe", "pb", "roe", "roce", "div_yield", "market_cap_cr"]
    active_cols = (tech_cols if fetch_tech else []) + (fund_cols if fetch_fund else [])

    comparison_table = {col: {r["symbol"]: r.get(col) for r in rows} for col in active_cols}

    return {
        "symbols":          [r["symbol"] for r in rows],
        "input_symbols":    [str(s).strip().upper() for s in symbols if str(s).strip()],
        "unresolved_symbols": [r["input_symbol"] for r in rows if "symbol_resolution" in (r.get("missing_evidence") or [])],
        "as_of":            date.today().isoformat(),
        "aspects":          aspects or ["both"],
        "comparison_table": comparison_table,
        "stock_details":    rows,
        "missing_evidence": list(dict.fromkeys(
            ev
            for r in rows
            for ev in (r.get("missing_evidence") or [])
        )),
        "evidence_coverage": "complete" if all(not (r.get("missing_evidence") or []) for r in rows) else "partial",
        "source_trail": [
            "resolve_symbol",
            *(["get_symbol_snapshot"] if fetch_tech else []),
            *(["scrape_screener_in"] if fetch_fund else []),
        ],
    }



# ── Intraday screener wrapper ────────────────────────────────────────────────

def scan_intraday_market(
    index: str = "NIFTY 50",
    interval: str = "15m",
    strategies: list[str] | None = None,
    direction_filter: str = "all",
    min_rr: float = 1.5,
    top_n: int = 5,
) -> dict:
    """Scan all stocks in an NSE index for intraday trading signals.

    Args:
        index:            Index name: NIFTY 50, NIFTY BANK, NIFTY IT, NIFTY 500, NIFTY PHARMA...
        interval:         Candle interval: 5m, 15m, 30m, 1h.
        strategies:       Strategy keys: macd, rsi, supertrend, bollinger, ema, vcp, volume.
        direction_filter: buy, sell, or all.
        min_rr:           Minimum R:R ratio (default 1.5).
        top_n:            Top signals to highlight.
    """
    # PG-SCAN-FALLBACK: use the shared helper (NSE live + local CSV fallback)
    try:
        stocks = _fetch_nse_index_constituents(index)
    except Exception as e:
        return {"error": f"Could not fetch {index} constituents: {e}"}

    if not stocks:
        return {"error": f"No stocks found for index: {index}"}

    result = _run_intraday_screener(
        symbols=stocks,
        interval=interval,
        strategies=strategies,
        direction_filter=direction_filter,
        min_rr=min_rr,
    )
    result["index"]    = index
    result["top_buy"]  = result["buy_signals"][:top_n]
    result["top_sell"] = result["sell_signals"][:top_n]
    return result


def scan_symbols_intraday(
    symbols: list[str],
    interval: str = "15m",
    strategies: list[str] | None = None,
    direction_filter: str = "all",
    min_rr: float = 1.3,
    top_n: int = 10,
) -> dict:
    """Scan a specific list of NSE stocks for intraday signals.

    Use this when you already have a list of stocks from EOD screening,
    breakout analysis, or any source — NOT just index constituents.
    Returns ranked BUY/SELL signals with entry/target/SL/R:R.

    Args:
        symbols:          List of NSE tickers, e.g. ['RELIANCE', 'TCS', 'CHENNPETRO'].
        interval:         Candle interval: 5m, 15m, 30m, 1h.
        strategies:       Strategy keys: macd, rsi, supertrend, bollinger, ema, vcp, volume. None = all.
        direction_filter: 'buy', 'sell', or 'all'.
        min_rr:           Minimum Risk:Reward ratio (default 1.3).
        top_n:            Max signals to return.
    """
    if not symbols:
        return {"error": "No symbols provided"}
    # Clean and deduplicate
    clean_syms = list(dict.fromkeys(s.strip().upper() for s in symbols if s.strip()))
    result = _run_intraday_screener(
        symbols=clean_syms,
        interval=interval,
        strategies=strategies,
        direction_filter=direction_filter,
        min_rr=min_rr,
    )
    result["symbols_scanned"] = clean_syms
    result["top_buy"]         = result["buy_signals"][:top_n]
    result["top_sell"]        = result["sell_signals"][:top_n]
    return result

TOOL_REGISTRY: dict[str, Any] = {
    "get_live_quote": (
        get_live_quote,
        (
            "Fetch live NSE price for a single stock — returns last price, VWAP, day OHLC, "
            "% change, volume, traded value, 52w high/low with dates, circuit limits (upper/lower), "
            "sector P/E, stock P/E, and exact NSE last-update timestamp. "
            "First-class NSE real-time data — no lag. "
            "Use for: 'current price of X', 'live quote', 'what is X trading at', 'show me X price'."
        ),
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "get_nse_quotes": (
        get_nse_quotes,
        (
            "Fetch live NSE prices for MULTIPLE stocks in parallel — returns last price, VWAP, "
            "% change, volume, circuit limits, and sector P/E for all symbols at once. "
            "Much faster than calling get_live_quote() individually. "
            "Use for: 'prices of RELIANCE, TCS, INFY', 'check my watchlist prices', "
            "'how are these stocks doing: X Y Z', 'batch price check'."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {"type": "array", "items": {"type": "string"},
                            "description": "List of NSE tickers, up to 20"},
            },
            "required": ["symbols"],
        },
    ),
    "nse_search": (
        nse_search,
        (
            "Search NSE by company name or keyword to find the NSE symbol + live price. "
            "Use when user asks for 'Larsen and Toubro price', 'find Adani companies', "
            "'what is the NSE symbol for X', 'search for [company name]'. "
            "Returns symbol, company name, current price, % change, and sector."
        ),
        {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "Company name or keyword, e.g. 'Tata Steel'"},
                "top_n": {"type": "integer", "default": 5},
            },
            "required": ["query"],
        },
    ),
    "get_live_market_overview": (
        get_live_market_overview,
        "Fetch live NSE index levels (Nifty 50, Bank Nifty, IT, Midcap, Smallcap) plus advances/declines. Use for 'how is the market', 'market today', 'live market' queries.",
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_intraday_market_recap": (
        get_intraday_market_recap,
        (
            "Summarize what changed in the live market over the last N minutes. "
            "Compares current NSE live index tape with stored intraday snapshots when available; "
            "otherwise returns the current live tape with a snapshot-only note. "
            "Use for: 'what happened in the last 15 minutes', 'last 30 min market recap', "
            "'what changed just now'."
        ),
        {
            "type": "object",
            "properties": {"minutes": {"type": "integer", "default": 15}},
            "required": [],
        },
    ),
    "get_nse_intraday_snapshot": (
        get_nse_intraday_snapshot,
        (
            "Fetch the NSE website live quote/index snapshot before any yfinance intraday fallback. "
            "For stocks this calls NSE quote-equity; for indices such as NIFTY50/BANKNIFTY it calls "
            "NSE allIndices. If NSE blocks stock quotes, returns a degraded yfinance intraday-candle "
            "snapshot with fallback_reason. NSE website APIs provide live snapshot fields, not full "
            "candle history."
        ),
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "get_top_gainers_losers": (
        get_top_gainers_losers,
        (
            "Get top gaining and/or losing stocks from an NSE index RIGHT NOW (live). "
            "Returns symbol, LTP, % change, volume, 52w range for each stock. "
            "Use for 'top gainers intraday', 'top losers right now', 'biggest movers now'. "
            "For end-of-day / historical top movers use `get_eod_top_movers` instead. "
            "index can be: 'NIFTY 50', 'NIFTY BANK', 'NIFTY IT', 'NIFTY 500', "
            "'NIFTY MIDCAP 100', 'NIFTY SMALLCAP 100', 'NIFTY PHARMA', etc."
        ),
        {
            "type": "object",
            "properties": {
                "index":     {"type": "string", "default": "NIFTY 500"},
                "top_n":     {"type": "integer", "default": 10},
                "direction": {"type": "string", "enum": ["gainers", "losers", "both"], "default": "both"},
            },
            "required": [],
        },
    ),
    "get_eod_top_movers": (
        get_eod_top_movers,
        (
            "Get top EOD gainers/losers from the latest daily snapshot "
            "(`scores.stage_snapshots`). Use for 'top gainers EOD', "
            "'top movers end of day', 'yesterday's top gainers', or any "
            "top-movers request when the agent is in historical mode. "
            "For LIVE intraday top movers use `get_top_gainers_losers` instead. "
            "Returns symbol, price, change_1d_pct, stage, RS, RSI, signal, sector."
        ),
        {
            "type": "object",
            "properties": {
                "index":     {"type": "string", "default": "NIFTY 500"},
                "top_n":     {"type": "integer", "default": 10},
                "direction": {"type": "string", "enum": ["gainers", "losers", "both"], "default": "both"},
            },
            "required": [],
        },
    ),
    "get_most_active_stocks": (
        get_most_active_stocks,
        (
            "Get the most actively traded stocks by volume or traded value from an NSE index. "
            "Use for 'most active', 'highest volume', 'most traded', 'activity leaders'. "
            "'by' can be 'volume' or 'value'."
        ),
        {
            "type": "object",
            "properties": {
                "by":    {"type": "string", "enum": ["volume", "value"], "default": "value"},
                "index": {"type": "string", "default": "NIFTY 500"},
                "top_n": {"type": "integer", "default": 10},
            },
            "required": [],
        },
    ),
    "get_52week_extremes": (
        get_52week_extremes,
        (
            "Find stocks nearest to their 52-week high or low in a given index. "
            "Use for '52-week high', '52-week low', 'new highs', 'new lows', "
            "'near 52w high', 'breakout candidates', 'stocks at lows'. "
            "direction: 'high' or 'low'."
        ),
        {
            "type": "object",
            "properties": {
                "direction": {"type": "string", "enum": ["high", "low"], "default": "high"},
                "index":     {"type": "string", "default": "NIFTY 500"},
                "top_n":     {"type": "integer", "default": 15},
            },
            "required": [],
        },
    ),
    "get_fii_dii_activity": (
        get_fii_dii_activity,
        (
            "Fetch today's FII (Foreign Institutional Investors) and DII (Domestic Institutional) "
            "buy/sell activity from NSE. Returns net buy/sell values in crores and sentiment. "
            "Use for 'FII', 'DII', 'institutional activity', 'foreign investors', 'FII buying/selling'."
        ),
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_bulk_block_deals": (
        get_bulk_block_deals,
        (
            "Fetch today's bulk deals (> 0.5% of shares) and block deals from NSE. "
            "Shows which institutional clients bought or sold large stakes. "
            "Use for 'bulk deals', 'block deals', 'institutional trades', 'large trades', 'who is buying'."
        ),
        {
            "type": "object",
            "properties": {"top_n": {"type": "integer", "default": 20}},
            "required": [],
        },
    ),
    "compare_stocks": (
        compare_stocks,
        (
            "Compare multiple NSE stocks SIDE-BY-SIDE on both technical AND fundamental metrics. "
            "Technical: stage (Weinstein 1-4), RSI, Relative Strength %, technical score, "
            "investment score, trading signal (BUY/SELL/HOLD), trend, supertrend, 1d/1w/1m returns. "
            "Fundamental (from screener.in): P/E, P/B, ROE, ROCE, dividend yield, market cap, book value. "
            "Returns a comparison_table dict and full stock_details list with pros/cons/narrative. "
            "Use for: 'compare X vs Y', 'X vs Y vs Z', 'which is better: X or Y', "
            "'peer comparison', 'compare IT stocks', 'TCS vs INFY', 'rank by ROE'."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of NSE symbols, e.g. ['TCS', 'INFY', 'WIPRO']",
                },
                "aspects": {
                    "type": "array",
                    "items": {"type": "string", "enum": ["technical", "fundamental", "both"]},
                    "description": "What to compare. Default is both.",
                },
            },
            "required": ["symbols"],
        },
    ),
    "get_intraday_analysis": (
        get_intraday_analysis,
        (
            "Deep intraday technical analysis of a SINGLE NSE stock. "
            "Computes: MACD crossover, RSI reversal, Supertrend direction, Bollinger Band bounce, "
            "EMA 9/21 crossover, VCP pattern, Volume spike signals — each with entry price, "
            "target, stop-loss, and Risk:Reward ratio. Also returns key support/resistance levels "
            "(pivot points, swing highs/lows, EMAs), full indicator readings, and overall bias. "
            "Use for: 'intraday setup for X', 'should I buy X today', 'entry target SL for X', "
            "'technical signals X', 'trading strategy for X'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":     {"type": "string"},
                "interval":   {"type": "string", "enum": ["5m","15m","30m","1h"], "default": "15m"},
                "strategies": {
                    "type": "array",
                    "items": {"type": "string",
                              "enum": ["macd","rsi","supertrend","bollinger","ema","vcp","volume"]},
                    "description": "Strategies to run. Default is all.",
                },
            },
            "required": ["symbol"],
        },
    ),
    "scan_intraday_market": (
        scan_intraday_market,
        (
            "Scan ALL stocks in an NSE index for live intraday trading signals RIGHT NOW. "
            "Runs MACD, RSI, Supertrend, Bollinger, EMA crossover, VCP, Volume spike strategies "
            "across every constituent stock. Returns ranked BUY and SELL signals with entry/target/SL/R:R. "
            "Use for: 'intraday screener', 'scan Nifty 50 for buy signals', 'which BANK NIFTY stocks "
            "have buy signals', 'best intraday stocks today', 'top momentum plays', '/scan'."
        ),
        {
            "type": "object",
            "properties": {
                "index":            {"type": "string", "default": "NIFTY 50",
                                     "description": "NIFTY 50, NIFTY BANK, NIFTY IT, NIFTY 500, NIFTY PHARMA, etc."},
                "interval":         {"type": "string", "enum": ["5m","15m","30m","1h"], "default": "15m"},
                "strategies":       {"type": "array",
                                     "items": {"type": "string",
                                               "enum": ["macd","rsi","supertrend","bollinger","ema","vcp","volume"]}},
                "direction_filter": {"type": "string", "enum": ["buy","sell","all"], "default": "all"},
                "min_rr":           {"type": "number", "default": 1.5},
                "top_n":            {"type": "integer", "default": 5},
            },
            "required": [],
        },
    ),
    "scan_symbols_intraday": (
        scan_symbols_intraday,
        (
            "Scan a SPECIFIC LIST of NSE stocks for intraday signals — use when you already know which stocks to check. "
            "Works for ANY stock (not just index constituents). Ideal after EOD screening, breakout hunting, or "
            "when user asks 'check these stocks intraday', 'intraday setup for CHENNPETRO, NAM-INDIA, YATHARTH', "
            "'scan my watchlist intraday', 'which of these have buy signals'. "
            "Includes market-session awareness and EOD fallback for pre-market queries."
        ),
        {
            "type": "object",
            "properties": {
                "symbols":          {"type": "array", "items": {"type": "string"},
                                     "description": "List of NSE tickers, e.g. ['RELIANCE','CHENNPETRO']"},
                "interval":         {"type": "string", "enum": ["5m","15m","30m","1h"], "default": "15m"},
                "strategies":       {"type": "array",
                                     "items": {"type": "string",
                                               "enum": ["macd","rsi","supertrend","bollinger","ema","vcp","volume"]}},
                "direction_filter": {"type": "string", "enum": ["buy","sell","all"], "default": "all"},
                "min_rr":           {"type": "number", "default": 1.3},
                "top_n":            {"type": "integer", "default": 10},
            },
            "required": ["symbols"],
        },
    ),
    "resolve_symbol": (
        resolve_symbol,
        "Resolve a company name, alias, or near-match to its canonical NSE ticker symbol. Call this before stock-specific tools.",
        {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]},
    ),
    "resolve_stock_entity": (
        resolve_stock_entity,
        "Resolve a stock/company mention to a canonical NSE equity symbol with status, confidence, and mismatch-safe output.",
        {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]},
    ),
    "resolve_company_alias": (
        resolve_company_alias,
        "Resolve a company alias/name such as USL or United Spirits to the canonical NSE equity symbol.",
        {"type": "object", "properties": {"alias": {"type": "string"}}, "required": ["alias"]},
    ),
    "validate_requested_symbols": (
        validate_requested_symbols,
        "Validate explicit ticker-looking symbols requested by the user against symbols used in executed evidence, ignoring indicators like RSI/ADX/MA.",
        {
            "type": "object",
            "properties": {
                "query": {"type": "string"},
                "executed_symbols": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["query"],
        },
    ),
    "detect_non_symbol_terms": (
        detect_non_symbol_terms,
        "Detect market or technical terms such as RSI, ADX, MA, MACD that must not be routed as stock symbols.",
        {"type": "object", "properties": {"text": {"type": "string"}}, "required": ["text"]},
    ),
    "resolve_index_or_stock": (
        resolve_index_or_stock,
        "Resolve index/derivative underlyings such as NIFTY/BANKNIFTY first, otherwise resolve the query as an NSE stock.",
        {"type": "object", "properties": {"query": {"type": "string"}}, "required": ["query"]},
    ),
    "assess_user_situation": (
        assess_user_situation,
        "Assess what the user is asking, prior context needs, resolved entities, evidence plan, and clarification needs before routing.",
        {
            "type": "object",
            "properties": {
                "user_input": {"type": "string"},
                "data_mode": {"type": "string", "default": "historical"},
            },
            "required": ["user_input"],
        },
    ),
    "resolve_conversation_reference": (
        resolve_conversation_reference,
        "Resolve contextual references like 'the report' or 'these' against prior turn context when available.",
        {"type": "object", "properties": {"user_input": {"type": "string"}}, "required": ["user_input"]},
    ),
    "resolve_entity_context": (
        resolve_entity_context,
        "Resolve entity plus topic from commands or natural prompts such as 'search USL growth strategy'.",
        {"type": "object", "properties": {"user_input": {"type": "string"}}, "required": ["user_input"]},
    ),
    "validate_intent_evidence_plan": (
        validate_intent_evidence_plan,
        "Validate a planned evidence-tool list against required tools for an intent.",
        {
            "type": "object",
            "properties": {
                "intent": {"type": "string"},
                "evidence_plan": {"type": "array", "items": {"type": "string"}},
                "required_tools": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["intent"],
        },
    ),
    "request_clarification": (
        request_clarification,
        "Return a structured clarification request when situation assessment cannot safely choose tools.",
        {
            "type": "object",
            "properties": {
                "question": {"type": "string"},
                "reason": {"type": "string"},
            },
            "required": ["question"],
        },
    ),
    "get_symbol_snapshot": (
        get_symbol_snapshot,
        "Get the latest DB snapshot for a symbol: stage, RS, RSI, trading signal, sector, price",
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "discover_financial_filings": (
        discover_financial_filings,
        "Discover ranked latest-results filing candidates from NSE, BSE, and Screener for a symbol.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "max_results": {"type": "integer", "default": 10},
            },
            "required": ["symbol"],
        },
    ),
    "ingest_financial_filing": (
        ingest_financial_filing,
        "Download/register a direct financial filing URL using the deterministic filing registry.",
        {
            "type": "object",
            "properties": {
                "url": {"type": "string"},
                "symbol": {"type": "string"},
                "period": {"type": "string", "default": "latest"},
                "root_dir": {"type": "string"},
                "force": {"type": "boolean", "default": False},
            },
            "required": ["url"],
        },
    ),
    "parse_financial_filing": (
        parse_financial_filing,
        "Parse a registered financial filing manifest and return extracted evidence/facts when supported.",
        {"type": "object", "properties": {"manifest_path": {"type": "string"}}, "required": ["manifest_path"]},
    ),
    "parse_xbrl_filing": (
        parse_xbrl_filing,
        "Parse an XBRL filing path when XBRL support is wired; currently returns explicit unsupported status.",
        {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]},
    ),
    "parse_results_pdf_filing": (
        parse_results_pdf_filing,
        "Parse a local PDF results filing path using deterministic text/table extraction.",
        {"type": "object", "properties": {"path": {"type": "string"}}, "required": ["path"]},
    ),
    "reconcile_filing_facts": (
        reconcile_filing_facts,
        "Reconcile parsed filing evidence with Screener quarterly tables without inventing missing revenue/PAT/EPS facts.",
        {
            "type": "object",
            "properties": {
                "parsed_filing": {"type": "object"},
                "screener_data": {"type": "object"},
            },
            "required": [],
        },
    ),
    "get_latest_results": (
        get_latest_results,
        "High-level latest-results evidence pack for a symbol: discovery, ingestion, parsing, reconciliation, missing facts, and source trail.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "period": {"type": "string", "default": "latest"},
                "ingest": {"type": "boolean", "default": True},
            },
            "required": ["symbol"],
        },
    ),
    "summarize_latest_results": (
        summarize_latest_results,
        "Summarize a latest-results evidence pack without inventing missing revenue/PAT/EPS facts.",
        {"type": "object", "properties": {"results_pack": {"type": "object"}}, "required": ["results_pack"]},
    ),
    "build_evidence_matrix": (
        build_evidence_matrix,
        "Build a semantic evidence matrix from executed tool results by category.",
        {"type": "object", "properties": {"tool_results": {"type": "array", "items": {"type": "object"}}}, "required": ["tool_results"]},
    ),
    "validate_answer_against_evidence": (
        validate_answer_against_evidence,
        "Validate rendered answer text against available evidence categories and report unsupported claim categories.",
        {
            "type": "object",
            "properties": {
                "answer": {"type": "string"},
                "tool_results": {"type": "array", "items": {"type": "object"}},
            },
            "required": ["answer", "tool_results"],
        },
    ),
    "render_missing_evidence_block": (
        render_missing_evidence_block,
        "Render a standard missing-evidence block for blocked unsupported conclusions.",
        {
            "type": "object",
            "properties": {
                "intent": {"type": "string"},
                "missing_categories": {"type": "array", "items": {"type": "string"}},
                "missing_tools": {"type": "array", "items": {"type": "string"}},
            },
            "required": ["intent"],
        },
    ),
    "validate_required_tools_executed": (
        validate_required_tools_executed,
        "Validate that all required evidence tools executed for a planned intent.",
        {
            "type": "object",
            "properties": {
                "required_tools": {"type": "array", "items": {"type": "string"}},
                "tool_results": {"type": "array", "items": {"type": "object"}},
            },
            "required": ["required_tools", "tool_results"],
        },
    ),
    "get_fno_overview": (
        get_fno_overview,
        "Composite F&O overview: option-chain PCR/max-pain/top OI, futures basis/carry, and a gated options-strategy recommendation.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string", "default": "NIFTY"},
                "expiry_index": {"type": "integer", "default": 0},
            },
            "required": [],
        },
    ),
    "get_option_chain_summary": (
        get_option_chain_summary,
        "Summarize option-chain PCR, max pain, and top OI strikes for an index/equity.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry_index": {"type": "integer", "default": 0},
            },
            "required": ["symbol"],
        },
    ),
    "get_composite_max_pain": (
        get_composite_max_pain,
        "Return max-pain summary from the composite option-chain wrapper.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry_index": {"type": "integer", "default": 0},
            },
            "required": ["symbol"],
        },
    ),
    "get_pcr_summary": (
        get_pcr_summary,
        "Return put-call ratio and regime from the composite option-chain wrapper.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry_index": {"type": "integer", "default": 0},
            },
            "required": ["symbol"],
        },
    ),
    "get_top_oi_strikes": (
        get_top_oi_strikes,
        "Return top call and put OI strikes from the composite option-chain wrapper.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry_index": {"type": "integer", "default": 0},
            },
            "required": ["symbol"],
        },
    ),
    "get_composite_futures_basis": (
        get_composite_futures_basis,
        "Return futures basis from the composite F&O wrapper.",
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "get_composite_cost_of_carry": (
        get_composite_cost_of_carry,
        "Return cost-of-carry from the composite F&O wrapper.",
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "analyze_youtube_video": (
        analyze_youtube_video,
        (
            "Analyze a YouTube market video from metadata and available captions. "
            "Audio speech-to-text is attempted only when transcribe=true. Never returns the full transcript."
        ),
        {
            "type": "object",
            "properties": {
                "source": {"type": "string"},
                "persist": {"type": "boolean", "default": True},
                "max_segments": {"type": "integer", "default": 12},
                "transcribe": {"type": "boolean", "default": False},
                "transcription_backend": {"type": "string", "enum": ["local", "auto"], "default": "local"},
            },
            "required": ["source"],
        },
    ),
    "list_youtube_channels": (
        list_youtube_channels,
        "List preset YouTube market channels from data/youtube_channels.json.",
        {"type": "object", "properties": {}, "required": []},
    ),
    "analyze_youtube_channel_latest": (
        analyze_youtube_channel_latest,
        (
            "Select a preset YouTube market channel by number/name/id, fetch its latest video, "
            "then analyze metadata/captions. Speech-to-text is opt-in with transcribe=true."
        ),
        {
            "type": "object",
            "properties": {
                "selection": {"type": "string"},
                "persist": {"type": "boolean", "default": True},
                "max_segments": {"type": "integer", "default": 12},
                "transcribe": {"type": "boolean", "default": False},
                "transcription_backend": {"type": "string", "enum": ["local", "auto"], "default": "local"},
            },
            "required": ["selection"],
        },
    ),
    "recommend_options_strategy": (
        recommend_options_strategy,
        "Recommend a gated research-only options strategy from option-chain and futures evidence.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "option_chain": {"type": "object"},
                "futures": {"type": "object"},
                "raw_strategy": {"type": "object"},
            },
            "required": ["symbol"],
        },
    ),
    "audit_company_search": (
        audit_company_search,
        "Audit company evidence search attempts with source group, query, result count, parse status, and gaps.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "alias": {"type": "string"},
                "include_external": {"type": "boolean", "default": False},
            },
            "required": ["symbol"],
        },
    ),
    "search_company_official_sources": (
        search_company_official_sources,
        "Search or report official company source attempts before external sources.",
        {
            "type": "object",
            "properties": {"symbol": {"type": "string"}, "alias": {"type": "string"}},
            "required": ["symbol"],
        },
    ),
    "search_company_filings": (
        audit_search_company_filings,
        "Search or report company filing-source attempts with auditable no-result gaps.",
        {
            "type": "object",
            "properties": {"symbol": {"type": "string"}, "alias": {"type": "string"}},
            "required": ["symbol"],
        },
    ),
    "promote_company_evidence_to_postgres": (
        promote_company_evidence_to_postgres,
        "Prepare/dry-run promotion of company evidence records into PostgreSQL with source URL/path metadata.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "evidence": {"type": "array", "items": {"type": "object"}},
                "dsn": {"type": "string"},
                "dry_run": {"type": "boolean", "default": True},
            },
            "required": ["symbol"],
        },
    ),
    "get_company_evidence_coverage": (
        get_company_evidence_coverage,
        "Report company evidence coverage counts and source gaps by category.",
        {
            "type": "object",
            "properties": {"symbol": {"type": "string"}, "alias": {"type": "string"}},
            "required": ["symbol"],
        },
    ),
    "get_technical_setup": (
        get_technical_setup,
        "Compute technical indicators (RSI, ADX, MACD, supertrend, MAs, 52w position) from price history",
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "get_sector_context": (
        get_sector_context,
        "Get sector breadth, top stocks, and performance. Pass a stock SYMBOL (e.g. 'BHEL') to auto-detect its sector, or a sector name like 'Pharma'",
        {"type": "object", "properties": {"sector_or_symbol": {"type": "string"}}, "required": ["sector_or_symbol"]},
    ),
    "run_screener_query": (
        run_screener_query,
        (
            "Run an EOD screener from DB snapshot data. "
            "Original: stage2, breakouts, supertrend_buy, strong_buy, new_entrants. "
            "New: momentum_52w (near-52W-high leaders), high_rs (top RS ≥ 1.15), "
            "turnaround (recovery setups), stage1_base (basing/coiling), "
            "tight_range (VCP-like consolidation), oversold_bounce (RSI < 40 in Stage 2 dip)."
        ),
        {
            "type": "object",
            "properties": {
                "screen_type": {
                    "type": "string",
                    "enum": [
                        "stage2", "breakouts", "supertrend_buy", "strong_buy", "new_entrants",
                        "momentum_52w", "high_rs", "turnaround", "stage1_base",
                        "tight_range", "oversold_bounce",
                    ],
                },
                "top_n": {"type": "integer", "default": 10},
            },
            "required": ["screen_type"],
        },
    ),
    "get_long_term_growth_candidates": (
        get_long_term_growth_candidates,
        (
            "Rank live NSE index constituents for long-term growth research using PostgreSQL "
            "technical/fundamental scores and optional screener.in enrichment. Use for prompts "
            "asking for long-term growth-potential stocks in midcap/smallcap/index universes."
        ),
        {
            "type": "object",
            "properties": {
                "index_scope": {"type": "string", "default": "MIDCAP"},
                "top_n": {"type": "integer", "default": 12},
                "include_research": {"type": "boolean", "default": True},
            },
            "required": [],
        },
    ),
    "validate_strength_watchlist": (
        validate_strength_watchlist,
        (
            "Data-direct validator for ranking a list of NSE stocks by CANSLIM score, "
            "relative strength, enhanced fundamental score, Piotroski F-score, Beneish, "
            "and Altman. Never assumes missing evidence: missing CANSLIM, RS, fundamental, "
            "or forensic data is returned in missing_evidence and evidence_coverage."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "NSE symbols to validate and rank",
                },
                "top_n": {"type": "integer", "default": 20},
            },
            "required": ["symbols"],
        },
    ),
    "get_index_snapshot": (
        get_index_snapshot,
        "Get the latest OHLCV and 10-day trend for a Nifty index",
        {"type": "object", "properties": {"index_name": {"type": "string"}}, "required": ["index_name"]},
    ),
    "get_market_breadth": (
        get_market_breadth,
        "Get overall NSE market breadth: advance/decline, RS distribution, stage breakdown",
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_global_market_assessment": (
        get_global_market_assessment,
        (
            "Assess global market cues for Indian markets from cached global index, FX, "
            "commodity, and correlation data. Returns risk regime, US/Asia/commodity/FX "
            "bias, India sector read-through, watch items, and Nifty correlation context. "
            "Use for global market assessment, overnight cues, US/Asia markets, crude, "
            "DXY, USDINR, and India read-through questions."
        ),
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_intraday_source_health": (
        get_intraday_source_health,
        (
            "Check PostgreSQL intraday source health. Reports whether intraday.quote_snapshots, "
            "intraday.ohlcv_bars, and related tables exist, row counts, latest timestamps, freshness, and "
            "overall intraday status. Use before intraday calculations."
        ),
        {
            "type": "object",
            "properties": {
                "max_age_minutes": {"type": "integer", "default": 30},
            },
            "required": [],
        },
    ),
    "get_intraday_bars": (
        get_intraday_bars,
        (
            "Read OHLCV bars from PostgreSQL intraday.ohlcv_bars for one symbol and timeframe. "
            "When PG has no bars, seed PG from yfinance candles and return source metadata."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "timeframe": {"type": "string", "enum": ["5m", "15m", "30m", "1h"], "default": "15m"},
                "lookback": {"type": "integer", "default": 120},
            },
            "required": ["symbol"],
        },
    ),
    "get_intraday_levels": (
        get_intraday_levels,
        (
            "Compute support, resistance, pivot, EMA, and latest close levels from "
            "PostgreSQL intraday.ohlcv_bars. Technical levels only; not a trade recommendation."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "timeframe": {"type": "string", "enum": ["5m", "15m", "30m", "1h"], "default": "15m"},
            },
            "required": ["symbol"],
        },
    ),
    "compute_intraday_indicators": (
        compute_intraday_indicators,
        (
            "Compute RSI, MACD, Supertrend, EMA, Bollinger, ATR, volume ratio, and "
            "a research setup score from PostgreSQL intraday.ohlcv_bars."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "timeframe": {"type": "string", "enum": ["5m", "15m", "30m", "1h"], "default": "15m"},
            },
            "required": ["symbol"],
        },
    ),
    "explain_intraday_setup": (
        explain_intraday_setup,
        (
            "Explain a symbol's PostgreSQL-backed intraday setup using analyzer-style evidence, "
            "research-only labels LONG_SETUP/SHORT_SETUP/WATCH/AVOID/SETUP_INVALIDATED, "
            "technical target zones, and setup invalidation levels."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "timeframe": {"type": "string", "enum": ["5m", "15m", "30m", "1h"], "default": "15m"},
            },
            "required": ["symbol"],
        },
    ),
    "run_intraday_screener": (
        run_intraday_screener,
        (
            "Run an intraday screener (PostgreSQL-backed or live yfinance fallback). "
            "Original types: momentum, breakouts, vcp, supertrend, levels, all. "
            "New types: opening_range_breakout (ORB — first 15-30min range break + volume), "
            "gap_and_go (gap continuation with MACD + volume), "
            "macd_crossover (fresh MACD signal line cross only), "
            "rsi_divergence (RSI extreme + Bollinger mean-reversion), "
            "bb_squeeze (Bollinger Band squeeze breakout), "
            "vwap_reclaim (short-EMA VWAP proxy reclaim or loss). "
            "Returns setup labels, support/resistance, target zones, confidence."
        ),
        {
            "type": "object",
            "properties": {
                "screen_type": {
                    "type": "string",
                    "enum": [
                        "momentum", "breakouts", "vcp", "supertrend", "levels", "all",
                        "opening_range_breakout", "gap_and_go", "macd_crossover",
                        "rsi_divergence", "bb_squeeze", "vwap_reclaim",
                    ],
                    "default": "momentum",
                },
                "timeframe": {"type": "string", "enum": ["5m", "15m", "30m", "1h"], "default": "15m"},
                "min_score": {"type": "number", "default": 55},
                "top_n": {"type": "integer", "default": 10},
                "symbols": {"type": "array", "items": {"type": "string"}},
            },
            "required": [],
        },
    ),
    "get_data_health": (
        get_data_health,
        "Check freshness of all local data sources (CSV files, DB snapshots)",
        {"type": "object", "properties": {}, "required": []},
    ),
    "find_latest_report": (
        find_latest_report,
        "List available generated reports (HTML/CSV) by type keyword",
        {"type": "object", "properties": {"report_type": {"type": "string"}}, "required": []},
    ),
    "search_latest_catalysts": (
        search_latest_catalysts,
        "Search for recent news and catalysts for a stock symbol via web search. "
        "Fetches actual article content from top 3 results — read the 'article_text' "
        "field to provide detailed summary and opinion.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "max_results": {"type": "integer", "default": 5},
            },
            "required": ["symbol"],
        },
    ),
    "search_market_knowledge": (
        search_market_knowledge,
        (
            "Answer financial-market education questions using source-backed search on "
            "Wikipedia and Investopedia. Use for definitions, explainers, and concept "
            "comparisons such as 'what is PE', 'explain Minervini strategy', "
            "'ROCE vs ROE', 'what is EBITDA', 'how does RSI work'. Never answer these "
            "from memory first; return source URLs and say when reliable sources are not found."
        ),
        {
            "type": "object",
            "properties": {
                "query": {"type": "string", "description": "The market concept question to answer"},
                "sources": {
                    "type": "array",
                    "items": {"type": "string", "enum": ["wikipedia", "investopedia"]},
                    "default": ["wikipedia", "investopedia"],
                },
            },
            "required": ["query"],
        },
    ),
    "scrape_screener_in": (
        scrape_screener_in,
        (
            "Scrape screener.in for deep fundamental data: key ratios (P/E, P/B, ROE, ROCE, "
            "market-cap, dividend yield), Screener's pros/cons analysis, last 6 quarters of "
            "financials, 5-year annual P&L, peer comparison table, shareholding pattern, "
            "BSE corporate-announcement PDF links, annual-report PDF links. "
            "Use for 'fundamentals', 'ratios', 'financial statements', 'peers', 'valuation' queries."
        ),
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "search_yahoo_finance": (
        search_yahoo_finance,
        (
            "Fetch Yahoo Finance data for an NSE stock: current price stats (52-week range, "
            "day range, prev close) and up to 6 recent news articles with real URLs. "
            "Use for 'yahoo finance', 'YF news', 'latest news', or as a supplementary "
            "news source alongside search_latest_catalysts."
        ),
        {"type": "object", "properties": {"symbol": {"type": "string"}}, "required": ["symbol"]},
    ),
    "multi_source_web_search": (
        multi_source_web_search,
        (
            "Search multiple finance websites (moneycontrol.com, screener.in, "
            "economictimes.indiatimes.com, nseindia.com, bseindia.com) for a stock. "
            "Returns up to 4 real results per site with decoded URLs. "
            "Use for 'moneycontrol', 'concall', 'transcript', 'BSE filings', 'announcements', "
            "or when comprehensive multi-source research is needed."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":       {"type": "string"},
                "company_name": {"type": "string", "description": "Full company name for better search"},
                "extra_query":  {"type": "string", "description": "Additional search terms e.g. 'concall Q4 2026'"},
            },
            "required": ["symbol"],
        },
    ),
    "comprehensive_stock_research": (
        comprehensive_stock_research,
        (
            "One-call deep research across ALL sources: screener.in fundamentals + Yahoo Finance "
            "stats + multi-site news search (moneycontrol, ET, NSE, BSE). Returns ratios, "
            "pros/cons, quarterly results, peer table, shareholding, filings, news, and direct "
            "deep-links to screener.in, NSE, BSE, Yahoo Finance pages. "
            "Use for broad 'research', 'full analysis', or when user asks for comprehensive info "
            "about a stock from multiple sources."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":  {"type": "string"},
                "aspects": {
                    "type": "array",
                    "items": {"type": "string",
                              "enum": ["fundamentals","news","concalls","peers","filings","ratios","all"]},
                    "description": "Which aspects to fetch. Defaults to all.",
                },
            },
            "required": ["symbol"],
        },
    ),

    # ── Deep Search Engine ─────────────────────────────────────────────────────

    "search_nse_announcements": (
        search_nse_announcements,
        (
            "Fetch live NSE corporate announcements for a stock (board meetings, results filings, "
            "regulatory disclosures, pledging changes). Source: nseindia.com live API. "
            "Use when user asks about 'NSE announcements', 'latest filings', 'company disclosures'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 15},
            },
            "required": ["symbol"],
        },
    ),

    "search_corporate_actions": (
        search_corporate_actions,
        (
            "Fetch upcoming and recent corporate actions from NSE: dividends, stock splits, "
            "bonus issues, rights issues, AGMs. Source: nseindia.com corporate actions API. "
            "Use when user asks about 'dividend', 'bonus', 'split', 'rights issue', 'ex-date'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 12},
            },
            "required": ["symbol"],
        },
    ),

    "search_insider_trades": (
        search_insider_trades,
        (
            "Fetch promoter / director / insider trading disclosures from NSE's SAST/PIT database. "
            "Shows who is buying or selling, quantities, values, and provides SEBI XBRL links. "
            "Use when user asks about 'insider trading', 'promoter buying/selling', 'insider activity'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 15},
            },
            "required": ["symbol"],
        },
    ),

    "search_bse_filings": (
        search_bse_filings,
        (
            "Search BSE India corporate filings: board meeting results, annual reports, "
            "concall notices, investor presentations. Uses DuckDuckGo site:bseindia.com. "
            "Use when user asks about 'BSE filings', 'annual report', 'board meeting results'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 10},
            },
            "required": ["symbol"],
        },
    ),

    "search_shareholding_analysis": (
        search_shareholding_analysis,
        (
            "Scrape quarterly shareholding pattern from screener.in: promoter %, FII %, DII %, "
            "public %, pledge alerts, and QoQ change trends. "
            "Use when user asks about 'shareholding', 'promoter holding', 'FII holding', "
            "'pledge', 'DII activity'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
            },
            "required": ["symbol"],
        },
    ),

    "search_analyst_coverage": (
        search_analyst_coverage,
        (
            "Aggregate analyst price targets, buy/sell/hold ratings, and brokerage recommendations "
            "from Moneycontrol, Economic Times, screener.in and general web. "
            "Use when user asks about 'analyst target', 'buy/sell/hold rating', 'brokerage view', "
            "'consensus target price'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 8},
            },
            "required": ["symbol"],
        },
    ),

    "search_concall_transcripts": (
        search_concall_transcripts,
        (
            "Search for earnings call transcripts, investor day presentations, and management "
            "commentary from screener.in, trendlyne, BSE, Moneycontrol, ET. "
            "Use when user asks about 'concall', 'earnings call', 'management commentary', "
            "'Q4 concall', 'investor day transcript'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 8},
            },
            "required": ["symbol"],
        },
    ),

    "search_sector_news": (
        search_sector_news,
        (
            "Aggregate sector-level news from 6 portals in parallel: Economic Times, "
            "Business Standard, Mint, Moneycontrol, Financial Express, Hindu BusinessLine. "
            "Use when user asks about 'sector news', 'industry news', 'latest news', "
            "or for macro sector context around a stock."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "sector":      {"type": "string", "description": "Optional sector override"},
                "max_results": {"type": "integer", "default": 5},
            },
            "required": ["symbol"],
        },
    ),

    "search_social_buzz": (
        search_social_buzz,
        (
            "Gauge retail investor sentiment from Indian investing communities: "
            "Reddit r/IndiaInvestments, r/IndianStockMarket, Valuepickr, Traderji, Tijori. "
            "Returns community discussions plus a rough bullish/bearish sentiment signal. "
            "Use when user asks about 'social sentiment', 'retail view', 'community buzz', "
            "'what are investors saying about X'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":      {"type": "string"},
                "max_results": {"type": "integer", "default": 5},
            },
            "required": ["symbol"],
        },
    ),

    "deep_search": (
        deep_search,
        (
            "Run a full deep-dive search on any NSE stock using up to 9 distinct parallel "
            "search verticals: NSE announcements, corporate actions, insider trades, BSE filings, "
            "shareholding analysis, analyst coverage, concall transcripts, sector news, social buzz. "
            "Intelligently selects verticals based on context (e.g. 'results' → concalls+announcements, "
            "'dividend' → corporate_actions, 'insider' → insider_trades+shareholding). "
            "Use for 'deep search', 'deep dive', 'full search', or when multiple search verticals needed."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":    {"type": "string"},
                "verticals": {
                    "type": "array",
                    "items": {
                        "type": "string",
                        "enum": [
                            "announcements", "corporate_actions", "insider_trades",
                            "bse_filings", "shareholding", "analyst_coverage",
                            "concalls", "sector_news", "social_buzz",
                        ],
                    },
                    "description": "Specific verticals to run. Leave empty to auto-select from context.",
                },
                "context": {
                    "type": "string",
                    "description": "Optional context hint to bias vertical selection (e.g. 'dividend ex-date', 'concall highlights', 'insider buying').",
                },
            },
            "required": ["symbol"],
        },
    ),
    "get_portfolio_exposure": (
        get_portfolio_exposure,
        "Return portfolio holdings summary (total stocks, sector distribution, top holdings). Optionally filter by sector.",
        {
            "type": "object",
            "properties": {
                "sector": {"type": "string", "description": "Optional sector name to filter holdings"},
            },
            "required": [],
        },
    ),
    "find_portfolio_overlap": (
        find_portfolio_overlap,
        "Find portfolio holdings that also appear in a screener (stage2, supertrend_buy, all)",
        {
            "type": "object",
            "properties": {
                "screener": {
                    "type": "string",
                    "enum": ["stage2", "supertrend_buy", "all"],
                    "default": "stage2",
                },
            },
            "required": [],
        },
    ),
}


# ─────────────────────────────────────────────────────────────────────────────
# F&O / Options Tool Wrappers
# ─────────────────────────────────────────────────────────────────────────────

def _postgres_options_chain_summary(symbol: str, expiry_index: int = 0) -> dict:
    """Build the rich /options response from PostgreSQL EOD option-chain data."""
    try:
        from terminal.fno_data import get_eod_option_chain
        # Pick expiry by index from PostgreSQL-backed EOD data.
        all_chain = get_eod_option_chain(symbol)
        if all_chain.empty:
            return {"error": f"No PostgreSQL EOD options data for {symbol}", "source": "postgres-eod"}
        expiries = sorted(all_chain["expiry_date"].dropna().astype(str).unique().tolist())
        expiry = expiries[min(expiry_index, len(expiries) - 1)] if expiries else None
        df = get_eod_option_chain(symbol, expiry_date=expiry)
        if df.empty:
            return {"error": f"No PostgreSQL EOD options data for {symbol}", "source": "postgres-eod"}
        trade_dates = df["trade_date"].dropna().astype(str).unique().tolist() if "trade_date" in df.columns else []

        calls: list[dict] = []
        puts: list[dict] = []
        for _, row in df.iterrows():
            item = {
                "strike": float(row.get("strike") or 0),
                "oi": int(row.get("oi") or 0),
                "chg_oi": int(row.get("oi_change") or 0),
                "volume": int(row.get("volume") or 0),
                "iv": None,
                "ltp": float(row.get("last_price") or row.get("close") or 0),
                "bid": None,
                "ask": None,
            }
            if row.get("option_type") == "CE":
                calls.append(item)
            elif row.get("option_type") == "PE":
                puts.append(item)

        calls.sort(key=lambda x: x["strike"])
        puts.sort(key=lambda x: x["strike"])
        total_call_oi = sum(c["oi"] for c in calls)
        total_put_oi = sum(p["oi"] for p in puts)
        pcr = round(total_put_oi / total_call_oi, 3) if total_call_oi else 0
        strikes = sorted(set(c["strike"] for c in calls) | set(p["strike"] for p in puts))
        call_map = {c["strike"]: c["oi"] for c in calls}
        put_map = {p["strike"]: p["oi"] for p in puts}
        max_pain_strike = 0
        min_pain = float("inf")
        for s in strikes:
            loss = sum(call_map.get(k, 0) * max(0, k - s) for k in strikes)
            loss += sum(put_map.get(k, 0) * max(0, s - k) for k in strikes)
            if loss < min_pain:
                min_pain = loss
                max_pain_strike = s
        underlying_vals = df["underlying"].dropna().tolist() if "underlying" in df.columns else []
        underlying = float(underlying_vals[0]) if underlying_vals else 0
        atm = min(strikes, key=lambda s: abs(s - underlying)) if strikes and underlying else 0
        return {
            "symbol": symbol.upper(),
            "expiry": expiry,
            "expiry_dates": expiries[:4],
            "underlying": underlying,
            "atm": atm,
            "pcr": pcr,
            "max_pain": max_pain_strike,
            "calls": calls,
            "puts": puts,
            "total_call_oi": total_call_oi,
            "total_put_oi": total_put_oi,
            "source": "postgres-eod",
            "as_of": trade_dates[0] if trade_dates else None,
        }
    except Exception as exc:
        return {"error": f"PostgreSQL options fallback failed: {exc}", "source": "postgres-eod"}


def get_options_chain(symbol: str = "NIFTY", expiry_index: int = 0) -> dict:
    """Fetch NSE options chain for a symbol with PCR, max pain, and IV data."""
    symbol = symbol.upper()

    try:
        data = None
        try:
            from nsepython import nse_optionchain_scrapper
            data = nse_optionchain_scrapper(symbol)
        except Exception:
            data = None

        if data is None:
            import requests
            session = requests.Session()
            session.headers.update({
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36",
                "Accept-Encoding": "gzip, deflate, br",
                "Accept-Language": "en-US,en;q=0.9",
                "Referer": "https://www.nseindia.com/option-chain",
            })
            session.get("https://www.nseindia.com", timeout=10)
            import time as _t; _t.sleep(1)
            if symbol in ("NIFTY", "BANKNIFTY", "FINNIFTY", "MIDCPNIFTY"):
                url = f"https://www.nseindia.com/api/option-chain-indices?symbol={symbol}"
            else:
                url = f"https://www.nseindia.com/api/option-chain-equities?symbol={symbol}"
            resp = session.get(url, timeout=15)
            data = resp.json()

        if not data or "records" not in data:
            return _postgres_options_chain_summary(symbol, expiry_index)

        records = data["records"]
        expiry_dates = records.get("expiryDates", [])
        if not expiry_dates:
            return _postgres_options_chain_summary(symbol, expiry_index)

        expiry = expiry_dates[min(expiry_index, len(expiry_dates) - 1)]
        underlying = records.get("underlyingValue", 0)

        calls: list[dict] = []
        puts: list[dict] = []
        for row in records.get("data", []):
            if row.get("expiryDate") != expiry:
                continue
            strike = row.get("strikePrice", 0)
            ce = row.get("CE", {})
            pe = row.get("PE", {})
            if ce:
                calls.append({
                    "strike": strike,
                    "oi": ce.get("openInterest", 0),
                    "chg_oi": ce.get("changeinOpenInterest", 0),
                    "volume": ce.get("totalTradedVolume", 0),
                    "iv": ce.get("impliedVolatility", 0),
                    "ltp": ce.get("lastPrice", 0),
                    "bid": ce.get("bidprice", 0),
                    "ask": ce.get("askPrice", 0),
                })
            if pe:
                puts.append({
                    "strike": strike,
                    "oi": pe.get("openInterest", 0),
                    "chg_oi": pe.get("changeinOpenInterest", 0),
                    "volume": pe.get("totalTradedVolume", 0),
                    "iv": pe.get("impliedVolatility", 0),
                    "ltp": pe.get("lastPrice", 0),
                    "bid": pe.get("bidprice", 0),
                    "ask": pe.get("askPrice", 0),
                })

        calls.sort(key=lambda x: x["strike"])
        puts.sort(key=lambda x: x["strike"])

        total_call_oi = sum(c["oi"] for c in calls)
        total_put_oi = sum(p["oi"] for p in puts)
        pcr = round(total_put_oi / total_call_oi, 3) if total_call_oi else 0

        strikes = sorted(set(c["strike"] for c in calls) | set(p["strike"] for p in puts))
        call_map = {c["strike"]: c["oi"] for c in calls}
        put_map = {p["strike"]: p["oi"] for p in puts}
        max_pain_strike = 0
        min_pain = float("inf")
        for s in strikes:
            loss = sum(call_map.get(k, 0) * max(0, k - s) for k in strikes)
            loss += sum(put_map.get(k, 0) * max(0, s - k) for k in strikes)
            if loss < min_pain:
                min_pain = loss
                max_pain_strike = s

        atm = min(strikes, key=lambda s: abs(s - underlying)) if strikes else 0

        return {
            "symbol": symbol,
            "expiry": expiry,
            "expiry_dates": expiry_dates[:4],
            "underlying": underlying,
            "atm": atm,
            "pcr": pcr,
            "max_pain": max_pain_strike,
            "calls": calls,
            "puts": puts,
            "total_call_oi": total_call_oi,
            "total_put_oi": total_put_oi,
        }
    except Exception:
        return _postgres_options_chain_summary(symbol, expiry_index)


def get_option_chain(symbol: str, expiry: str | None = None,
                     use_live: bool = True) -> dict:
    """
    Fetch live option chain for an index or stock.
    Returns OI, IV, PCR, max pain, ATM greeks, OI buildup and support/resistance levels.
    Falls back to EOD data outside market hours.
    """
    return analyze_option_chain(symbol, expiry, use_live=use_live)


def get_oi_analysis(symbol: str, expiry: str | None = None) -> dict:
    """
    Focused open-interest analysis: PCR, max pain, CE/PE OI concentration,
    OI buildup and unwinding at key strikes.  Good for support/resistance.
    """
    result = analyze_option_chain(symbol, expiry, use_live=True)
    if "error" in result:
        return result
    return {
        "symbol":            result["symbol"],
        "underlying":        result["underlying"],
        "expiry":            result["expiry"],
        "dte":               result["dte"],
        "source":            result.get("source"),
        "pcr":               result["pcr"],
        "max_pain":          result["max_pain"],
        "max_pain_vs_spot":  result.get("max_pain_vs_spot"),
        "top_ce_oi_strikes": result["top_ce_oi_strikes"],   # resistance
        "top_pe_oi_strikes": result["top_pe_oi_strikes"],   # support
        "oi_buildup":        result["oi_buildup"],
        "total_ce_oi":       result.get("total_ce_oi"),
        "total_pe_oi":       result.get("total_pe_oi"),
    }


def get_futures_analysis(symbol: str) -> dict:
    """
    Futures basis, cost-of-carry, and rollover analysis for an index or stock.
    """
    return analyze_futures(symbol, use_live=True)


def get_options_strategy(symbol: str, strategy: str,
                          expiry: str | None = None) -> dict:
    """
    Build a specific options strategy with live pricing.
    Returns legs, entry costs, max risk/reward, breakevens, and payoff curve.

    Available strategies: long_call, long_put, bull_call_spread, bear_put_spread,
    long_straddle, long_strangle, iron_condor, covered_call, protective_put,
    calendar_spread.
    """
    return build_strategy(symbol, strategy, expiry, use_live=True)


def get_strategy_recommendations(symbol: str, expiry: str | None = None) -> dict:
    """
    Analyse current option chain context (PCR, IV, DTE, max pain) and
    recommend the top 3 options strategies with rationale.
    """
    return recommend_strategies(symbol, expiry, use_live=True)


def refresh_fno_eod_data() -> dict:
    """
    Download the latest F&O EOD bhavcopy from NSE archives and store in PostgreSQL.
    Returns summary: trade_date, rows_stored, options count, futures count.
    """
    return _fno_load_latest()


def get_fno_data_status() -> dict:
    """
    Return the status of the local F&O EOD database:
    available dates, record counts, and last download time.
    """
    dates = _fno_available_dates()
    if not dates:
        return {
            "status": "no_data",
            "message": "F&O EOD database is empty. Run refresh_fno_eod_data() first.",
            "available_strategies": list(STRATEGY_CATALOG.keys()),
        }
    return {
        "status":             "ok",
        "latest_date":        dates[0],
        "available_dates":    dates[:10],
        "total_dates":        len(dates),
        "available_strategies": list(STRATEGY_CATALOG.keys()),
    }


def analyze_options_buying(symbol: str, direction: str = "bullish",
                            expiry: str | None = None) -> dict:
    """
    Deep options buying analysis for a symbol.
    Returns: ATM IV + regime, IV rank (estimated), expected move, strike selection guide
    (ITM/ATM/OTM with delta/theta/breakeven/probability), theta decay profile,
    OI support/resistance context, buying verdict (BUY / SPREAD / AVOID).
    direction: 'bullish' | 'bearish' | 'volatile'
    """
    return analyze_buying_opportunity(symbol, direction, expiry, use_live=True)


def scan_options_buys(direction: str = "bullish",
                       max_iv: float = 25.0,
                       min_oi: int = 500_000,
                       top_n: int = 10) -> dict:
    """
    Scan all F&O-eligible stocks and indices for the best options buying opportunities.
    Ranks by: low ATM IV (cheap options) + adequate OI liquidity + ideal DTE.
    Returns top N symbols with ATM IV, straddle cost, expected move, and buying score.
    direction: 'bullish' | 'bearish' | 'volatile'
    max_iv: IV ceiling (default 25%) — lower = cheaper options only
    """
    return scan_options_buying_opportunities(direction, min_oi, max_iv, top_n)


def get_fno_analytics(symbol: str | None = None, top_n: int = 20) -> dict:
    """Return PostgreSQL-backed F&O analytics: PCR, max pain, buildup, signal."""
    try:
        import psycopg2
        conn = psycopg2.connect("dbname=nse_market user=nse_admin host=/tmp")
        try:
            params: list[Any] = []
            where = ""
            if symbol:
                where = "WHERE symbol = %s"
                params.append(symbol.upper())
            else:
                params.append(int(top_n))
            sql = f"""
                SELECT
                    trade_date::text AS trade_date,
                    symbol,
                    options_expiry::text AS options_expiry,
                    futures_expiry::text AS futures_expiry,
                    underlying_price,
                    futures_close,
                    futures_price_change_pct,
                    futures_oi_change_pct,
                    pcr_oi,
                    pcr_volume,
                    max_call_oi_strike,
                    max_put_oi_strike,
                    max_pain,
                    distance_from_max_pain_pct,
                    buildup,
                    fno_signal
                FROM derivatives.mv_fno_symbol_analytics
                {where}
                ORDER BY
                    CASE fno_signal
                        WHEN 'BULL' THEN 1
                        WHEN 'BEAR' THEN 2
                        WHEN 'MILD_BULL' THEN 3
                        WHEN 'MILD_BEAR' THEN 4
                        ELSE 5
                    END,
                    abs(distance_from_max_pain_pct) DESC NULLS LAST,
                    symbol
                """ + ("" if symbol else "LIMIT %s")
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", UserWarning)
                df = pd.read_sql_query(sql, conn, params=tuple(params))
        finally:
            conn.close()
    except Exception as exc:
        return {"error": f"PostgreSQL F&O analytics unavailable: {exc}", "symbol": symbol}

    if df.empty:
        return {"status": "no_data", "symbol": symbol, "rows": []}
    return {
        "status": "ok",
        "source": "postgres.derivatives.mv_fno_symbol_analytics",
        "symbol": symbol.upper() if symbol else None,
        "count": int(len(df)),
        "rows": df.where(pd.notna(df), None).to_dict("records"),
    }


def run_option_payoff_scenario(
    symbol: str,
    option_type: str = "CE",
    strike: float | None = None,
    expiry_date: str | None = None,
    entry_premium: float | None = None,
    lots: int = 1,
    move_start: float = -5.0,
    move_end: float = 5.0,
    move_step: float = 2.5,
) -> dict:
    """Run PostgreSQL option payoff what-if scenarios."""
    try:
        import psycopg2
        conn = psycopg2.connect("dbname=nse_market user=nse_admin host=/tmp")
        try:
            with warnings.catch_warnings():
                warnings.simplefilter("ignore", UserWarning)
                df = pd.read_sql_query(
                    """
                    SELECT *
                    FROM derivatives.option_payoff(%s, %s, %s, %s, %s, %s, %s, %s, %s, false, NULL)
                    """,
                    conn,
                    params=(
                        symbol.upper(),
                        option_type.upper(),
                        strike,
                        expiry_date,
                        entry_premium,
                        lots,
                        move_start,
                        move_end,
                        move_step,
                    ),
                )
        finally:
            conn.close()
    except Exception as exc:
        return {"error": f"Option payoff scenario failed: {exc}", "symbol": symbol}

    return {
        "status": "ok",
        "source": "postgres.derivatives.option_payoff",
        "symbol": symbol.upper(),
        "option_type": option_type.upper(),
        "count": int(len(df)),
        "rows": df.where(pd.notna(df), None).to_dict("records"),
    }


# Register F&O tools
TOOL_REGISTRY.update({
    "get_option_chain": (
        get_option_chain,
        (
            "Fetch live option chain for a NIFTY/BANKNIFTY index or any NSE stock. "
            "Returns: underlying price, PCR (put-call ratio), max pain, top CE/PE OI strikes "
            "(key resistance/support), ATM±2 greeks (delta/theta/vega/gamma/IV), OI buildup "
            "and unwinding, IV skew, DTE. Falls back to EOD data outside market hours. "
            "Use for: 'option chain', 'OI analysis', 'PCR', 'max pain', 'support resistance', "
            "'greeks', 'IV', 'option data for NIFTY/BANKNIFTY/<stock>'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":   {"type": "string", "description": "NSE symbol e.g. NIFTY, BANKNIFTY, RELIANCE"},
                "expiry":   {"type": "string", "description": "Expiry date YYYY-MM-DD or NSE format e.g. 08-May-2026"},
                "use_live": {"type": "boolean", "default": True},
            },
            "required": ["symbol"],
        },
    ),
    "get_oi_analysis": (
        get_oi_analysis,
        (
            "Focused open interest analysis: PCR, max pain, top CE/PE OI concentration strikes, "
            "OI buildup and unwinding. Identifies key support (PE OI) and resistance (CE OI) levels. "
            "Use for: 'where is support/resistance', 'OI buildup', 'PCR signal', 'max pain level', "
            "'where is call writing', 'where is put writing'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry": {"type": "string"},
            },
            "required": ["symbol"],
        },
    ),
    "get_futures_analysis": (
        get_futures_analysis,
        (
            "Futures chain analysis: basis (futures vs spot), cost of carry (annualised), "
            "rollover percentage, OI across expiries. Good for gauging institutional positioning. "
            "Use for: 'futures basis', 'cost of carry', 'rollover', 'futures premium/discount', "
            "'long/short build-up in futures', 'what is NIFTY futures price'."
        ),
        {
            "type": "object",
            "properties": {"symbol": {"type": "string"}},
            "required": ["symbol"],
        },
    ),
    "get_options_strategy": (
        get_options_strategy,
        (
            "Build a specific options strategy with live pricing. Returns legs, strikes, "
            "entry costs, max risk, max reward, breakeven points, and payoff curve. "
            "Strategies: long_call, long_put, bull_call_spread, bear_put_spread, "
            "long_straddle, long_strangle, iron_condor, covered_call, protective_put, calendar_spread. "
            "Use for: 'set up a bull call spread on NIFTY', 'long straddle for earnings', "
            "'buy a call option', 'iron condor setup', 'options strategy for <stock>'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":   {"type": "string"},
                "strategy": {
                    "type": "string",
                    "enum": [
                        "long_call", "long_put", "bull_call_spread", "bear_put_spread",
                        "long_straddle", "long_strangle", "iron_condor",
                        "covered_call", "protective_put", "calendar_spread",
                    ],
                },
                "expiry": {"type": "string"},
            },
            "required": ["symbol", "strategy"],
        },
    ),
    "get_strategy_recommendations": (
        get_strategy_recommendations,
        (
            "Analyse current PCR, IV, DTE, and max pain to recommend the top 3 options strategies "
            "with rationale. Tailors recommendations to current market regime (trending/volatile/range-bound). "
            "Use for: 'what options strategy should I use', 'best strategy for NIFTY now', "
            "'options setup recommendation', 'which strategy fits current conditions'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "expiry": {"type": "string"},
            },
            "required": ["symbol"],
        },
    ),
    "refresh_fno_eod_data": (
        refresh_fno_eod_data,
        (
            "Download and store the latest F&O EOD bhavcopy from NSE. "
            "Returns summary of rows stored, options count, futures count, and trade date. "
            "Use when asked to 'update F&O data', 'download bhavcopy', 'refresh options data'."
        ),
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_fno_data_status": (
        get_fno_data_status,
        (
            "Check status of local F&O EOD database: available dates and available strategy list. "
            "Use at the start of any F&O session to confirm data availability."
        ),
        {"type": "object", "properties": {}, "required": []},
    ),
    "get_fno_analytics": (
        get_fno_analytics,
        (
            "PostgreSQL-backed F&O analytics for one symbol or top symbols. "
            "Returns PCR, max pain, max call/put OI strikes, futures OI/price change, "
            "buildup classification, and composite F&O signal. "
            "Use for: 'F&O analytics', 'FNO signal', 'long buildup', 'short buildup', "
            "'max pain distance', 'top F&O bearish/bullish names'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "top_n": {"type": "integer", "default": 20},
            },
            "required": [],
        },
    ),
    "run_option_payoff_scenario": (
        run_option_payoff_scenario,
        (
            "Run an option what-if payoff scenario from PostgreSQL F&O data. "
            "Returns P&L across underlying move percentages for a CE/PE strike. "
            "Use for: 'what if NIFTY moves 5%', 'option payoff', 'scenario analysis', "
            "'breakeven for call/put'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "option_type": {"type": "string", "enum": ["CE", "PE"], "default": "CE"},
                "strike": {"type": "number"},
                "expiry_date": {"type": "string"},
                "entry_premium": {"type": "number"},
                "lots": {"type": "integer", "default": 1},
                "move_start": {"type": "number", "default": -5.0},
                "move_end": {"type": "number", "default": 5.0},
                "move_step": {"type": "number", "default": 2.5},
            },
            "required": ["symbol"],
        },
    ),
})


# ── Chart tool wrapper ────────────────────────────────────────────────────────

def get_chart_summary(symbol: str, timeframe: str = "3mo") -> dict:
    """
    Get chart data summary for a symbol: current price, change%, RSI, MACD,
    EMA positions, period high/low, trend context.
    Does NOT render the ASCII chart — call render_chart() for that.
    """
    return chart_summary(symbol, timeframe)


def open_html_chart(symbol: str, timeframe: str = "3mo",
                    indicators: list | None = None) -> dict:
    """
    Generate and open a full interactive Plotly HTML chart in the browser.
    Returns the file path.
    """
    inds = indicators or ["volume", "rsi", "macd"]
    try:
        fpath = render_html_chart(symbol, timeframe, inds, open_browser=True)
        return {"status": "opened", "file": fpath, "symbol": symbol, "timeframe": timeframe}
    except Exception as e:
        return {"error": str(e)}

# Register chart + options buying tools
TOOL_REGISTRY.update({
    "get_chart_summary": (
        get_chart_summary,
        (
            "Get chart data and technical summary for a stock or index. Returns current price, "
            "daily change%, RSI(14), MACD signal (bullish/bearish), EMA20/EMA50 positions, "
            "period high/low. Use for: 'chart for X', 'technical levels', 'trend analysis', "
            "'is X above its moving average', 'show me the chart', 'price action for X'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":    {"type": "string", "description": "NSE symbol e.g. RELIANCE, NIFTY"},
                "timeframe": {
                    "type": "string",
                    "enum": ["1d", "5d", "1mo", "3mo", "6mo", "1y", "2y"],
                    "description": "Chart timeframe (default 3mo)",
                },
            },
            "required": ["symbol"],
        },
    ),
    "open_html_chart": (
        open_html_chart,
        (
            "Generate and open a full interactive HTML chart (Plotly) in the browser. "
            "Shows candlestick + EMA20/50/200 + Bollinger Bands + Volume + RSI + MACD "
            "in a professional dark TradingView-style layout. Fully interactive: zoom, pan, hover. "
            "Use for: 'open chart', 'interactive chart', 'show html chart', 'open in browser', "
            "'full chart for X', 'detailed chart'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":     {"type": "string"},
                "timeframe":  {"type": "string", "enum": ["1d", "5d", "1mo", "3mo", "6mo", "1y", "2y"]},
                "indicators": {"type": "array", "items": {"type": "string"},
                               "description": "['volume','rsi','macd'] — default: all three"},
            },
            "required": ["symbol"],
        },
    ),
    "run_visual_scan": (
        run_visual_scan,
        (
            "Generate a grounded swing/EOD visual scan report for one NSE symbol with "
            "annotated charts, deterministic pattern evidence, MTF alignment, optional "
            "TradingView corroboration, and a research stance."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "capture_tradingview": {"type": "boolean", "default": True},
            },
            "required": ["symbol"],
        },
    ),
    "analyze_options_buying": (
        analyze_options_buying,
        (
            "Deep options buying analysis for a symbol. Returns ATM IV + regime (cheap/fair/rich), "
            "IV rank estimate, expected move (±1σ/±2σ), strike selection guide (ITM/ATM/OTM with "
            "delta/theta/breakeven/probability), theta decay profile, OI context, and buying verdict "
            "(BUY / USE SPREAD / AVOID). "
            "Use for: 'should I buy calls on X', 'best strike to buy', 'options buying setup for X', "
            "'what call should I buy', 'options trade idea for X', 'options opportunity for X'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol":    {"type": "string"},
                "direction": {
                    "type": "string",
                    "enum": ["bullish", "bearish", "volatile"],
                    "description": "Trade direction (default: bullish)",
                },
                "expiry":    {"type": "string", "description": "Expiry date YYYY-MM-DD (optional)"},
            },
            "required": ["symbol"],
        },
    ),
    "scan_options_buys": (
        scan_options_buys,
        (
            "Scan all F&O-eligible stocks/indices for the best options buying opportunities. "
            "Ranks by: low ATM IV (cheap options) + OI liquidity + ideal DTE. "
            "Returns top symbols with ATM IV, straddle cost, expected move, and buying score. "
            "Use for: 'scan for options buying opportunities', 'which stocks have cheap options', "
            "'best stocks for buying calls today', 'low IV options scan', 'options buying scan'."
        ),
        {
            "type": "object",
            "properties": {
                "direction": {"type": "string", "enum": ["bullish", "bearish", "volatile"]},
                "max_iv":    {"type": "number", "description": "Max ATM IV% (default 25)"},
                "min_oi":    {"type": "integer", "description": "Min OI for liquidity (default 500000)"},
                "top_n":     {"type": "integer", "description": "Number of results to return (default 10)"},
            },
            "required": [],
        },
    ),
})


# Register forensic accounting + event calendar tools
TOOL_REGISTRY.update({
    "run_forensic_analysis": (
        run_forensic_analysis,
        (
            "Run D5 Forensic Accounting analysis on an NSE stock. "
            "Computes three quantitative red-flag models: "
            "(1) Beneish M-score — detects earnings manipulation (M > -1.78 = manipulation risk); "
            "(2) Piotroski F-score — financial health 0-9 (7-9 = strong, 0-3 = weak); "
            "(3) Altman Z'-score — bankruptcy/distress risk (Z' < 1.1 = distress zone). "
            "Data from screener.in annual balance sheet + P&L + cash flow statements. "
            "Use for: forensic analysis, earnings quality, manipulation risk, financial health, "
            "due diligence, red flags, balance sheet quality, accounting risk."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string", "description": "NSE ticker e.g. RELIANCE, ADANIENT"},
            },
            "required": ["symbol"],
        },
    ),
    "screen_forensic_watchlist": (
        screen_forensic_watchlist,
        (
            "Run forensic accounting screening across multiple NSE stocks in parallel. "
            "Returns all three forensic scores (Beneish, Piotroski, Altman) for each stock, "
            "ranked by risk level. Use for portfolio-level forensic due diligence, "
            "pre-buy checklist, or identifying high-risk holdings."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of NSE tickers to screen (max 8)",
                },
            },
            "required": ["symbols"],
        },
    ),
    "search_broker_research": (
        search_broker_research,
        (
            "Search for broker house research reports, institutional analyst ratings, "
            "and consensus price targets for an NSE stock. "
            "Searches: Trendlyne (consensus estimates), Moneycontrol (broker radar), "
            "Economic Times (analyst reports), major brokers (Motilal/Kotak/ICICI/HDFC/Edelweiss/Axis). "
            "Extracts price targets from report titles where available. "
            "Use for: broker reports, analyst ratings, price targets, buy/sell recommendations, "
            "institutional views, consensus estimate."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string", "description": "NSE ticker e.g. TCS"},
            },
            "required": ["symbol"],
        },
    ),
    "search_mf_holdings": (
        search_mf_holdings,
        (
            "Search for mutual fund holdings and FII/DII institutional ownership data for a stock. "
            "Combines screener.in direct shareholding scrape (promoter/FII/DII quarterly trend) "
            "with web search across Trendlyne, Moneycontrol, Tijori Finance. "
            "Use for: MF holdings, institutional ownership, FII activity, DII buying/selling, "
            "promoter pledge, shareholding pattern changes."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string", "description": "NSE ticker e.g. HDFCBANK"},
            },
            "required": ["symbol"],
        },
    ),
    "get_upcoming_events": (
        get_upcoming_events,
        (
            "E4 Event-Driven Alert Engine: fetch upcoming corporate action events for a list of "
            "symbols or an entire index. Tracks dividends, bonus issues, stock splits, rights issues, "
            "AGMs, board meetings, and results calendar. Returns events grouped by date and type, "
            "with days-until countdown. "
            "Use for: upcoming dividends, ex-date, bonus issues, stock splits, results dates, "
            "event calendar, corporate actions calendar, board meeting dates."
        ),
        {
            "type": "object",
            "properties": {
                "symbols":    {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "NSE tickers to check (leave empty to use index symbols)",
                },
                "index":      {
                    "type": "string",
                    "description": "Index to scan when symbols is empty (default 'NIFTY 50')",
                    "default": "NIFTY 50",
                },
                "days_ahead": {
                    "type": "integer",
                    "description": "Calendar days ahead to look (default 30)",
                    "default": 30,
                },
                "event_types": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Filter to specific types e.g. ['Dividend', 'Results', 'Bonus', 'Split']",
                },
            },
            "required": [],
        },
    ),
    "get_event_calendar_summary": (
        get_event_calendar_summary,
        (
            "Quick event calendar overview for an index — compact list of upcoming dividends, "
            "splits, bonuses, results announcements, and board meetings in the next N days. "
            "Lighter than get_upcoming_events. Use for: what events this week/fortnight, "
            "upcoming ex-dates, results season, event summary for NIFTY."
        ),
        {
            "type": "object",
            "properties": {
                "index":      {
                    "type": "string",
                    "description": "Index name e.g. 'NIFTY 50' (default)",
                    "default": "NIFTY 50",
                },
                "days_ahead": {
                    "type": "integer",
                    "description": "Days ahead to look (default 14)",
                    "default": 14,
                },
            },
            "required": [],
        },
    ),
    "get_latest_results_feed": (
        get_latest_results_feed,
        (
            "Market-wide feed of companies that have just declared quarterly financial "
            "results across NSE. Pulls NSE corporates-financial-results JSON (cached 30 min) "
            "with screener.in /results/latest fallback. "
            "Use for: 'latest results', 'who reported today', 'results this week', "
            "'companies that announced results', 'results posted recently'. Do NOT use "
            "when the user names a specific symbol — use stock_results path instead."
        ),
        {
            "type": "object",
            "properties": {
                "days_back": {
                    "type": "integer",
                    "description": "Calendar days back to include (default 7).",
                    "default": 7,
                },
                "limit": {
                    "type": "integer",
                    "description": "Max rows to return (default 50).",
                    "default": 50,
                },
            },
            "required": [],
        },
    ),
    "get_forthcoming_results": (
        get_forthcoming_results,
        (
            "Forthcoming results / earnings events from NSE event-calendar — companies "
            "with scheduled board meetings to declare quarterly financial results. "
            "Use for: 'results due this week', 'who is reporting tomorrow', 'upcoming "
            "earnings', 'forthcoming results', 'results calendar this week'. Returns "
            "symbol, company, scheduled date, purpose, board-meeting description; "
            "sorted earliest first."
        ),
        {
            "type": "object",
            "properties": {
                "days_ahead": {
                    "type": "integer",
                    "description": "Calendar days ahead to include (default 14).",
                    "default": 14,
                },
                "limit": {
                    "type": "integer",
                    "description": "Max rows to return (default 50).",
                    "default": 50,
                },
            },
            "required": [],
        },
    ),
})

# ── B3, B5, D4, P2-2, P2-4, P3-2: New tool registrations ────────────────────
TOOL_REGISTRY.update({
    "get_sector_heat_calendar": (
        get_sector_heat_calendar,
        (
            "B3 Sectoral Heat Calendar — 12-month seasonal return heatmap for NSE sector indices. "
            "Shows which sectors historically perform best in each month (TAILWIND/HEADWIND/NEUTRAL). "
            "Use for: sector rotation timing, seasonal investing, which sectors favour current month, "
            "historical monthly patterns, seasonal tailwinds and headwinds."
        ),
        {
            "type": "object",
            "properties": {
                "month": {
                    "type": "integer",
                    "description": "Target month 1-12 (default: current month)",
                },
            },
            "required": [],
        },
    ),
    "get_economic_cycle_assessment": (
        get_economic_cycle_assessment,
        (
            "B5 Economic Cycle Tracker — detect current macro cycle phase from proxy signals. "
            "Returns EARLY_EXPANSION / LATE_EXPANSION / SLOWDOWN / RECOVERY with confidence, "
            "preferred sectors to overweight, sectors to avoid, and full macro snapshot. "
            "Use for: macro-driven sector allocation, positioning in current cycle phase, "
            "understanding where we are in the business cycle, rate/inflation/commodity regime."
        ),
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
    ),
    "analyze_concall_sentiment": (
        analyze_concall_sentiment,
        (
            "D4 Concall NLP — extract management tone, key themes, risk flags, and sentiment "
            "from the most recent earnings call or investor day transcript. "
            "Returns: sentiment (Bullish/Cautious/Bearish/Neutral), tone score (-1 to +1), "
            "top themes, risk flags, key management quotes, guidance summary. "
            "Use for: post-results management tone check, earnings quality assessment, "
            "concall digest, what management said about margins/growth/guidance."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {
                    "type": "string",
                    "description": "NSE ticker symbol (e.g. 'TCS', 'RELIANCE')",
                },
            },
            "required": ["symbol"],
        },
    ),
    "run_scenario_analysis": (
        run_scenario_analysis,
        (
            "P2-2 Scenario Engine — what-if price analysis for a stock at various levels. "
            "For each hypothetical price: % change, RSI estimate, stage implication (Stage 2/3/4), "
            "proximity to key levels (support, resistance, 50-DMA, 200-DMA). "
            "Use for: what happens if stock drops 10%? which level triggers Stage 4? "
            "risk/reward analysis, stop-loss placement, bull/base/bear case planning."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {
                    "type": "string",
                    "description": "NSE ticker symbol (e.g. 'TCS')",
                },
                "price_scenarios": {
                    "type": "array",
                    "items": {"type": "number"},
                    "description": "List of price levels to evaluate. Default: ±5/10/20% from current.",
                },
                "scenario_labels": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Labels for each price scenario (e.g. ['Base', 'Bull', 'Bear'])",
                },
            },
            "required": ["symbol"],
        },
    ),
    "generate_portfolio_narratives": (
        generate_portfolio_narratives,
        (
            "P2-4 Portfolio Narrative Engine — per-stock investment narratives combining "
            "stage analysis, RSI, RS, signals, and fundamental snapshot. "
            "Returns bull thesis, bear case, and action hint for each stock. "
            "Use for: portfolio review, morning briefing narrative, stock-by-stock commentary, "
            "investment thesis validation, portfolio health check."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of symbols to narrate. Default: top portfolio holdings.",
                },
                "top_n": {
                    "type": "integer",
                    "description": "Max number of stocks to narrate (default 5)",
                    "default": 5,
                },
            },
            "required": [],
        },
    ),
    "generate_voice_briefing": (
        generate_voice_briefing,
        (
            "P3-2 Voice Briefing — convert market summary to an MP3 audio briefing using OpenAI TTS. "
            "Auto-generates a 60-second daily market briefing if no text provided. "
            "Saves to data/voice_briefing.mp3. Requires OPENAI_API_KEY. "
            "Use for: hands-free morning briefing, audio market update, voice digest."
        ),
        {
            "type": "object",
            "properties": {
                "text": {
                    "type": "string",
                    "description": "Custom text to convert. If None, auto-generates from market overview.",
                },
                "voice": {
                    "type": "string",
                    "description": "TTS voice: alloy, echo, fable, onyx, nova, shimmer (default: alloy)",
                    "default": "alloy",
                },
                "save_path": {
                    "type": "string",
                    "description": "File path for MP3 output. Default: data/voice_briefing.mp3",
                },
            },
            "required": [],
        },
    ),
})


# ── Watchlist alert tools ──────────────────────────────────────────────────


def get_watchlist_alerts(_args: dict) -> dict:
    from terminal.alerts import list_alerts
    alerts = list_alerts()
    return {"alerts": alerts, "count": len(alerts)}


def add_watchlist_alert(args: dict) -> dict:
    from terminal.alerts import add_alert
    sym = args.get("symbol", "").upper()
    trigger = args.get("trigger", "price_above")
    value = float(args.get("value", 0))
    note = args.get("note", "")
    alert = add_alert(sym, trigger, value, note)
    return {"added": alert}


def delete_watchlist_alert(args: dict) -> dict:
    from terminal.alerts import delete_alert
    aid = int(args.get("alert_id", 0))
    ok = delete_alert(aid)
    return {"deleted": ok, "alert_id": aid}


def check_watchlist_alerts(_args: dict) -> dict:
    from terminal.alerts import check_alerts
    triggered = check_alerts()
    return {"triggered": triggered, "count": len(triggered)}


TOOL_REGISTRY.update({
    "get_watchlist_alerts": (
        get_watchlist_alerts,
        "List all active price/RSI alerts",
        {"type": "object", "properties": {}},
    ),
    "add_watchlist_alert": (
        add_watchlist_alert,
        "Add a price or RSI alert for a symbol",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "trigger": {"type": "string", "enum": ["price_above", "price_below", "rsi_above", "rsi_below"]},
                "value": {"type": "number"},
                "note": {"type": "string"},
            },
        },
    ),
    "delete_watchlist_alert": (
        delete_watchlist_alert,
        "Delete an alert by ID",
        {"type": "object", "properties": {"alert_id": {"type": "integer"}}},
    ),
    "check_watchlist_alerts": (
        check_watchlist_alerts,
        "Check all alerts against live prices/RSI and fire macOS notifications for triggered ones",
        {"type": "object", "properties": {}},
    ),
})


TOOL_REGISTRY.update({
    "get_options_chain": (
        get_options_chain,
        (
            "Fetch live NSE options chain with full OI, IV, LTP, PCR, and max pain. "
            "Returns side-by-side calls/puts data for all strikes near ATM, total OI, "
            "PCR (put-call ratio), max pain strike, and available expiry dates. "
            "Use for: 'options chain', 'OI buildup', 'PCR', 'max pain', 'IV skew', "
            "'option chain for NIFTY/BANKNIFTY/<stock>', 'show me calls and puts'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {
                    "type": "string",
                    "description": "Index (NIFTY/BANKNIFTY/FINNIFTY) or equity symbol",
                },
                "expiry_index": {
                    "type": "integer",
                    "description": "0=nearest expiry, 1=next, etc.",
                    "default": 0,
                },
            },
            "required": [],
        },
    ),
})


def get_portfolio_pnl(_args: dict = None) -> dict:
    """Compute live unrealised P&L for all holdings in data/holdings.csv."""
    try:
        from terminal.portfolio_pnl import compute_pnl
        return compute_pnl()
    except Exception as e:
        return {"error": str(e)}


TOOL_REGISTRY.update({
    "get_portfolio_pnl": (
        get_portfolio_pnl,
        (
            "Live Portfolio P&L — compute unrealised gains/losses for all holdings in data/holdings.csv. "
            "Fetches live prices for each position and returns per-stock P&L, day P&L, invested value, "
            "current value, P&L %, and portfolio totals. "
            "Use for: 'portfolio P&L', 'my holdings', 'unrealised gains', 'unrealised losses', "
            "'how is my portfolio doing', 'portfolio performance', 'check my holdings'."
        ),
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
    ),
})


def export_session(args: dict) -> dict:
    """Export current session to HTML or PDF."""
    fmt = args.get("format", "html")
    symbol = args.get("symbol", "")
    return {"message": f"Use /export {fmt} {symbol} command to export the session."}


TOOL_REGISTRY.update({
    "export_session": (
        export_session,
        "Export the current research session to HTML or PDF report",
        {
            "type": "object",
            "properties": {
                "format": {"type": "string", "enum": ["html", "pdf"], "default": "html"},
                "symbol": {"type": "string"},
            },
        },
    ),
    "fetch_article_content": (
        fetch_article_content,
        (
            "Fetch and read the full text of a news article from its URL. "
            "Use this to get deeper context from articles found via search tools "
            "when you need to provide a thorough summary and opinion."
        ),
        {
            "type": "object",
            "properties": {
                "url":       {"type": "string", "description": "Full URL of the article to fetch"},
                "max_chars": {"type": "integer", "default": 3000},
            },
            "required": ["url"],
        },
    ),
    "fetch_pdf_text": (
        fetch_pdf_text,
        (
            "Download and extract text from a PDF at any URL — BSE financial results, "
            "annual reports, concall transcripts, SEBI filings, NSE circulars. "
            "Use whenever you have a direct PDF URL (e.g. from search_bse_filings, "
            "search_nse_announcements, or scrape_screener_in) and the user asks to "
            "read, summarise, or analyse the document contents. "
            "Returns page-by-page extracted text up to max_pages (default 15). "
            "Works on BSE/NSE hosted PDFs, screener.in transcript PDFs, and broker reports."
        ),
        {
            "type": "object",
            "properties": {
                "url":       {"type": "string",  "description": "Direct URL to the PDF file"},
                "max_pages": {"type": "integer", "default": 15,
                              "description": "Maximum pages to extract (default 15)"},
            },
            "required": ["url"],
        },
    ),
})


# ─────────────────────────────────────────────────────────────────────────────
# /analyze — Local document reader (PDF, DOCX, TXT) + web page scraper
# PG: First-class document analysis capability for Agent Adda
# ─────────────────────────────────────────────────────────────────────────────

def analyze_document(source: str, max_pages: int = 50,
                     vision_fallback: bool = True,
                     vision_threshold: int = 200) -> dict:
    """Read and extract text from a local file (PDF, DOCX, HTML, TXT, CSV) or a web URL.

    Detects the source type automatically:
      - URL (http/https) → scrape web page or download PDF
      - .pdf             → extract text with PyMuPDF page by page
      - .docx            → extract text with python-docx paragraph by paragraph
      - .html / .htm     → strip tags via BeautifulSoup, extract headings + tables
      - .txt / .csv / .md → read as plain text

    For PDFs, pages whose text extraction yields < ``vision_threshold`` chars
    are automatically re-transcribed using the OpenAI vision model so that
    scanned pages and image-only tables are recovered.

    Returns structured dict with text content, metadata, and page/section info.
    """
    import os
    source = source.strip()

    # ── URL path ──────────────────────────────────────────────────────────
    if source.lower().startswith(("http://", "https://")):
        source = _normalise_http_url(source)
        resolved_pdf_url = _resolve_embedded_pdf_url(source)
        # Check if it's a PDF URL
        if resolved_pdf_url or source.lower().endswith(".pdf") or "/pdf/" in source.lower():
            return fetch_pdf_text(source, max_pages=max_pages,
                                  vision_fallback=vision_fallback,
                                  vision_threshold=vision_threshold)
        # Otherwise scrape as web page
        result = fetch_article_content(source, max_chars=15000)
        result["source_type"] = "web_page"
        return result

    # ── Local file path ───────────────────────────────────────────────────
    path = os.path.expanduser(source)
    if not os.path.isabs(path):
        # Try relative to common locations
        for base in [os.getcwd(), os.path.expanduser("~/Documents"),
                     os.path.expanduser("~/Downloads")]:
            candidate = os.path.join(base, path)
            if os.path.isfile(candidate):
                path = candidate
                break

    if not os.path.isfile(path):
        return {"error": f"File not found: {source}", "source": source}

    ext = os.path.splitext(path)[1].lower()
    file_size = os.path.getsize(path)
    file_name = os.path.basename(path)

    # ── PDF ───────────────────────────────────────────────────────────────
    if ext == ".pdf":
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(path)
            total_pages = len(doc)
            pages_to_read = min(total_pages, max_pages)
            pages_text = []
            full_text_parts = []
            for i in range(pages_to_read):
                page = doc[i]
                text = page.get_text("text").strip()
                method = "text"
                if vision_fallback and len(text) < vision_threshold:
                    vision_text = _vision_transcribe_page(page, i + 1)
                    if vision_text and len(vision_text) > len(text):
                        text = vision_text
                        method = "vision"
                if text:
                    pages_text.append({"page": i + 1, "text": text, "extraction_method": method})
                    full_text_parts.append(f"--- Page {i + 1} ({method}) ---\n{text}")
            doc.close()
            full_text = "\n\n".join(full_text_parts)
            return {
                "source": path, "source_type": "pdf", "file_name": file_name,
                "file_size_kb": round(file_size / 1024, 1),
                "total_pages": total_pages, "pages_read": pages_to_read,
                "truncated": total_pages > max_pages,
                "text": full_text, "pages": pages_text,
            }
        except ImportError:
            return {"error": "PyMuPDF (fitz) not installed. Run: pip install pymupdf",
                    "source": path}
        except Exception as e:
            return {"error": f"PDF read error: {e}", "source": path}

    # ── DOCX ──────────────────────────────────────────────────────────────
    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    style = para.style.name if para.style else "Normal"
                    paragraphs.append({"style": style, "text": text})

            # Extract tables
            tables_text = []
            for t_idx, table in enumerate(doc.tables):
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    tables_text.append(f"--- Table {t_idx + 1} ---\n" + "\n".join(rows))

            # Build full text
            full_parts = [p["text"] for p in paragraphs]
            if tables_text:
                full_parts.append("\n\n=== TABLES ===\n" + "\n\n".join(tables_text))
            full_text = "\n\n".join(full_parts)

            return {
                "source": path, "source_type": "docx", "file_name": file_name,
                "file_size_kb": round(file_size / 1024, 1),
                "paragraph_count": len(paragraphs),
                "table_count": len(doc.tables),
                "text": full_text,
                "sections": paragraphs[:100],  # Cap section detail
            }
        except ImportError:
            return {"error": "python-docx not installed. Run: pip install python-docx",
                    "source": path}
        except Exception as e:
            return {"error": f"DOCX read error: {e}", "source": path}

    # ── HTML ──────────────────────────────────────────────────────────────
    if ext in (".html", ".htm"):
        try:
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                raw = f.read()
            title = ""
            sections: list[dict] = []
            tables_text: list[str] = []
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(raw, "lxml") if "lxml" in str(BeautifulSoup) else BeautifulSoup(raw, "html.parser")
                for tag in soup(["script", "style", "noscript"]):
                    tag.decompose()
                if soup.title and soup.title.string:
                    title = soup.title.string.strip()
                for h in soup.find_all(["h1", "h2", "h3", "h4"]):
                    htext = h.get_text(" ", strip=True)
                    if htext:
                        sections.append({"style": h.name.upper(), "text": htext})
                for t_idx, table in enumerate(soup.find_all("table")):
                    rows = []
                    for tr in table.find_all("tr"):
                        cells = [c.get_text(" ", strip=True) for c in tr.find_all(["th", "td"])]
                        if cells:
                            rows.append(" | ".join(cells))
                    if rows:
                        tables_text.append(f"--- Table {t_idx + 1} ---\n" + "\n".join(rows))
                body_text = soup.get_text("\n", strip=True)
            except ImportError:
                import re as _re
                body_text = _re.sub(r"<script[^>]*>.*?</script>", " ", raw, flags=_re.S | _re.I)
                body_text = _re.sub(r"<style[^>]*>.*?</style>", " ", body_text, flags=_re.S | _re.I)
                body_text = _re.sub(r"<[^>]+>", " ", body_text)
                import html as _html_mod
                body_text = _html_mod.unescape(_re.sub(r"\s+\n", "\n", body_text)).strip()

            body_text = "\n".join(line for line in (ln.strip() for ln in body_text.splitlines()) if line)
            max_chars = 200_000
            truncated = len(body_text) > max_chars
            if truncated:
                body_text = body_text[:max_chars]
            full_parts = []
            if title:
                full_parts.append(f"# {title}")
            full_parts.append(body_text)
            if tables_text:
                full_parts.append("\n=== TABLES ===\n" + "\n\n".join(tables_text))
            full_text = "\n\n".join(full_parts)
            return {
                "source": path, "source_type": "html", "file_name": file_name,
                "file_size_kb": round(file_size / 1024, 1),
                "title": title,
                "section_count": len(sections),
                "table_count": len(tables_text),
                "truncated": truncated,
                "text": full_text,
                "sections": sections[:100],
            }
        except Exception as e:
            return {"error": f"HTML read error: {e}", "source": path}

    # ── Plain text (TXT, CSV, MD, etc.) ───────────────────────────────────
    if ext in (".txt", ".csv", ".md", ".log", ".json", ".yaml", ".yml"):
        try:
            max_bytes = 50_000  # ~50KB cap for plain text
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read(max_bytes)
            truncated = file_size > max_bytes
            return {
                "source": path, "source_type": ext.lstrip("."), "file_name": file_name,
                "file_size_kb": round(file_size / 1024, 1),
                "truncated": truncated,
                "text": text,
            }
        except Exception as e:
            return {"error": f"Read error: {e}", "source": path}

    return {"error": f"Unsupported file type: {ext}. Supported: .pdf, .docx, .html, .txt, .csv, .md",
            "source": path}


TOOL_REGISTRY.update({
    "analyze_document": (
        analyze_document,
        ("Read and extract text from a local file (PDF, DOCX, HTML, TXT, CSV) or a web URL. "
         "Use this for document analysis — annual reports, research papers, filings, "
         "concall transcripts, presentations, dashboards. Returns full text content with metadata. "
         "For PDFs: page-by-page extraction with automatic vision/OCR fallback "
         "for image-only or scanned pages. For DOCX: paragraphs + tables. "
         "For HTML: heading sections + tables via BeautifulSoup. "
         "For URLs: smart web scraping or PDF download."),
        {
            "type": "object",
            "properties": {
                "source": {
                    "type": "string",
                    "description": "File path (local PDF/DOCX/TXT) or URL (http/https web page or PDF link)",
                },
                "max_pages": {
                    "type": "integer",
                    "description": "Max pages to read for PDFs (default 50)",
                    "default": 50,
                },
                "vision_fallback": {
                    "type": "boolean",
                    "description": "If true, pages with little extractable text are re-transcribed via OpenAI vision (default true).",
                    "default": True,
                },
            },
            "required": ["source"],
        },
    ),
})


# ─────────────────────────────────────────────────────────────────────────────
# Multi-timeframe (MTF) analysis tools — PG 2026-05-22
# Engine lives in terminal/mtf.py. These wrappers shape inputs/outputs for the
# LLM tool-calling surface and add a multi-symbol scanner that reuses the same
# scoring as the single-symbol analyzer.
# ─────────────────────────────────────────────────────────────────────────────


def analyze_mtf(
    symbol: str,
    timeframes: list[str] | None = None,
    days: int = 800,
) -> dict:
    """Aligned multi-timeframe analysis for a single symbol.

    Returns the per-timeframe indicator stack (RSI/MACD/EMA20/EMA50/SMA stack)
    for monthly, weekly, daily, 60m, 15m, a weighted confluence score (0-100),
    a deterministic BUY/WATCH/AVOID/SELL verdict, and the rationale used.

    Missing timeframes are reported with status="missing" and never inferred.
    """
    from terminal.mtf import compute_mtf, DEFAULT_TIMEFRAMES

    sym = _canonical_symbol(symbol)
    tfs = tuple(timeframes) if timeframes else DEFAULT_TIMEFRAMES
    try:
        result = compute_mtf(sym, timeframes=tfs, days=days)
    except Exception as exc:
        return {"symbol": sym, "error": f"MTF computation failed: {exc}"}
    out = result.as_dict()
    out["data_source"] = "PostgreSQL market.equity_eod (daily resampled) + PG intraday.ohlcv_bars"
    return out


def scan_mtf_aligned(
    symbols: list[str] | None = None,
    index: str | None = None,
    direction: str = "bullish",
    min_score: int = 70,
    timeframes: list[str] | None = None,
    top_n: int = 10,
) -> dict:
    """Rank a universe by MTF confluence in a given direction.

    Provide EITHER an explicit ``symbols`` list OR an NSE ``index`` name
    (e.g. 'NIFTY 50', 'NIFTY 500'). Symbols path is preferred for tight
    fan-outs (faster, cheaper). Returns ranked top_n symbols whose MTF
    direction matches and whose confluence_score >= min_score.
    """
    from terminal.mtf import compute_mtf, DEFAULT_TIMEFRAMES

    direction = (direction or "bullish").lower()
    if direction not in {"bullish", "bearish"}:
        return {"error": "direction must be 'bullish' or 'bearish'"}

    universe: list[str] = []
    if symbols:
        universe = [_canonical_symbol(s) for s in symbols if s and s.strip()]
    elif index:
        # PG-SCAN-FALLBACK: use the shared helper (NSE live + local CSV fallback)
        try:
            universe = _fetch_nse_index_constituents(index)
        except Exception as exc:
            return {"error": f"Could not fetch {index} constituents: {exc}"}
        if not universe:
            return {"error": f"No stocks found for index: {index}"}
    else:
        return {"error": "Provide either 'symbols' or 'index'."}

    universe = list(dict.fromkeys(universe))[:200]
    if not universe:
        return {"error": "Empty universe after dedup."}

    tfs = tuple(timeframes) if timeframes else DEFAULT_TIMEFRAMES

    ranked: list[dict] = []
    errors: list[dict] = []
    for sym in universe:
        try:
            res = compute_mtf(sym, timeframes=tfs)
        except Exception as exc:
            errors.append({"symbol": sym, "error": str(exc)[:120]})
            continue
        if res.direction != direction:
            continue
        if res.confluence_score < int(min_score):
            continue
        ranked.append(
            {
                "symbol": sym,
                "direction": res.direction,
                "confluence_score": res.confluence_score,
                "verdict": res.verdict,
                "aligned_tfs": res.aligned_tfs,
                "missing_tfs": res.missing_tfs,
            }
        )

    ranked.sort(key=lambda r: (-r["confluence_score"], r["symbol"]))
    return {
        "direction": direction,
        "min_score": int(min_score),
        "timeframes": list(tfs),
        "universe_size": len(universe),
        "matches_total": len(ranked),
        "top": ranked[: int(top_n)],
        "errors": errors[:10],
        "data_source": "PostgreSQL market.equity_eod (daily resampled) + PG intraday.ohlcv_bars",
    }


TOOL_REGISTRY.update({
    "analyze_mtf": (
        analyze_mtf,
        (
            "Multi-timeframe (MTF) technical analysis for a single NSE stock. "
            "Computes RSI/MACD/EMA20/EMA50/SMA stack across monthly, weekly, daily, 60m, 15m, "
            "produces a weighted confluence score (0-100) and a deterministic "
            "BUY/WATCH/AVOID/SELL verdict with per-timeframe rationale. "
            "Missing timeframes are reported, never inferred. "
            "Use for: 'multi timeframe analysis of X', 'MTF view on X', "
            "'are weekly/daily aligned on X', 'is X a confluent buy'."
        ),
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string", "description": "NSE ticker"},
                "timeframes": {
                    "type": "array",
                    "items": {"type": "string", "enum": ["monthly", "weekly", "daily", "60m", "15m"]},
                    "description": "Subset of timeframes to analyse; defaults to all five.",
                },
                "days": {
                    "type": "integer",
                    "description": "Daily history depth in days (default 800 ≈ 3 years).",
                    "default": 800,
                },
            },
            "required": ["symbol"],
        },
    ),
    "scan_mtf_aligned": (
        scan_mtf_aligned,
        (
            "Rank an NSE universe by multi-timeframe confluence in a chosen direction. "
            "Provide EITHER 'symbols' (preferred — explicit list, fast) OR 'index' "
            "(NSE index name, fetches constituents live). Returns the top_n stocks whose "
            "MTF direction matches and whose confluence_score >= min_score. "
            "Use for: 'top picks where weekly+daily agree', "
            "'recommendation report — bullish across timeframes', 'MTF scan NIFTY 50'."
        ),
        {
            "type": "object",
            "properties": {
                "symbols": {"type": "array", "items": {"type": "string"}},
                "index": {"type": "string", "description": "NSE index name, e.g. 'NIFTY 50'"},
                "direction": {"type": "string", "enum": ["bullish", "bearish"], "default": "bullish"},
                "min_score": {"type": "integer", "default": 70},
                "timeframes": {
                    "type": "array",
                    "items": {"type": "string", "enum": ["monthly", "weekly", "daily", "60m", "15m"]},
                },
                "top_n": {"type": "integer", "default": 10},
            },
            "required": [],
        },
    ),
})


# ─────────────────────────────────────────────────────────────────────────────
# /report — First-class report generation (PDF, HTML, Markdown)
# PG: Enables any LLM output to be written as a formatted report file.
# ─────────────────────────────────────────────────────────────────────────────

def _research_council_state_from_payload(payload: Any):
    from terminal.research_council.schemas import CouncilState

    if isinstance(payload, CouncilState):
        return payload
    if isinstance(payload, dict):
        return CouncilState.from_dict(payload)
    raise ValueError("Expected CouncilState or CouncilState dict")


def _research_council_state_summary(state: Any) -> dict:
    decision = getattr(state, "decision", None)
    flags = getattr(state, "flags", {}) or {}
    return {
        "ok": True,
        "run_id": getattr(state, "run_id", None),
        "mode": getattr(state, "mode", None),
        "stage": getattr(state, "stage", None),
        "horizon": getattr(state, "horizon", None),
        "risk_budget": getattr(state, "risk_budget", None),
        "evidence_only": bool(flags.get("evidence_only")),
        "final_label": getattr(decision, "final_label", None) if decision else None,
        "confidence": getattr(decision, "confidence", None) if decision else None,
        "candidate_count": len(getattr(decision, "candidates", []) or []) if decision else 0,
        "report_paths": {
            "markdown": flags.get("markdown_report_path"),
            "html": flags.get("html_report_path") or getattr(state, "html_path", None),
        },
        "events": list(getattr(state, "events", []) or []),
    }


def build_research_evidence_pack(
    mode: str = "market_council",
    universe_filter: str = "liquid",
    symbols: list[str] | None = None,
    max_stock_candidates: int = 50,
) -> dict:
    """Build the Research Council evidence pack from PostgreSQL-backed sources."""
    try:
        from terminal.research_council.evidence_pack_builder import build_research_evidence_pack as _build

        pack = _build(
            mode=mode,
            universe_filter=universe_filter,
            symbols=symbols or [],
            max_stock_candidates=int(max_stock_candidates),
        )
        return {"ok": True, "pack": pack.to_dict(), "pack_id": pack.pack_id}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def run_research_council(
    objective: str,
    mode: str | None = None,
    symbols: list[str] | None = None,
    horizon: str | None = None,
    risk_budget: str | None = None,
    dry_run: bool = False,
    output_format: str = "md",
    **flags: Any,
) -> dict:
    """Run the Research Council state machine and return a compact summary."""
    try:
        from terminal.research_council.engine import run_council

        run_flags = dict(flags)
        if mode:
            run_flags["mode"] = mode
        if symbols is not None:
            run_flags["symbols"] = symbols
        if horizon:
            run_flags["horizon"] = horizon
        if risk_budget:
            run_flags["risk_budget"] = risk_budget
        run_flags["dry_run"] = bool(dry_run)
        run_flags["output_format"] = output_format
        state = run_council(objective, **run_flags)
        return _research_council_state_summary(state)
    except Exception as exc:
        return {"ok": False, "error": str(exc), "objective": objective}


def run_data_steward_check(mode: str = "market_council") -> dict:
    """Run the Research Council data-steward freshness/readiness gate."""
    try:
        from terminal.research_council.states.data_steward import run_check

        verdict = run_check(mode=mode)
        return {"ok": True, "mode": mode, "verdict": verdict.to_dict()}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "mode": mode}


def compose_plan(objective: str, mode: str = "market_council", **flags: Any) -> dict:
    """Compose a deterministic Research Council evidence plan."""
    try:
        from terminal.research_council.engine import initialize_state
        from terminal.research_council.states import plan_build

        state = initialize_state(objective, mode=mode, **flags)
        state = plan_build.run(state)
        return {"ok": True, "run_id": state.run_id, "plans": [plan.to_dict() for plan in state.plans]}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def execute_plan(state: dict | Any) -> dict:
    """Execute the latest Research Council plan in a serialized state."""
    try:
        from terminal.research_council.states import plan_execute

        updated = plan_execute.run(_research_council_state_from_payload(state))
        return {
            "ok": True,
            "run_id": updated.run_id,
            "plans": [plan.to_dict() for plan in updated.plans],
            "execution_results": {
                plan_id: {step_id: result.to_dict() for step_id, result in results.items()}
                for plan_id, results in updated.execution_results.items()
            },
        }
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def review_plan_execution(state: dict | Any) -> dict:
    """Review plan execution output and decide whether more evidence is needed."""
    try:
        from terminal.research_council.states import plan_review

        updated = plan_review.run(_research_council_state_from_payload(state))
        return {"ok": True, "run_id": updated.run_id, "plan_reviews": [review.to_dict() for review in updated.plan_reviews]}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def run_critic_review(state: dict | Any) -> dict:
    """Run Research Council critic reviews against a serialized state."""
    try:
        from terminal.research_council.states import critic_review

        updated = critic_review.run(_research_council_state_from_payload(state))
        reviews = [[review.to_dict() for review in group] for group in updated.critic_reviews]
        return {"ok": True, "run_id": updated.run_id, "critic_reviews": reviews, "flags": updated.flags}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def apply_revision_round(state: dict | Any) -> dict:
    """Apply Research Council convergence/revision rules."""
    try:
        from terminal.research_council.states import revision

        updated = revision.run(_research_council_state_from_payload(state))
        return {"ok": True, "run_id": updated.run_id, "revision_history": [item.to_dict() for item in updated.revision_history], "flags": updated.flags}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def synthesize_council_decision(state: dict | Any) -> dict:
    """Synthesize the final Research Council decision from findings and reviews."""
    try:
        from terminal.research_council.states import synthesis

        updated = synthesis.run(_research_council_state_from_payload(state))
        return {"ok": True, "run_id": updated.run_id, "decision": updated.decision.to_dict() if updated.decision else None}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def render_research_council_report(
    state: dict | Any | None = None,
    run_id: str = "latest",
    output_format: str = "html",
) -> dict:
    """Render or locate a Research Council report."""
    try:
        if state is not None:
            from terminal.research_council.states import render_html as render_state

            updated = render_state.run(_research_council_state_from_payload(state))
            flags = updated.flags
            return {
                "ok": True,
                "run_id": updated.run_id,
                "report_path": flags.get("html_report_path") if output_format == "html" else flags.get("markdown_report_path"),
                "report_paths": {"markdown": flags.get("markdown_report_path"), "html": flags.get("html_report_path")},
            }
        resumed = resume_council_run(run_id=run_id)
        if not resumed.get("ok"):
            return resumed
        return {"ok": True, "run_id": resumed.get("run_id"), "report_path": resumed.get("report_path"), "report_paths": resumed.get("report_paths", {})}
    except Exception as exc:
        return {"ok": False, "error": str(exc), "run_id": run_id}


def persist_research_council_run(state: dict | Any) -> dict:
    """Persist Research Council run metadata for a serialized state."""
    try:
        from terminal.research_council.persistence import save_council_run_metadata

        council_state = _research_council_state_from_payload(state)
        saved = save_council_run_metadata(council_state)
        return {"ok": True, **saved}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}


def resume_council_run(run_id: str = "latest", include_debug: bool = False, output_format: str = "json") -> dict:
    """Load compact Research Council run metadata from PostgreSQL."""
    try:
        from terminal.research_council.persistence import connect

        with connect() as conn:
            with conn.cursor() as cur:
                if run_id == "latest":
                    cur.execute(
                        """
                        SELECT run_id, generated_at, report_path, final_label, council_status,
                               council_mode, horizon, risk_budget
                        FROM recommendation_reports.runs
                        ORDER BY generated_at DESC
                        LIMIT 1
                        """
                    )
                else:
                    cur.execute(
                        """
                        SELECT run_id, generated_at, report_path, final_label, council_status,
                               council_mode, horizon, risk_budget
                        FROM recommendation_reports.runs
                        WHERE run_id = %s
                        """,
                        (run_id,),
                    )
                row = cur.fetchone()
        if not row:
            return {"ok": False, "error": f"Research Council run not found: {run_id}", "run_id": run_id}
        values = list(row)
        result = {
            "ok": True,
            "run_id": values[0],
            "generated_at": str(values[1]),
            "report_path": values[2],
            "report_paths": {"markdown": values[2], "html": str(values[2]).replace(".md", ".html") if values[2] else None},
            "final_label": values[3],
            "stage": values[4],
            "mode": values[5],
            "horizon": values[6],
            "risk_budget": values[7],
        }
        if output_format == "json":
            result["export"] = dict(result)
        if include_debug:
            result["debug"] = {"source": "recommendation_reports.runs"}
        return result
    except Exception as exc:
        return {"ok": False, "error": str(exc), "run_id": run_id}


def _generate_report_tool(content: str, report_type: str = "research",
                          symbol: str = "", title: str = "",
                          output_format: str = "html", filename: str = "") -> dict:
    """Generate a styled report file from analysis content."""
    from terminal.reports import generate_report
    return generate_report(
        content=content,
        report_type=report_type,
        symbol=symbol,
        title=title or None,
        output_format=output_format,
        filename=filename or None,
    )


def run_recommendation_report(
    symbols: list[str] | None = None,
    indices: list[str] | None = None,
    sectors: list[str] | None = None,
    watchlist: list[str] | None = None,
    output_format: str = "html",
    top_n: int = 25,
    include_portfolio: bool = False,
    persist: bool = True,
) -> dict:
    """Generate a grounded recommendation report scoped by symbols, indices, or sectors."""
    from terminal.recommendation_report import RecommendationReportOptions, generate_recommendation_report

    opts = RecommendationReportOptions(
        output_format=output_format,
        top_n=int(top_n or 25),
        include_portfolio=bool(include_portfolio),
        watchlist=[str(symbol).upper().strip() for symbol in (watchlist or []) if str(symbol).strip()],
        symbols=[str(symbol).upper().strip() for symbol in (symbols or []) if str(symbol).strip()],
        indices=[str(index).upper().strip() for index in (indices or []) if str(index).strip()],
        sectors=[str(sector).strip() for sector in (sectors or []) if str(sector).strip()],
    )
    return generate_recommendation_report(options=opts, persist=bool(persist))


TOOL_REGISTRY.update({
    "build_research_evidence_pack": (
        build_research_evidence_pack,
        "Build a Research Council evidence pack from PostgreSQL-backed market, sector, stock, F&O, fundamentals, events, and report sources.",
        {
            "type": "object",
            "properties": {
                "mode": {"type": "string", "default": "market_council"},
                "universe_filter": {"type": "string", "default": "liquid"},
                "symbols": {"type": "array", "items": {"type": "string"}},
                "max_stock_candidates": {"type": "integer", "default": 50},
            },
            "required": [],
        },
    ),
    "run_research_council": (
        run_research_council,
        "Run the Research Council state machine for `/council` market, stock, strategy, intraday, and report-review objectives.",
        {
            "type": "object",
            "properties": {
                "objective": {"type": "string"},
                "mode": {"type": "string"},
                "symbols": {"type": "array", "items": {"type": "string"}},
                "horizon": {"type": "string"},
                "risk_budget": {"type": "string"},
                "report_path": {"type": "string"},
                "dry_run": {"type": "boolean", "default": False},
                "evidence_only": {"type": "boolean", "default": False},
                "output_format": {"type": "string", "enum": ["md", "html"], "default": "md"},
            },
            "required": ["objective"],
        },
    ),
    "run_data_steward_check": (
        run_data_steward_check,
        "Run Research Council data-steward freshness and universe-readiness checks.",
        {
            "type": "object",
            "properties": {"mode": {"type": "string", "default": "market_council"}},
            "required": [],
        },
    ),
    "compose_plan": (
        compose_plan,
        "Compose the deterministic Research Council evidence plan for an objective.",
        {
            "type": "object",
            "properties": {"objective": {"type": "string"}, "mode": {"type": "string", "default": "market_council"}},
            "required": ["objective"],
        },
    ),
    "execute_plan": (
        execute_plan,
        "Execute the latest Research Council plan from a serialized CouncilState.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "review_plan_execution": (
        review_plan_execution,
        "Review Research Council plan execution and decide whether more evidence is required.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "run_critic_review": (
        run_critic_review,
        "Run deterministic Research Council critics over a serialized CouncilState.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "apply_revision_round": (
        apply_revision_round,
        "Apply Research Council revision and convergence rules to a serialized CouncilState.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "synthesize_council_decision": (
        synthesize_council_decision,
        "Synthesize the final Research Council decision from findings, branch summaries, plan results, and critic reviews.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "render_research_council_report": (
        render_research_council_report,
        "Render a Research Council report from state or locate a persisted report by run_id.",
        {
            "type": "object",
            "properties": {
                "state": {"type": "object"},
                "run_id": {"type": "string", "default": "latest"},
                "output_format": {"type": "string", "enum": ["html", "md"], "default": "html"},
            },
            "required": [],
        },
    ),
    "persist_research_council_run": (
        persist_research_council_run,
        "Persist Research Council run metadata for a serialized CouncilState.",
        {"type": "object", "properties": {"state": {"type": "object"}}, "required": ["state"]},
    ),
    "resume_council_run": (
        resume_council_run,
        "Resume or inspect compact Research Council run metadata from PostgreSQL by run_id or latest.",
        {
            "type": "object",
            "properties": {
                "run_id": {"type": "string", "default": "latest"},
                "include_debug": {"type": "boolean", "default": False},
                "output_format": {"type": "string", "enum": ["json"], "default": "json"},
            },
            "required": [],
        },
    ),
    "list_generated_reports": (
        list_generated_reports,
        "List generated report artifacts with type, symbol, path, timestamp, and size metadata.",
        {
            "type": "object",
            "properties": {
                "project_root": {"type": "string"},
                "report_type": {"type": "string", "default": "any"},
                "limit": {"type": "integer", "default": 20},
            },
            "required": [],
        },
    ),
    "get_last_report": (
        get_last_report,
        "Return the last generated report path/context, or request clarification if none is remembered.",
        {
            "type": "object",
            "properties": {
                "last_report_path": {"type": "string"},
                "project_root": {"type": "string"},
            },
            "required": [],
        },
    ),
    "open_report": (
        open_report,
        "Open a generated report file by path and return a structured status message.",
        {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "project_root": {"type": "string"},
            },
            "required": ["path"],
        },
    ),
    "read_report": (
        read_report,
        "Read a generated Markdown/HTML/JSON/CSV report and return content plus report metadata.",
        {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "project_root": {"type": "string"},
                "max_chars": {"type": "integer", "default": 12000},
            },
            "required": ["path"],
        },
    ),
    "summarize_report": (
        summarize_report,
        "Summarize an existing report while preserving symbol, report type, recommendation, and source path.",
        {
            "type": "object",
            "properties": {
                "path": {"type": "string"},
                "project_root": {"type": "string"},
            },
            "required": ["path"],
        },
    ),
    "compare_reports": (
        compare_reports,
        "Compare two generated reports and highlight recommendation changes.",
        {
            "type": "object",
            "properties": {
                "first_path": {"type": "string"},
                "second_path": {"type": "string"},
                "project_root": {"type": "string"},
            },
            "required": ["first_path", "second_path"],
        },
    ),
    "run_recommendation_report": (
        run_recommendation_report,
        "Generate a grounded recommendation report scoped by stock symbols, indices, sectors, watchlist, or portfolio.",
        {
            "type": "object",
            "properties": {
                "symbols": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional NSE stock symbols to include, e.g. ['DIXON', 'DMART'].",
                },
                "indices": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional NSE index names to include, e.g. ['NIFTY 50', 'NIFTY BANK'].",
                },
                "sectors": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional sector names to include, e.g. ['IT', 'Capital Goods'].",
                },
                "watchlist": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional watchlist symbols to add as portfolio/watchlist evidence.",
                },
                "output_format": {
                    "type": "string",
                    "enum": ["html", "pdf", "md"],
                    "default": "html",
                },
                "top_n": {"type": "integer", "default": 25},
                "include_portfolio": {"type": "boolean", "default": False},
                "persist": {"type": "boolean", "default": True},
            },
            "required": [],
        },
    ),
    "get_postgres_health": (
        get_postgres_health,
        "Check PostgreSQL connectivity, DSN/socket details, required schemas/tables, and row counts.",
        {
            "type": "object",
            "properties": {"dsn": {"type": "string", "description": "Optional PostgreSQL DSN override"}},
            "required": [],
        },
    ),
    "ensure_postgres_schema": (
        ensure_postgres_schema,
        "Idempotently create core Agent Adda PostgreSQL schemas/tables required by runtime tools.",
        {
            "type": "object",
            "properties": {"dsn": {"type": "string", "description": "Optional PostgreSQL DSN override"}},
            "required": [],
        },
    ),
    "audit_postgres_coverage": (
        audit_postgres_coverage,
        "Audit PostgreSQL table existence and row-count coverage for market, intraday, scores, and report data.",
        {
            "type": "object",
            "properties": {"dsn": {"type": "string", "description": "Optional PostgreSQL DSN override"}},
            "required": [],
        },
    ),
    "load_historical_eod_to_postgres": (
        load_historical_eod_to_postgres,
        "Report or trigger historical EOD loading into PostgreSQL once the load orchestrator is wired.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "days": {"type": "integer", "default": 0},
                "dsn": {"type": "string"},
            },
            "required": [],
        },
    ),
    "load_intraday_ohlcv_to_postgres": (
        load_intraday_ohlcv_to_postgres,
        "Report or trigger intraday OHLCV loading into PostgreSQL once the load orchestrator is wired.",
        {
            "type": "object",
            "properties": {
                "symbol": {"type": "string"},
                "timeframe": {"type": "string", "default": "15m"},
                "dsn": {"type": "string"},
            },
            "required": [],
        },
    ),
    "get_data_source_manifest": (
        get_data_source_manifest,
        "Return Agent Adda's active data-source manifest, PostgreSQL primary-store policy, and fallback rules.",
        {"type": "object", "properties": {}, "required": []},
    ),
    "generate_report": (
        _generate_report_tool,
        ("Generate a formatted report file (HTML, PDF, or Markdown) from analysis content. "
         "Use this after completing any analysis to save results as a professional report. "
         "Supports report types: technical, fundamental, forensic, research, intraday, "
         "canslim, ric, sector. The report is saved to reports/generated/ directory."),
        {
            "type": "object",
            "properties": {
                "content": {
                    "type": "string",
                    "description": "The full analysis content in Markdown format to include in the report",
                },
                "report_type": {
                    "type": "string",
                    "enum": ["technical", "fundamental", "forensic", "research",
                             "intraday", "canslim", "ric", "sector"],
                    "description": "Type of report — determines styling and badge",
                    "default": "research",
                },
                "symbol": {
                    "type": "string",
                    "description": "Stock symbol or sector name for the report title",
                },
                "title": {
                    "type": "string",
                    "description": "Custom report title (optional — auto-generated from type+symbol if omitted)",
                },
                "output_format": {
                    "type": "string",
                    "enum": ["html", "pdf", "md"],
                    "description": "Output file format: html (styled), pdf (via weasyprint), or md (plain markdown)",
                    "default": "html",
                },
                "filename": {
                    "type": "string",
                    "description": "Custom filename without extension (optional — auto-generated if omitted)",
                },
            },
            "required": ["content"],
        },
    ),
})


def call_tool(name: str, args: dict) -> dict:
    """Execute a registered tool by name with given arguments.

    Defensively drops any kwargs not present in the target function's
    signature (unless the function accepts **kwargs). This guards against
    LLM-generated tool plans that pass extra/hallucinated kwargs
    (e.g. ``get_live_market_overview(timeframe=...)``) which would
    otherwise crash the executor with a TypeError.
    """
    import inspect

    if name not in TOOL_REGISTRY:
        return {"error": f"Unknown tool: {name}"}
    fn = TOOL_REGISTRY[name][0]
    safe_args = dict(args or {})
    try:
        sig = inspect.signature(fn)
        params = sig.parameters
        accepts_var_kw = any(
            p.kind == inspect.Parameter.VAR_KEYWORD for p in params.values()
        )
        if not accepts_var_kw:
            allowed = set(params.keys())
            unknown = [k for k in safe_args if k not in allowed]
            if unknown:
                safe_args = {k: v for k, v in safe_args.items() if k in allowed}
    except (TypeError, ValueError):
        # Builtins or C-extension callables may not expose a signature; fall
        # through and let the call surface its own error.
        pass
    try:
        return fn(**safe_args)
    except Exception as e:
        return {"error": str(e), "tool": name}


def openai_tool_schemas() -> list[dict]:
    """Return OpenAI-compatible tool schemas for all registered tools."""
    schemas = []
    for name, (_, description, params) in TOOL_REGISTRY.items():
        schemas.append({
            "type": "function",
            "function": {"name": name, "description": description, "parameters": params},
        })
    return schemas
